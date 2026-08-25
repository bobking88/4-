from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document

from update_formal_report_v2 import (
    FORMAL_REPORT,
    _add_heading_before,
    _add_paragraph_before,
    _add_picture_before,
    _add_table_before,
    replace_paragraph_text,
)


ROOT = Path(__file__).resolve().parents[1]
SUMMARY = ROOT / "outputs" / "business_metrics" / "rsg_hrgv" / "source_holdout" / "rsg_three_seed_summary.json"
PAIRED = ROOT / "outputs" / "business_metrics" / "rsg_hrgv" / "source_holdout" / "paired_rsg_complete_vs_hrgv_reference.json"
FIGURE = ROOT / "outputs" / "paper_figures_v2" / "fig_rsg_source_holdout.png"


def pct(value: float, std: float) -> str:
    return f"{100 * value:.2f}% ± {100 * std:.2f}%"


def metric(summary: dict, setting: str, key: str) -> str:
    values = summary[setting][key]
    return pct(float(values["mean"]), float(values["sample_std"]))


def update_report(input_path: Path, output_path: Path) -> Path:
    summary = json.loads(SUMMARY.read_text(encoding="utf-8"))
    paired = json.loads(PAIRED.read_text(encoding="utf-8"))
    document = Document(input_path)
    if any(p.text.strip() == "5.19 RSG-HRGV-Net 的摄影者留出泛化验证" for p in document.paragraphs):
        document.save(output_path)
        return output_path

    headings = {p.text.strip(): p for p in document.paragraphs if p.text.strip()}
    anchor = headings["5.19 实验结果的综合判断"]
    replace_paragraph_text(anchor, "5.20 实验结果的综合判断")
    _add_heading_before(document, anchor, "5.19 RSG-HRGV-Net 的摄影者留出泛化验证", 2)
    _add_paragraph_before(
        document,
        anchor,
        "固定随机划分已按重复图像和 Mindat 图片编号控制泄漏，但仍可能受摄影背景、布光与上传者风格影响。为检验后悔监督机制在此类来源变化下是否保持其预期方向，本节沿用摄影者留出清单：5,639张具有摄影者字段的图像按摄影者与重复图像组整体划分为训练/验证/测试集3,947/846/846张，三划分的摄影者与 split_group_id 均不交叉；缺失摄影者字段的2,890张图像不进入本实验。该实验是公开 Mindat 标本图片上的来源留出验证，不是独立工业现场或未知矿物（OOD）测试。",
    )
    _add_paragraph_before(
        document,
        anchor,
        "为避免使用试跑结果进行重复计数，种子20260728只用于检查执行与冻结流程，不进入统计推断；正式确认比较使用独立种子20260727、20260729和20260730。两组模型使用同一清单、同一 EfficientNet-B0 初始化、优化器、早停规则与耦合残差验证器，唯一差异为完整 RSG-HRGV 加入式（17）的局部梯度隔离软后悔门控监督。",
    )
    _add_table_before(
        document,
        anchor,
        ["设置", "Accuracy", "Macro F1", "目标召回", "含钛误入", "金属误入", "一对一错路由", "平均路由后悔"],
        [
            [
                "耦合残差 HRGV（参考）",
                metric(summary, "hrgv_reference", "accuracy"),
                metric(summary, "hrgv_reference", "macro_f1"),
                metric(summary, "hrgv_reference", "target_recall"),
                metric(summary, "hrgv_reference", "ti_to_target_intrusion_rate"),
                metric(summary, "hrgv_reference", "metallic_to_target_intrusion_rate"),
                metric(summary, "hrgv_reference", "one_right_gate_selection_accuracy"),
                metric(summary, "hrgv_reference", "mean_routing_regret_nll"),
            ],
            [
                "RSG-HRGV（完整）",
                metric(summary, "rsg_complete", "accuracy"),
                metric(summary, "rsg_complete", "macro_f1"),
                metric(summary, "rsg_complete", "target_recall"),
                metric(summary, "rsg_complete", "ti_to_target_intrusion_rate"),
                metric(summary, "rsg_complete", "metallic_to_target_intrusion_rate"),
                metric(summary, "rsg_complete", "one_right_gate_selection_accuracy"),
                metric(summary, "rsg_complete", "mean_routing_regret_nll"),
            ],
        ],
        [2.45, 1.45, 1.45, 1.55, 1.55, 1.55, 2.05, 1.65],
        "表 27 摄影者留出下 RSG-HRGV-Net 与 HRGV 的独立三随机种子比较",
    )
    _add_picture_before(
        document,
        anchor,
        FIGURE,
        16.0,
        "图 18 摄影者留出条件下 RSG-HRGV 的角色指标与路由机制比较",
    )
    routing = paired["routing_regret"]
    classification = paired["classification"]
    _add_paragraph_before(
        document,
        anchor,
        f"来源留出下，RSG 的平均 Macro F1 由{metric(summary, 'hrgv_reference', 'macro_f1')}变为{metric(summary, 'rsg_complete', 'macro_f1')}，目标召回由{metric(summary, 'hrgv_reference', 'target_recall')}变为{metric(summary, 'rsg_complete', 'target_recall')}；两者的成对簇 Bootstrap 区间均跨0。含钛干扰误入目标的均值下降，但区间仍跨0；金属光泽干扰误入目标的均值上升，且区间亦跨0。因此，该实验不能声称 RSG 在来源变化下稳定提高总体角色分类或所有业务风险指标。",
    )
    _add_paragraph_before(
        document,
        anchor,
        f"相反，直接对应命题6的路由机制指标在三个独立种子中方向一致：平均路由后悔由{metric(summary, 'hrgv_reference', 'mean_routing_regret_nll')}降至{metric(summary, 'rsg_complete', 'mean_routing_regret_nll')}，RSG−HRGV 的差异为{100 * float(routing['difference']):.2f}个百分点，95%成对簇 Bootstrap 区间为[{100 * float(routing['ci_low']):.2f}, {100 * float(routing['ci_high']):.2f}]个百分点，未跨0，方向性重采样概率为{float(routing['probability_favorable']):.4f}。同时，在恰有一位专家正确的样本中，门控选择正确专家的比例由{metric(summary, 'hrgv_reference', 'one_right_gate_selection_accuracy')}提高至{metric(summary, 'rsg_complete', 'one_right_gate_selection_accuracy')}。这为“后悔监督改善跨粒度证据路由”提供了比固定随机划分更严格、但仍限于公开标本来源的补充证据。",
    )
    _add_paragraph_before(
        document,
        anchor,
        f"上述判断依赖于路由后悔的预定义方向，而非事后挑选指标。作为边界，{classification['macro_f1']['label']}、{classification['target_recall']['label']}与两类误入率的区间均未给出稳定总体增益；论文应把该来源留出结果定位为机制泛化验证，不将其替代真实颗粒、矿石品位或工业回收率验证。",
    )

    replacements = {
        "表 27 本阶段方法贡献及其解释边界": "表 28 本阶段方法贡献及其解释边界",
        "表 28 当前局限性及处理原则": "表 29 当前局限性及处理原则",
    }
    for paragraph in document.paragraphs:
        if paragraph.text.strip() in replacements:
            replace_paragraph_text(paragraph, replacements[paragraph.text.strip()])

    old_last = next(
        p for p in document.paragraphs if p.text.strip().startswith("13. 本研究的理论结果建立在固定种类")
    )
    replace_paragraph_text(old_last, old_last.text.replace("13. 本研究", "14. 本研究", 1))
    _add_paragraph_before(
        document,
        old_last,
        "13. 摄影者留出实验在不交叉摄影者和重复图像组的条件下，使用独立三随机种子复现了 RSG 路由后悔降低的方向；但总体分类指标的区间仍跨0，且数据源仍为公开标本图片。因此，本研究将该结果作为跨来源机制证据，而不扩大为工业场景泛化结论。",
    )
    for paragraph in document.paragraphs:
        if "outputs/training/rsg_controlled/" in paragraph.text:
            replace_paragraph_text(
                paragraph,
                paragraph.text + " 摄影者留出确认集的清单、独立种子和汇总统计位于 outputs/training/rsg_source_holdout/、outputs/business_metrics/rsg_hrgv/source_holdout/ 和 docs/experiment_records/2026-08-24_rsg_source_holdout.md。",
            )
            break
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Integrate RSG source-held-out evidence into the formal report.")
    parser.add_argument("--input", type=Path, default=FORMAL_REPORT)
    parser.add_argument("--output", type=Path, default=FORMAL_REPORT)
    args = parser.parse_args()
    print(update_report(args.input.resolve(), args.output.resolve()))


if __name__ == "__main__":
    main()
