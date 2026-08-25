from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

from update_formal_report_v2 import FORMAL_REPORT, _set_run_font


ROOT = Path(__file__).resolve().parents[1]
ABLATION_DIR = ROOT / "outputs" / "business_metrics" / "rsg_hrgv" / "theory_ablation"


def _add_body(document: Document, text: str, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        _set_run_font(lead)
        rest = paragraph.add_run(text[len(bold_lead) :])
        _set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run)


def _percent(value: float, digits: int = 2) -> str:
    return f"{100 * value:.{digits}f}%"


def _load_comparison(configuration: str, ablation_dir: Path) -> dict[str, object]:
    path = ablation_dir / f"paired_{configuration}_vs_rsg_complete.json"
    with path.open(encoding="utf-8") as handle:
        return json.load(handle)


def _summary_by_configuration(ablation_dir: Path) -> dict[str, dict[str, object]]:
    with (ablation_dir / "rsg_theory_ablation_summary.json").open(encoding="utf-8") as handle:
        return json.load(handle)


def _add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run)
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.bold = True
    document.add_paragraph()


def update_report(
    input_path: Path,
    output_path: Path,
    ablation_dir: Path = ABLATION_DIR,
) -> Path:
    document = Document(input_path)
    heading = "附录 C RSG-HRGV 三随机种子理论消融"
    if any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.save(output_path)
        return output_path

    summaries = _summary_by_configuration(ablation_dir)
    hard = _load_comparison("rsg_hard_target", ablation_dir)
    unweighted = _load_comparison("rsg_unweighted", ablation_dir)
    coupled = _load_comparison("rsg_coupled_gate", ablation_dir)
    comparisons = {
        "硬目标": hard,
        "取消置信度加权": unweighted,
        "取消局部梯度隔离": coupled,
    }

    document.add_heading(heading, level=1)
    _add_body(
        document,
        "本附录使用与完整 RSG 相同的三随机种子、固定数据划分、EfficientNet-B0 主干、验证器耦合方式和训练预算，只替换一个后悔门控部件。各配置按图片 split_group_id 聚类、按随机种子重采样的两阶段 Bootstrap（2,000 次）与完整 RSG 进行逐图成对比较。该消融用于检验已定义的机制边界，不用于声称一般分类精度或工业分选性能。",
    )
    document.add_heading("C.1 受控配置与指标", level=2)
    _add_table(
        document,
        [
            ["配置", "仅改变的部件", "理论对应"],
            ["完整 RSG", "软目标 + 差距权重 + 后悔分支 stopgrad", "完整方法"],
            ["硬目标", "将 g*=sigmoid(Delta/Tr) 改为 I(Delta>=0)", "式(B-2)的连续软目标"],
            ["取消置信度加权", "将 w=tanh(|Delta|/Tw) 改为 w=1", "式(B-1)的差距风险代理"],
            ["取消局部梯度隔离", "允许 Lreg 经门控输入回传至主干/专家", "式(B-3)的局部梯度边界"],
        ],
    )
    _add_body(
        document,
        "主要机制指标为平均路由后悔（越低越好）和一对一错样本的门控选择正确率（越高越好）；同时报告 Accuracy、Macro F1、目标类召回率及两类误入率，以检查路由改善是否伴随不可接受的主任务风险。",
    )

    document.add_heading("C.2 三随机种子汇总", level=2)
    table_rows = [["配置", "Macro F1", "门控选择正确率", "平均路由后悔"]]
    for configuration, label in (
        ("rsg_complete", "完整 RSG"),
        ("rsg_hard_target", "硬目标"),
        ("rsg_unweighted", "取消置信度加权"),
        ("rsg_coupled_gate", "取消局部梯度隔离"),
    ):
        summary = summaries[configuration]
        macro = summary["macro_f1"]
        selection = summary["one_right_gate_selection_accuracy"]
        regret = summary["mean_routing_regret_nll"]
        table_rows.append(
            [
                label,
                f"{_percent(macro['mean'])} +/- {_percent(macro['sample_std'])}",
                f"{_percent(selection['mean'])} +/- {_percent(selection['sample_std'])}",
                f"{_percent(regret['mean'])} +/- {_percent(regret['sample_std'])}",
            ]
        )
    _add_table(document, table_rows)

    document.add_heading("C.3 机制检验与可支持的结论", level=2)
    hard_regret = hard["routing_regret"]
    _add_body(
        document,
        "（1）软目标：相对完整 RSG，硬目标的平均路由后悔增加"
        f"{_percent(hard_regret['difference'])}，95% Bootstrap 区间为"
        f"[{_percent(hard_regret['ci_low'])}, {_percent(hard_regret['ci_high'])}]，"
        f"有利于硬目标的重采样概率为{hard_regret['probability_favorable']:.3f}。"
        "该区间不跨零，说明在当前固定公开标本数据上，连续软目标优于将近似等价专家强制二值化；这与式(B-2)的动机一致。",
    )
    unweighted_regret = unweighted["routing_regret"]
    _add_body(
        document,
        "（2）置信度加权：取消加权相对完整 RSG 的平均路由后悔差异为"
        f"{_percent(unweighted_regret['difference'])}，区间为"
        f"[{_percent(unweighted_regret['ci_low'])}, {_percent(unweighted_regret['ci_high'])}]，跨越零。"
        "因此，w=tanh(|Delta|/Tw) 可作为由式(B-1)风险上界导出的有界设计，但本次三随机种子实验尚不能证明它在该数据集上带来稳定的额外性能增益；报告保留该负结果，不把它表述为已验证的普适优势。",
    )
    coupled_regret = coupled["routing_regret"]
    coupled_macro = coupled["classification"]["macro_f1"]
    _add_body(
        document,
        "（3）局部梯度隔离：取消隔离后平均路由后悔变化为"
        f"{_percent(coupled_regret['difference'])}，区间为"
        f"[{_percent(coupled_regret['ci_low'])}, {_percent(coupled_regret['ci_high'])}]；"
        f"Macro F1 变化为{_percent(coupled_macro['difference'])}，区间为"
        f"[{_percent(coupled_macro['ci_low'])}, {_percent(coupled_macro['ci_high'])}]。"
        "前者下降而后者呈下降趋势，表明让后悔监督直接改写专家证据可能以主任务稳定性为代价交换路由指标。式(B-3)是计算图的严格局部梯度结论；本实验只将其作为模块化优化边界的经验权衡，不声称隔离必然提高所有分类指标。",
    )
    document.add_heading("C.4 结论边界", level=2)
    _add_body(
        document,
        "本组结果支持的最强结论是：在当前固定公开标本图像数据与冻结训练协议下，RSG 的软门控目标可稳定降低已定义的路由后悔；后悔分支的差距加权和梯度隔离分别体现为风险建模与优化边界，但其单独的性能贡献存在数据依赖性。该结论不等同于矿物品位预测、真实产线回收率提升或未知矿物拒识能力。",
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Append RSG theory ablation results to the formal report.")
    parser.add_argument("--input", type=Path, default=FORMAL_REPORT)
    parser.add_argument("--output", type=Path, default=FORMAL_REPORT)
    parser.add_argument("--ablation-dir", type=Path, default=ABLATION_DIR)
    args = parser.parse_args()
    print(update_report(args.input.resolve(), args.output.resolve(), args.ablation_dir.resolve()))


if __name__ == "__main__":
    main()
