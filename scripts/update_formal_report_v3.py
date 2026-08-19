from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from update_formal_report_v2 import (
    FORMAL_REPORT,
    _add_caption_before,
    _add_heading_before,
    _add_paragraph_before,
    _add_picture_before,
    _add_table_before,
    replace_paragraph_text,
)


ROOT = Path(__file__).resolve().parents[1]
STATISTICS_JSON = (
    ROOT
    / "outputs"
    / "paper_experiments_v3"
    / "statistical_inference"
    / "paired_cluster_bootstrap_summary.json"
)
PAIRED_FIGURE = ROOT / "outputs" / "paper_figures_v3" / "fig_paired_cluster_effects.png"
ASSET_DIR = ROOT / "outputs" / "report_assets_v3"


def render_bootstrap_equation(path: Path) -> Path:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    equation = (
        r"$\widehat{\Delta}^{(b)}_m=\frac{1}{S}\sum_{s\in\mathcal{S}^{(b)}}"
        r"\left[m\!\left(H;\mathcal{D}^{*(b)}_s\right)-m\!\left(B;\mathcal{D}^{*(b)}_s\right)\right],"
        r"\quad \mathcal{D}^{*(b)}_s\sim\operatorname{SCB}_{\mathrm{role,group}}(\mathcal{D}_s),"
        r"\quad \mathrm{CI}_{0.95}=\left[Q_{0.025}(\widehat{\Delta}_m),Q_{0.975}(\widehat{\Delta}_m)\right]$"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    figure = plt.figure(figsize=(13.8, 1.2), dpi=220, facecolor="white")
    figure.text(0.5, 0.5, equation, ha="center", va="center", fontsize=17, color="#1F2937")
    figure.savefig(path, bbox_inches="tight", pad_inches=0.12, facecolor="white")
    plt.close(figure)
    return path


def _percentage(value: float) -> str:
    return f"{100.0 * value:.2f}%"


def _points(value: float, signed: bool = True) -> str:
    return f"{100.0 * value:+.2f}" if signed else f"{100.0 * value:.2f}"


def paired_summary_rows(payload: dict) -> list[list[str]]:
    summary = payload["summary"]
    row_specs = (
        ("accuracy", "Accuracy", "区间跨0"),
        ("macro_f1", "Macro F1", "区间跨0"),
        ("target_recall", "目标召回", "稳定改善"),
        ("target_miss_rate", "目标漏选率", "稳定下降"),
        ("ti_to_target_intrusion", "含钛干扰误入目标", "不利上升"),
        ("metallic_to_target_intrusion", "金属光泽干扰误入目标", "不利上升"),
    )
    rows = []
    for key, label, interpretation in row_specs:
        result = summary[key]
        rows.append(
            [
                label,
                _percentage(float(result["baseline_mean"])),
                _percentage(float(result["comparison_mean"])),
                _points(float(result["difference"])),
                f"[{_points(float(result['ci_low']), False)}, {_points(float(result['ci_high']), False)}]",
                interpretation,
            ]
        )
    return rows


def update_report(input_path: Path, output_path: Path) -> Path:
    payload = json.loads(STATISTICS_JSON.read_text(encoding="utf-8"))
    equation = render_bootstrap_equation(ASSET_DIR / "paired_cluster_bootstrap.png")
    document = Document(input_path)

    if any(p.text.strip() == "5.16 成对簇 Bootstrap 与配对统计推断" for p in document.paragraphs):
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path

    headings = {p.text.strip(): p for p in document.paragraphs if p.text.strip()}
    anchor = headings["5.16 实验结果的综合判断"]
    replace_paragraph_text(anchor, "5.17 实验结果的综合判断")

    _add_heading_before(document, anchor, "5.16 成对簇 Bootstrap 与配对统计推断", 2)
    _add_paragraph_before(
        document,
        anchor,
        "为避免仅凭三随机种子均值判断模型优劣，本研究在固定测试集上按 image_id 对齐普通 EfficientNet-B0 与分层模型的逐图预测。每次重复首先对三个训练随机种子有放回抽样，再在每个真实角色内部以 split_group_id 为不可拆分簇进行有放回抽样；同一簇内的原图、重复图或近重复图始终共同进入样本。设 B 和 H 分别表示基线与分层模型，m 为评价指标，则第 b 次模型效应及其百分位区间定义为：",
    )
    _add_picture_before(
        document,
        anchor,
        equation,
        16.0,
        "式（11） 两阶段成对簇 Bootstrap 的模型效应估计",
    )
    _add_paragraph_before(
        document,
        anchor,
        "正式实验采用10,000次重复，区间为2.5%和97.5%分位数。Accuracy、Macro F1 和目标召回以正差值为改善；目标漏选率及两类干扰误入目标的比例以负差值为改善。该方法同时反映当前三个训练种子和测试图像簇的抽样不确定性，但不外推到未观测的矿区、网站或工业现场。",
    )
    _add_table_before(
        document,
        anchor,
        ["指标", "基线均值", "分层模型均值", "差值/百分点", "95%簇区间/百分点", "判断"],
        paired_summary_rows(payload),
        [3.4, 2.4, 2.8, 2.7, 3.6, 2.5],
        "表 24 固定测试集上的成对统计推断",
    )
    _add_picture_before(
        document,
        anchor,
        PAIRED_FIGURE,
        15.8,
        "图 16 分层模型相对基线的成对效应及95%簇 Bootstrap 区间",
    )
    _add_paragraph_before(
        document,
        anchor,
        "结果显示，分层模型的目标召回提高9.16个百分点，95%簇 Bootstrap 区间为[5.18, 13.28]个百分点；目标漏选率相应下降9.16个百分点。与此同时，含钛干扰和金属光泽干扰误入目标类的比例分别上升3.19和4.73个百分点，且对应区间均未跨0。Accuracy仅提高0.62个百分点、Macro F1仅提高0.39个百分点，两者区间均跨0。因此，分层模型的主要效果是改变目标召回—难负样本误收之间的风险分配，而非证明总体性能全面提高。",
    )
    _add_paragraph_before(
        document,
        anchor,
        "作为逐图正确/错误变化的补充检验，每个随机种子分别进行双侧精确 McNemar 检验，并对三个 p 值作 Holm 校正。三个种子的校正后 p 值依次为0.094、0.110和0.110，三组经 Holm 校正的 p 值均大于0.05；同时第一个种子中基线独有正确135张、分层模型独有正确101张，而另两个种子的方向相反。这表明总体正确率变化存在明显种子异质性，不能将单个未校正结果解释为稳定优势。",
    )
    _add_paragraph_before(
        document,
        anchor,
        "该统计结果进一步支持本报告的风险敏感理论主线：层级辅助任务可以显著降低目标漏选，但不会自动控制含钛和金属光泽难负样本的误收。后续推理必须结合5.13节的概率校准、风险证书或类别相关阈值，在目标保留与干扰进入之间进行显式权衡。",
    )

    caption_replacements = {
        "表 24 本阶段方法贡献及其解释边界": "表 25 本阶段方法贡献及其解释边界",
        "表 25 当前局限性及处理原则": "表 26 当前局限性及处理原则",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in caption_replacements:
            replace_paragraph_text(paragraph, caption_replacements[text])

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("10. 本研究的理论结果建立在固定种类—角色映射"):
            replace_paragraph_text(
                paragraph,
                "11. 本研究的理论结果建立在固定种类—角色映射、同分布校准样本和公开矿物标本图像上；当前结论不涉及钒钛品位预测、工业回收率、真实XRF送检成本或现场生产线性能。阶段条件化选矿决策图保留为获得流程阶段与成本数据后的后续研究方向。",
            )
            conclusion = document.add_paragraph()
            conclusion_run = conclusion.add_run(
                "10. 成对簇 Bootstrap 表明，分层模型稳定提高目标召回并降低目标漏选，但同时增加两类困难负样本误入目标的比例；总体 Accuracy 和 Macro F1 的区间跨0，McNemar 校正结果也未形成跨种子的总体优势证据。因此，方法贡献应解释为可量化的风险重分配与后续选择性决策基础。"
            )
            from update_formal_report_v2 import _set_run_font

            _set_run_font(conclusion_run)
            conclusion.paragraph_format.line_spacing = 1.5
            paragraph._p.addprevious(conclusion._element)
            break

    for paragraph in document.paragraphs:
        if "层级风险汇总位于 outputs/theory_validation/hierarchy_consistency/" in paragraph.text:
            replace_paragraph_text(
                paragraph,
                "项目保留数据清单、采集与筛选脚本、数据质量审计记录、训练脚本、实验指标、混淆矩阵、逐图预测和图表源数据。层级风险汇总位于 outputs/theory_validation/hierarchy_consistency/；校准风险认证汇总位于 outputs/theory_validation/calibrated_selective_recognition/；来源留出与代理消融汇总位于 outputs/paper_experiments_v2/training_summaries/；成对簇 Bootstrap、McNemar 结果和重采样分布位于 outputs/paper_experiments_v3/statistical_inference/；对应论文图位于 outputs/paper_figures_v3/。",
            )
            break

    reference_anchor = next(p for p in document.paragraphs if p.text.strip().startswith("附录 A"))
    _add_paragraph_before(
        document,
        reference_anchor,
        "[20] Efron B, Tibshirani R J. An Introduction to the Bootstrap[M]. New York: Chapman & Hall, 1993.",
    )

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Add paired statistical inference to the formal report.")
    parser.add_argument("--input", type=Path, default=FORMAL_REPORT)
    parser.add_argument("--output", type=Path, default=FORMAL_REPORT)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    print(update_report(args.input.resolve(), args.output.resolve()))


if __name__ == "__main__":
    main()
