from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from update_formal_report_v2 import (
    FORMAL_REPORT,
    _add_heading_before,
    _add_paragraph_before,
    _add_picture_before,
    _add_table_before,
    _set_run_font,
    replace_paragraph_text,
)


ROOT = Path(__file__).resolve().parents[1]
SUMMARY_JSON = ROOT / "outputs" / "business_metrics" / "hrgv_network" / "hrgv_three_seed_summary.json"
NO_CONTRAST_JSON = (
    ROOT / "outputs" / "business_metrics" / "hrgv_network" / "paired_no_contrast_vs_residual_complete.json"
)
HIERARCHICAL_JSON = (
    ROOT
    / "outputs"
    / "paper_experiments_v4"
    / "statistical_inference"
    / "hrgv_vs_hierarchical"
    / "paired_cluster_bootstrap_summary.json"
)
ARCHITECTURE = ROOT / "outputs" / "paper_figures_v2" / "fig_hrgv_architecture.png"
EFFECT_FIGURE = ROOT / "outputs" / "paper_figures_v2" / "fig_hrgv_vs_hierarchical_effects.png"
ASSET_DIR = ROOT / "outputs" / "report_assets_v4"


def _pct(mean: float, std: float | None = None) -> str:
    if std is None:
        return f"{100.0 * mean:.2f}%"
    return f"{100.0 * mean:.2f}% ± {100.0 * std:.2f}%"


def _metric(summary: dict, setting: str, key: str) -> str:
    values = summary[setting][key]
    return _pct(float(values["mean"]), float(values["sample_std"]))


def hrgv_result_rows(summary: dict) -> list[list[str]]:
    settings = (
        ("residual_complete", "耦合残差 HRGV（主模型）"),
        ("equal_fusion", "等权融合"),
        ("no_contrast", "无角色感知对比约束"),
        ("decoupled_residual", "验证器梯度隔离"),
        ("gate_only", "仅门控、无验证器"),
        ("complete", "旧乘法验证"),
    )
    return [
        [
            label,
            _metric(summary, key, "accuracy"),
            _metric(summary, key, "macro_f1"),
            _metric(summary, key, "target_recall"),
            _metric(summary, key, "ti_to_target_intrusion_rate"),
            _metric(summary, key, "metallic_to_target_intrusion_rate"),
        ]
        for key, label in settings
    ]


def render_hrgv_equations(output_dir: Path) -> dict[str, Path]:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    equations = {
        "fusion": (
            r"$\mathbf{p}_d=\mathrm{softmax}(W_dh),\quad "
            r"\mathbf{p}_s=\mathrm{softmax}(W_sh),\quad "
            r"\mathbf{p}_m=A\mathbf{p}_s,\quad "
            r"\mathbf{p}_f=g\mathbf{p}_d+(1-g)\mathbf{p}_m$"
        ),
        "correction": (
            r"$c_h(v_h)=[\tau_h-v_h]_+/\tau_h,\quad "
            r"\eta=\beta_{Ti}c_{Ti}+\beta_{Met}c_{Met},\quad "
            r"q_T\propto p_{f,T}e^{-\eta},\quad q_j\propto p_{f,j}\ (j\ne T)$"
        ),
        "objective": (
            r"$\mathcal{L}=\mathcal{L}_{role}+\lambda_d\mathcal{L}_{direct}"
            r"+\lambda_s\mathcal{L}_{species}+\lambda_c\mathcal{L}_{KL}"
            r"+\lambda_v(\mathcal{L}_{Ti}+\mathcal{L}_{Met})"
            r"+\lambda_{con}\mathcal{L}_{contrast}$"
        ),
        "properties": (
            r"$-\log p_{f,y}\leq g[-\log p_{d,y}]+(1-g)[-\log p_{m,y}],\qquad "
            r"\frac{\partial}{\partial c_h}\log\frac{q_T}{q_j}=-\beta_h\leq0$"
        ),
        "selection": (
            r"$R_{act}(x)=\min_a\sum_y C(y,a)q(y\mid x),\qquad "
            r"\delta^*(x)=\mathrm{reject}\ \Longleftrightarrow\ R_{act}(x)>c_{rej},\qquad "
            r"\frac{\partial\mathcal{L}_v}{\partial\theta}=0\ \mathrm{under}\ \mathrm{stopgrad}(h)$"
        ),
    }
    output_dir.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    for name, equation in equations.items():
        figure = plt.figure(figsize=(13.6, 1.15), dpi=220, facecolor="white")
        figure.text(0.5, 0.5, equation, ha="center", va="center", fontsize=16.5, color="#1F2937")
        path = output_dir / f"hrgv_{name}.png"
        figure.savefig(path, bbox_inches="tight", pad_inches=0.12, facecolor="white")
        plt.close(figure)
        paths[name] = path
    return paths


def _remove_figure_and_caption(caption) -> None:
    caption_element = caption._p
    previous = caption_element.getprevious()
    parent = caption_element.getparent()
    parent.remove(caption_element)
    if previous is not None and previous.xpath(".//w:drawing"):
        previous.getparent().remove(previous)


def _insert_numbered_conclusion(document, anchor, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(run)
    paragraph.paragraph_format.line_spacing = 1.5
    anchor._p.addprevious(paragraph._element)


def update_report(input_path: Path, output_path: Path) -> Path:
    summary = json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))
    no_contrast = json.loads(NO_CONTRAST_JSON.read_text(encoding="utf-8"))
    hierarchical = json.loads(HIERARCHICAL_JSON.read_text(encoding="utf-8"))
    equations = render_hrgv_equations(ASSET_DIR)
    document = Document(input_path)

    if any(p.text.strip() == "4.6 层级风险门控验证网络（HRGV-Net）" for p in document.paragraphs):
        for paragraph in document.paragraphs:
            if paragraph.text.startswith("既有分层模型能够利用17类矿物种类监督") and "不等同于工业分选回收率验证" not in paragraph.text:
                replace_paragraph_text(
                    paragraph,
                    paragraph.text + " 本方法仍是公开矿物标本图像上的视觉代理任务，不等同于工业分选回收率验证。",
                )
                break
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path

    old_figure_caption = next(
        p for p in document.paragraphs if p.text.strip().startswith("图 5 理论感知的分层网络")
    )
    _remove_figure_and_caption(old_figure_caption)

    headings = {p.text.strip(): p for p in document.paragraphs if p.text.strip()}
    method_anchor = headings["4.6 固定映射下的角色可识别性命题"]
    replace_paragraph_text(method_anchor, "4.7 固定映射下的角色可识别性命题")
    replace_paragraph_text(headings["4.7 选择性识别规则"], "4.8 选择性识别规则")

    _add_heading_before(document, method_anchor, "4.6 层级风险门控验证网络（HRGV-Net）", 2)
    _add_paragraph_before(
        document,
        method_anchor,
        "既有分层模型能够利用17类矿物种类监督提升目标召回，但固定测试集的成对分析表明，两类困难负样本误入目标的比例也随之上升。为直接处理这一失效模式，本研究提出层级风险门控验证网络（HRGV-Net）。网络以 EfficientNet-B0 为共享主干，设置直接角色专家和矿物种类专家；后者经固定映射矩阵 A 聚合为角色后验。可靠性门控根据共享特征、两个专家的熵及 Jensen-Shannon 分歧进行逐图融合，随后由“目标—含钛干扰”和“目标—金属光泽干扰”两个验证器对目标后验进行中性区残差校正。本方法仍是公开矿物标本图像上的视觉代理任务，不等同于工业分选回收率验证。",
    )
    _add_picture_before(
        document,
        method_anchor,
        ARCHITECTURE,
        16.2,
        "图 5 HRGV-Net 的层级证据融合与困难负样本风险验证结构",
    )
    _add_paragraph_before(
        document,
        method_anchor,
        "设共享特征为 h，直接角色后验、17类种类后验及其角色映射分别为 pd、ps 和 pm，学习门控为 g，则层级证据融合为：",
    )
    _add_picture_before(document, method_anchor, equations["fusion"], 16.0, "式（12） HRGV-Net 的双专家可靠性融合")
    _add_paragraph_before(
        document,
        method_anchor,
        "对验证器 h∈{Ti,Met}，当验证分数不低于阈值 τh 时不修改融合后验；只有反对证据进入中性区以下时，才按强度 βh 单调降低目标相对赔率：",
    )
    _add_picture_before(document, method_anchor, equations["correction"], 16.0, "式（13） 中性区残差后验校正")
    _add_paragraph_before(document, method_anchor, "主任务、双专家、层级一致性、双验证器和角色感知对比约束的完整目标函数为：")
    _add_picture_before(document, method_anchor, equations["objective"], 16.0, "式（14） HRGV-Net 的完整训练目标")

    propositions = (
        (
            "命题1（凸融合对数损失上界）",
            "对任意真实角色 y 和 g∈[0,1]，融合后验的负对数损失不大于两个专家负对数损失的同权凸组合。证明直接来自 −log 在正数域上的凸性与 Jensen 不等式。该性质说明概率空间融合具有确定的上界，但不意味着学习门控必然显著优于等权融合。",
        ),
        (
            "命题2（中性区残差校正的恒等性与单调性）",
            "若两个验证器均给出支持或中性证据，则残差项为0且最终后验 q 与融合后验 pf 完全相同；当任一反对残差增加时，目标对任意非目标类的对数赔率以斜率 −βh 单调下降，非目标类之间的相对赔率保持不变。该性质由式（13）的后验比值直接得到。",
        ),
        (
            "命题3（验证式目标接纳的误收上界）",
            "若部署时要求基础分类器判为目标且相关困难负样本验证器同时通过，则最终误收事件是两事件的交集，其概率不超过任一组成事件的概率；只有在给定真实类别与来源后条件独立时才能进一步写成概率乘积。该结论适用于严格阈值接纳层，不能被扩大解释为软校正必然降低所有误收。",
        ),
        (
            "命题4（带拒识的贝叶斯最优选择性决策）",
            "给定校准后验 q、错误代价矩阵 C 和拒识代价 crej，应逐图比较所有直接动作的最小条件风险与 crej；当前者更大时拒识，否则执行最小风险动作。这是标准贝叶斯逐点风险最小化结果。本研究只验证了公开标本图像上的选择性识别，不宣称真实 XRF 送检成本已得到验证。",
        ),
        (
            "命题5（验证器梯度隔离）",
            "若验证器输入采用 stopgrad(h)，则验证器损失通过该路径对共享主干参数 θ 的梯度严格为0，而验证器自身参数仍正常更新。命题本身是计算图上的确定性性质；是否有利于分类性能必须通过消融实验判断。",
        ),
    )
    for title, body in propositions:
        _add_paragraph_before(document, method_anchor, title, bold_lead=title)
        _add_paragraph_before(document, method_anchor, body)
    _add_picture_before(document, method_anchor, equations["properties"], 16.0, "式（15） 凸融合上界与残差校正单调性")
    _add_picture_before(document, method_anchor, equations["selection"], 16.0, "式（16） 选择性风险规则与梯度隔离性质")
    _add_paragraph_before(
        document,
        method_anchor,
        "上述五条中，命题1、2和5是网络结构产生的确定性数学性质；命题3依赖严格阈值接纳规则及可选的条件独立假设；命题4是标准选择性决策性质。报告把这些性质与对应消融分开陈述，避免把数学恒等式、经验趋势和统计显著性混为一谈。",
    )

    headings = {p.text.strip(): p for p in document.paragraphs if p.text.strip()}
    result_anchor = headings["5.17 实验结果的综合判断"]
    replace_paragraph_text(result_anchor, "5.18 实验结果的综合判断")
    _add_heading_before(document, result_anchor, "5.17 HRGV-Net 三随机种子实验与组件验证", 2)
    _add_paragraph_before(
        document,
        result_anchor,
        "HRGV-Net 与所有组件消融沿用同一数据划分、相同三随机种子和统一评价口径。主模型采用耦合梯度、学习门控、中性区残差双验证器和角色感知对比约束。结果如下：",
    )
    _add_table_before(
        document,
        result_anchor,
        ["设置", "Accuracy", "Macro F1", "目标召回", "含钛误入目标", "金属光泽误入目标"],
        hrgv_result_rows(summary),
        [3.4, 2.2, 2.2, 2.2, 2.5, 2.7],
        "表 25 HRGV-Net 三随机种子组件消融",
    )

    bootstrap = no_contrast["bootstrap_summary"]
    accuracy = bootstrap["accuracy"]
    macro_f1 = bootstrap["macro_f1"]
    acc_gain = -100.0 * float(accuracy["difference"])
    acc_low = -100.0 * float(accuracy["ci_high"])
    acc_high = -100.0 * float(accuracy["ci_low"])
    f1_gain = -100.0 * float(macro_f1["difference"])
    f1_low = -100.0 * float(macro_f1["ci_high"])
    f1_high = -100.0 * float(macro_f1["ci_low"])
    _add_paragraph_before(
        document,
        result_anchor,
        f"角色感知对比约束给出了当前最明确的组件贡献：主模型相对无对比约束设置的 Accuracy 提高{acc_gain:.2f}个百分点，95%区间为[{acc_low:.2f}, {acc_high:.2f}]个百分点；Macro F1 提高{f1_gain:.2f}个百分点，95%区间为[{f1_low:.2f}, {f1_high:.2f}]个百分点，两个区间均未跨0。学习门控相对等权融合的 Macro F1 提高0.58个百分点，两类误入率分别下降0.61和0.76个百分点，但区间跨0，只能视为有利趋势。",
    )
    _add_paragraph_before(
        document,
        result_anchor,
        "完整 HRGV 相对 gate-only 将含钛干扰和金属光泽干扰误入目标的均值分别降低约0.62和0.19个百分点，但目标召回降低约1.99个百分点。两个验证器的 ROC-AUC 约为0.91至0.92，说明其学到可分信息；它们当前属于显式风险控制模块，不是无代价的精度增强模块。旧乘法校正与梯度隔离消融进一步表明，残差设计更稳定，而完全隔离验证器梯度会以总体性能和方差为代价提高目标召回。",
    )
    _add_picture_before(
        document,
        result_anchor,
        EFFECT_FIGURE,
        15.8,
        "图 17 HRGV-Net 相对既有分层模型的成对效应及95%簇 Bootstrap 区间",
    )

    comparison = hierarchical["summary"]
    _add_paragraph_before(
        document,
        result_anchor,
        "与既有分层 EfficientNet-B0 相比，HRGV 的 Accuracy 和 Macro F1 均值分别提高"
        f"{100.0 * float(comparison['accuracy']['difference']):.2f}和"
        f"{100.0 * float(comparison['macro_f1']['difference']):.2f}个百分点，两类困难负样本误入目标分别下降"
        f"{-100.0 * float(comparison['ti_to_target_intrusion']['difference']):.2f}和"
        f"{-100.0 * float(comparison['metallic_to_target_intrusion']['difference']):.2f}个百分点；"
        f"但目标召回下降{-100.0 * float(comparison['target_recall']['difference']):.2f}个百分点，且主要95%区间跨0。"
        "因此 HRGV 的证据应解释为更均衡、可审计的召回—误入风险重分配，不能解释为总体性能全面显著提高。",
    )

    caption_replacements = {
        "表 25 本阶段方法贡献及其解释边界": "表 26 本阶段方法贡献及其解释边界",
        "表 26 当前局限性及处理原则": "表 27 当前局限性及处理原则",
    }
    for paragraph in document.paragraphs:
        if paragraph.text.strip() in caption_replacements:
            replace_paragraph_text(paragraph, caption_replacements[paragraph.text.strip()])

    keyword_anchor = next(p for p in document.paragraphs if p.text.strip().startswith("关键词："))
    _add_paragraph_before(
        document,
        keyword_anchor,
        "在进一步的网络改进中，HRGV-Net 的三随机种子 Accuracy 为76.48%±0.23%、Macro F1 为74.59%±0.34%。角色感知对比约束的 Accuracy 与 Macro F1 成对簇区间均未跨0；门控、双验证器和残差校正则表现为可解释但多数尚未达到显著性的风险取舍。",
    )

    old_last_conclusion = next(
        p for p in document.paragraphs if p.text.strip().startswith("11. 本研究的理论结果建立在固定种类—角色映射")
    )
    replace_paragraph_text(
        old_last_conclusion,
        old_last_conclusion.text.replace("11. 本研究", "12. 本研究", 1),
    )
    _insert_numbered_conclusion(
        document,
        old_last_conclusion,
        "11. HRGV-Net 将层级双专家、逐图可靠性门控、两类困难负样本验证和中性区残差校正统一于同一计算图。三随机种子实验表明角色感知对比约束对 Accuracy 和 Macro F1 的改善具有未跨0的成对簇区间；其余组件主要体现召回与误入之间的可审计取舍，不能解释为全面显著优势。",
    )

    for paragraph in document.paragraphs:
        if "对应论文图位于 outputs/paper_figures_v3/" in paragraph.text:
            replace_paragraph_text(
                paragraph,
                paragraph.text
                + " HRGV 三随机种子汇总、消融差值、验证器 AUC 与成对统计位于 outputs/business_metrics/hrgv_network/；"
                + "HRGV 相对既有分层模型的成对统计位于 outputs/paper_experiments_v4/statistical_inference/hrgv_vs_hierarchical/；"
                + "网络结构图和效应图位于 outputs/paper_figures_v2/。",
            )
            break

    future_anchor = next(p for p in document.paragraphs if p.text.strip() == "7.3 中期研究方向")
    next_paragraph = future_anchor._p.getnext()
    if next_paragraph is not None:
        from docx.text.paragraph import Paragraph

        paragraph = Paragraph(next_paragraph, future_anchor._parent)
        replace_paragraph_text(
            paragraph,
            paragraph.text
            + " 阶段条件化选矿决策图仍保留为后续研究方向：只有获得真实流程阶段、矿物—动作映射、品位/回收率及送检成本后，才研究矿物种类与阶段共同决定处理动作的问题。",
        )

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Integrate HRGV-Net theory and experiments into the formal report.")
    parser.add_argument("--input", type=Path, default=FORMAL_REPORT)
    parser.add_argument("--output", type=Path, default=FORMAL_REPORT)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    print(update_report(args.input.resolve(), args.output.resolve()))


if __name__ == "__main__":
    main()
