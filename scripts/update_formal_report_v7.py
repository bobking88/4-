from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Cm

from update_formal_report_v2 import FORMAL_REPORT, _set_run_font


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "outputs" / "report_assets_v7"


def _add_body(document: Document, text: str, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        _set_run_font(lead)
        rest = paragraph.add_run(text[len(bold_lead):])
        _set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run)


def _add_formula(document: Document, path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(15.5))
    _add_body(document, caption)


def render_formulas() -> dict[str, Path]:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    formulas = {
        "bound": r"$0\leq -\log p_g+\log M\leq |g-g_o|\,|a-b|/\varepsilon$",
        "soft": r"$|\sigma(\Delta/T_r)-\mathbb{I}(\Delta\geq0)|\leq \exp(-|\Delta|/T_r)$",
        "gradient": r"$\partial\mathcal{L}_{reg}/\partial\theta_e=0,\qquad \partial\mathcal{L}_{reg}/\partial\phi\ne0$",
    }
    paths: dict[str, Path] = {}
    for name, formula in formulas.items():
        figure = plt.figure(figsize=(14, 1.0), dpi=220, facecolor="white")
        figure.text(0.5, 0.5, formula, ha="center", va="center", fontsize=17, color="#1F2937")
        path = ASSET_DIR / f"rsg_proof_{name}.png"
        figure.savefig(path, bbox_inches="tight", pad_inches=0.10, facecolor="white")
        plt.close(figure)
        paths[name] = path
    return paths


def update_report(input_path: Path, output_path: Path) -> Path:
    document = Document(input_path)
    if any(p.text.strip() == "附录 B RSG-HRGV 门控后悔理论证明" for p in document.paragraphs):
        document.save(output_path)
        return output_path
    formulas = render_formulas()
    document.add_heading("附录 B RSG-HRGV 门控后悔理论证明", level=1)
    _add_body(document, "本附录只证明当前 RSG-HRGV 计算图中后悔门控的三个局部结论。它们规定了优化对象与梯度范围，不推出网络在任意数据集上必然提高 Accuracy 或 Macro F1。令带真实角色标签 y 的图片 x 满足 a=pd(y|x)、b=pm(y|x)∈[ε,1]，ε>0；令 g∈[0,1] 为直接角色专家的门控权重，pg=ga+(1−g)b，M=max(a,b)，go=I(a≥b)。")
    document.add_heading("B.1 门控误差的路由后悔上界", level=2)
    _add_body(document, "定理 B.1（路由后悔上界）", bold_lead="定理 B.1（路由后悔上界）")
    _add_body(document, "在上述假设下，RSG 门控相对选择真实类概率更大专家的对数后悔满足下式。")
    _add_formula(document, formulas["bound"], "式（B-1） 门控误差的路由后悔上界")
    _add_body(document, "证明：pg 是 a 与 b 的凸组合，故 pg≤M，左端非负。若 a≥b，则 M−pg=(1−g)(a−b)=|g−go||a−b|；若 a<b，则 M−pg=g(b−a)=|g−go||a−b|。又因 −log t 在 [ε,1] 上的导数绝对值不超过1/ε，均值定理给出 −log pg+log M≤(M−pg)/ε，代入前述恒等式即得结论。证毕。")
    _add_body(document, "该上界说明，当两专家对真实类证据接近时，选错专家引起的后悔也受限；当差距显著时，门控误差的代价增大。因此 RSG 采用 tanh(|Δ|/Tw) 作为有界差距权重，而非仅凭经验放大所有样本。")
    document.add_heading("B.2 软门控目标的指数逼近", level=2)
    _add_body(document, "令 Δ=log(a/b)，g*=σ(Δ/Tr)，Tr>0。")
    _add_formula(document, formulas["soft"], "式（B-2） 软目标对硬最优门控的指数逼近")
    _add_body(document, "证明：Δ>0 时，go=1，且1−σ(Δ/Tr)=1/(1+exp(Δ/Tr))≤exp(−Δ/Tr)。Δ<0 时，go=0，且σ(Δ/Tr)=1/(1+exp(|Δ|/Tr))≤exp(−|Δ|/Tr)。证毕。该结论说明专家优劣明确时软目标接近硬选择，而专家接近时保留连续不确定性。")
    document.add_heading("B.3 后悔监督分支的局部梯度隔离", level=2)
    _add_body(document, "当前实现中，门控输入 z=[h,Hbar(pd),Hbar(pm),DJS(pd||pm)]、软目标 g* 与权重 w 都在后悔分支使用 stopgrad；门控为 g=σ(fφ(stopgrad(z)))，后悔损失为 Lreg=w BCE(g,g*)。")
    _add_formula(document, formulas["gradient"], "式（B-3） RSG 后悔分支的局部梯度隔离")
    _add_body(document, "证明：对产生 h、pd 与 pm 的任意共享主干或专家参数 θe，stopgrad 使 ∂z/∂θe、∂g*/∂θe 与∂w/∂θe 在 Lreg 分支均为0。链式法则遂使 ∂Lreg/∂θe=0。只要 0<g<1、w>0 且g≠g*，二元交叉熵关于g的导数非零，故门控参数 φ 仍可更新。证毕。")
    _add_body(document, "重要边界：式（B-3）只针对 Lreg 分支。角色、种类、一致性、验证器和对比损失仍依原计算图更新共享主干；因此局部梯度隔离不是“停止训练主干”。")
    document.add_heading("B.4 证伪性与实验对应", level=2)
    _add_body(document, "定理 B.1 对应平均路由后悔，定理 B.2 对应软/硬目标与一对一错路由选择率，定理 B.3 对应取消局部隔离的受控消融。固定测试划分中，RSG−HRGV 的平均路由后悔为−1.77个百分点，95%簇 Bootstrap 区间[−2.86, −0.69]；摄影者留出独立确认中为−3.53个百分点，区间[−4.75, −2.17]。这些结果支持已定义的路由机制，但不将理论性质扩大为工业分选、品位预测或整体分类精度的普适保证。")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Append formal RSG proofs to the technical report.")
    parser.add_argument("--input", type=Path, default=FORMAL_REPORT)
    parser.add_argument("--output", type=Path, default=FORMAL_REPORT)
    args = parser.parse_args()
    print(update_report(args.input.resolve(), args.output.resolve()))


if __name__ == "__main__":
    main()
