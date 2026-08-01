from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "结题"
OUTPUT = OUT_DIR / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（初稿）.docx"
ROUTE_IMAGE = OUT_DIR / "技术报告_配图_总体技术路线.png"
FIG_DIR = ROOT / "outputs" / "paper_figures_v1"
ROLE_IDENTIFIABILITY_JSON = ROOT / "outputs" / "theory_validation" / "role_identifiability" / "role_identifiability_summary.json"
SELECTIVE_RECOGNITION_JSON = ROOT / "outputs" / "theory_validation" / "selective_recognition" / "selective_recognition_summary.json"
COMPONENT_ABLATION_JSON = ROOT / "outputs" / "business_metrics" / "hierarchical_component_ablation" / "hierarchical_component_ablation_summary.json"
SELECTIVE_RECOGNITION_FIGURE = FIG_DIR / "fig9_selective_recognition.png"
THEORY_ARCHITECTURE_FIGURE = FIG_DIR / "fig10_theory_aware_hierarchical_architecture_cn.png"

FONT = "宋体"
HEADING_FONT = "黑体"
FONT_FILE = Path(r"C:\Windows\Fonts\msyh.ttc")


def load_theory_report_inputs():
    paths = {
        "role_identifiability": ROLE_IDENTIFIABILITY_JSON,
        "selective_recognition": SELECTIVE_RECOGNITION_JSON,
        "component_ablation": COMPONENT_ABLATION_JSON,
    }
    inputs = {
        name: json.loads(path.read_text(encoding="utf-8"))
        for name, path in paths.items()
    }
    if inputs["role_identifiability"]["validation_type"] != "controlled logical validation":
        raise ValueError("Role-identifiability evidence must be controlled logical validation.")
    if not inputs["selective_recognition"]["fixed_split_consumed_unchanged"]:
        raise ValueError("Selective-recognition evidence must use the unchanged fixed split.")
    if not isinstance(inputs["component_ablation"], list) or len(inputs["component_ablation"]) != 3:
        raise ValueError("Expected three committed component-ablation settings.")
    return inputs


def theory_equation_specs():
    return {
        "aggregation": r"$\widetilde{\mathbf{p}}_r = \mathbf{A}\mathbf{p}_s$",
        "joint_loss": (
            r"$\mathcal{L}=\mathcal{L}_{\mathrm{role}}"
            r"+\alpha\mathcal{L}_{\mathrm{species}}"
            r"+\beta\mathcal{L}_{\mathrm{cons}}"
            r"+\gamma\mathcal{L}_{\mathrm{binary}}"
            r"+\eta\mathcal{L}_{\mathrm{hard}}$"
        ),
        "consistency": (
            r"$\mathcal{L}_{\mathrm{cons}}="
            r"D_{\mathrm{KL}}\!\left(\mathbf{p}_r\,\Vert\,\widetilde{\mathbf{p}}_r\right)$"
        ),
        "confidence": r"$q(x)=\max_{r\in\mathcal{R}}p_r(r\mid x)$",
        "coverage": (
            r"$\mathrm{Coverage}(\tau)=\frac{1}{N}"
            r"\sum_{i=1}^{N}\mathbf{1}[q(x_i)\geq\tau]$"
        ),
        "risk": (
            r"$\mathrm{Risk}(\tau)="
            r"\frac{\sum_i\mathbf{1}[q(x_i)\geq\tau]\mathbf{1}[\hat{y}_i\neq y_i]}"
            r"{\sum_i\mathbf{1}[q(x_i)\geq\tau]}$"
        ),
    }


def _mathtext_python_candidates():
    candidates = []
    configured = os.environ.get("MATH_TEXT_PYTHON")
    if configured:
        candidates.append(Path(configured))
    candidates.extend([
        Path(sys.executable),
        ROOT / ".venv-training" / "Scripts" / "python.exe",
        ROOT.parent.parent / ".venv-training" / "Scripts" / "python.exe",
    ])
    for path_entry in os.environ.get("PATH", "").split(os.pathsep):
        if path_entry:
            candidates.append(Path(path_entry) / "python.exe")
    fallback = shutil.which("python")
    if fallback:
        candidates.append(Path(fallback))
    unique = []
    for candidate in candidates:
        try:
            if not candidate.exists():
                continue
            resolved = candidate.resolve()
        except OSError:
            continue
        if resolved not in unique:
            unique.append(resolved)
    return unique


def find_mathtext_python():
    for candidate in _mathtext_python_candidates():
        probe = subprocess.run(
            [str(candidate), "-c", "import matplotlib"],
            capture_output=True,
            text=True,
            check=False,
        )
        if probe.returncode == 0:
            return candidate
    raise RuntimeError(
        "Matplotlib is required to render report equations. Set MATH_TEXT_PYTHON "
        "to a Python interpreter with matplotlib installed."
    )


def render_theory_equations(output_dir):
    output_dir.mkdir(parents=True, exist_ok=True)
    manifest_path = output_dir / "equations.json"
    manifest_path.write_text(json.dumps(theory_equation_specs()), encoding="utf-8")
    renderer = r'''
import json
import sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

manifest = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
output_dir = Path(sys.argv[2])
for name, equation in manifest.items():
    fig = plt.figure(figsize=(11.5, 1.0), dpi=220, facecolor="white")
    fig.text(0.5, 0.5, equation, ha="center", va="center", fontsize=19, color="#1F1F1F")
    fig.savefig(output_dir / f"{name}.png", bbox_inches="tight", pad_inches=0.12, facecolor="white")
    plt.close(fig)
'''
    subprocess.run(
        [str(find_mathtext_python()), "-c", renderer, str(manifest_path), str(output_dir)],
        check=True,
    )
    return {name: output_dir / f"{name}.png" for name in theory_equation_specs()}


def format_percent(mean, sample_std):
    return f"{100 * mean:.2f} ± {100 * sample_std:.2f}"


def set_run_font(run, name=FONT, size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def table_border(table, color="808080", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        elem = borders.find(tag)
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_table_layout_fixed(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_paragraph_format(paragraph, before=0, after=6, line=1.5, first_indent=0, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_indent:
        fmt.first_line_indent = Cm(first_indent)
    if align is not None:
        paragraph.alignment = align


def add_text(paragraph, text, size=12, bold=False, font=FONT, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold, color)
    return run


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    set_paragraph_format(p, after=5, line=1.5, first_indent=0.74 if indent else 0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_text(p, text, 12)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_format(p, before=16, after=8, line=1.2)
        add_text(p, text, 16, True, HEADING_FONT, (31, 78, 120))
    elif level == 2:
        set_paragraph_format(p, before=12, after=6, line=1.2)
        add_text(p, text, 14, True, HEADING_FONT, (31, 78, 120))
    else:
        set_paragraph_format(p, before=8, after=4, line=1.2)
        add_text(p, text, 12, True, HEADING_FONT, (31, 78, 120))
    set_keep_with_next(p)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=4, after=8, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, text, 10.5, False, FONT)
    return p


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(table)
    table_border(table, "B7C7D6", "8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F7FA")
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    set_paragraph_format(p, after=0, line=1.35)
    add_text(p, label, 11, True, HEADING_FONT, (31, 78, 120))
    add_text(p, text, 11)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths_cm, font_size=10.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(table)
    table_border(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths_cm[i])
        set_cell_margins(cell)
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_paragraph_format(p, after=0, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, header, font_size, True, HEADING_FONT)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths_cm[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_paragraph_format(p, after=0, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            add_text(p, str(value), font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_figure(doc, image_path, width_cm, caption):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=4, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_equation(doc, image_path, width_cm=14.0):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=2, after=5, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))


def make_route_image(path):
    canvas = Image.new("RGB", (1800, 700), "white")
    draw = ImageDraw.Draw(canvas)
    font_big = ImageFont.truetype(str(FONT_FILE), 32)
    font_small = ImageFont.truetype(str(FONT_FILE), 22)
    title_font = ImageFont.truetype(str(FONT_FILE), 42)
    draw.text((580, 35), "钒钛矿相关矿物图像识别总体技术路线", font=title_font, fill="#17365D")
    boxes = [
        (80, 220, 390, 510, "公开矿物\n图像采集\nMindat 等公开页面\n记录编号与页面信息", "#EAF2F8"),
        (505, 220, 815, 510, "数据质控\n与标注\n筛选、混合样本隔离\n重复检测与分层复核", "#EEF5E9"),
        (930, 220, 1240, 510, "固定数据\n划分\n图片编号与近重复组分组\n训练、验证、测试隔离", "#FFF2CC"),
        (1355, 220, 1665, 510, "模型训练\n与评价\nResNet50 与 EfficientNet-B0\n三随机种子与错误分析", "#FCE4D6"),
    ]
    for x1, y1, x2, y2, text, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=26, fill=fill, outline="#5B9BD5", width=4)
        lines = text.split("\n")
        y = y1 + 32
        for i, line in enumerate(lines):
            f = font_big if i == 0 else font_small
            bbox = draw.textbbox((0, 0), line, font=f)
            draw.text(((x1+x2-(bbox[2]-bbox[0]))/2, y), line, font=f, fill="#1F1F1F")
            y += 46 if i < 2 else 38
    for x1, _, x2, _, _, _ in boxes[:-1]:
        start = x2 + 10
        end = x2 + 100
        draw.line((start, 365, end, 365), fill="#5B9BD5", width=7)
        draw.polygon([(end, 365), (end-20, 351), (end-20, 379)], fill="#5B9BD5")
    draw.rounded_rectangle((345, 565, 1455, 640), radius=18, fill="#F4F7FA", outline="#A6A6A6", width=2)
    foot = "输出：可追溯四分类数据集、训练基线、消融结果、混淆矩阵与后续真实矿石图像验证基础"
    bbox = draw.textbbox((0, 0), foot, font=font_small)
    draw.text(((1800-(bbox[2]-bbox[0]))/2, 584), foot, font=font_small, fill="#1F1F1F")
    canvas.save(path)


def set_document_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)
    section.header_distance = Cm(1.3)
    section.footer_distance = Cm(1.3)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.5

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(header, after=0, line=1.0)
    add_text(header, "钒钛关键战略材料四川省重点实验室开放项目技术报告", 9.5, False, FONT, (90, 90, 90))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(footer, after=0, line=1.0)
    add_text(footer, "技术报告初稿  |  2026 年 7 月", 9, False, FONT, (90, 90, 90))


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    theory_inputs = load_theory_report_inputs()
    role_validation = theory_inputs["role_identifiability"]
    selective_validation = theory_inputs["selective_recognition"]
    component_ablation = theory_inputs["component_ablation"]
    equation_temp = tempfile.TemporaryDirectory(prefix="technical_report_equations_")
    equations = render_theory_equations(Path(equation_temp.name))
    make_route_image(ROUTE_IMAGE)
    doc = Document()
    set_document_styles(doc)

    # Cover
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_format(p, after=20, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "钒钛关键战略材料四川省重点实验室", 20, True, HEADING_FONT, (31, 78, 120))
    p = doc.add_paragraph()
    set_paragraph_format(p, after=40, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "开放项目技术报告", 26, True, HEADING_FONT, (31, 78, 120))
    p = doc.add_paragraph()
    set_paragraph_format(p, after=50, line=1.35, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "基于深度学习的钒钛矿相关矿物图像识别方法研究", 22, True, HEADING_FONT)
    cover_rows = [
        ("项目编号", "[待填]"),
        ("项目承担单位", "[待填]"),
        ("项目负责人", "[待填]"),
        ("报告完成日期", "2026 年 7 月"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(table)
    table_border(table, "B7C7D6", "8")
    for key, val in cover_rows:
        cells = table.add_row().cells
        for i, value in enumerate((key, val)):
            set_cell_width(cells[i], 4.0 if i == 0 else 8.0)
            set_cell_margins(cells[i], 160, 180, 160, 180)
            if i == 0:
                set_cell_shading(cells[i], "E8EEF5")
            p = cells[i].paragraphs[0]
            set_paragraph_format(p, after=0, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(p, value, 12, i == 0, HEADING_FONT if i == 0 else FONT)
    doc.add_page_break()

    # Abstract
    add_heading(doc, "摘  要", 1)
    add_body(doc, "针对钒钛矿相关组分矿物在公开矿物标本图像中存在来源异构、伴生矿物复杂和困难负样本易混淆等问题，本文开展了基于深度学习的四分类图像识别方法研究。以公开矿物数据库中的矿物标本照片为主要来源，建立了包含目标矿物、含钛干扰矿物、脉石/废石和金属光泽干扰矿物四类样本的数据集，并为图像保留矿物标签、Mindat 图片编号、来源页面及筛选决策等元数据。通过主体清晰度筛选、混合样本隔离、重复图像检测和分层复核，形成了可追溯的固定训练、验证与测试清单。")
    role_consistent = role_validation["summary"]["role_consistent"]
    role_conflict = role_validation["summary"]["role_conflict"]
    threshold_090 = next(row for row in selective_validation["threshold_summary"] if row["threshold"] == 0.9)
    add_body(doc, f"在 8,529 张图像上，采用 ImageNet 预训练的 ResNet50 和 EfficientNet-B0 进行三随机种子重复训练。实验表明，EfficientNet-B0 的平均 Macro F1 为 72.87%±1.21%，ResNet50 为 72.06%±0.92%。进一步以 17 类矿物种类和四类选矿角色的固定映射为基础，对共享视觉特征上的角色头、种类头、目标代理头和投影头进行理论化分层建模，并使用种类概率到角色概率的聚合与 KL 一致性约束。受控逻辑条件验证中，{role_consistent['row_count']} 个角色一致候选集的角色唯一可识别率为 {100 * role_consistent['role_unique_rate']:.2f}%，{role_conflict['row_count']} 个角色冲突候选集为 {100 * role_conflict['role_unique_rate']:.2f}%；该实验不使用图像标签推断或模型预测。固定测试划分上的选择性识别结果显示，阈值 0.90 时三随机种子平均覆盖率为 {100 * threshold_090['mean']['coverage']:.2f}%，保留样本错误率为 {100 * threshold_090['mean']['risk']:.2f}%。这些结果仅支持理论形式化、受控逻辑验证和固定划分上的覆盖率—风险取舍，不构成工业分选、品位/回收率、真实 XRF 成本最优、外部验证或统计显著性结论。")
    p = doc.add_paragraph()
    set_paragraph_format(p, before=4, after=12, line=1.35)
    add_text(p, "关键词：", 12, True, HEADING_FONT)
    add_text(p, "钒钛矿；矿物图像识别；深度学习；数据集构建；困难负样本；分层一致性学习", 12)

    add_heading(doc, "目  录", 1)
    toc_items = [
        "1 研究背景与目标", "2 总体技术路线", "3 数据集构建与质量控制", "4 模型与实验设计",
        "5 实验结果与讨论", "6 主要创新与阶段性成果", "7 局限性与后续研究", "8 结论", "参考文献", "附录 A 数据字段与复现说明",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        set_paragraph_format(p, after=3, line=1.3, first_indent=0)
        add_text(p, item, 12)
    doc.add_page_break()

    add_heading(doc, "1 研究背景与目标", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_body(doc, "钒、钛、铁等战略金属的高效利用依赖于对矿石组分、伴生矿物和脉石的有效识别。传统矿物鉴定通常需要显微镜、X 射线荧光光谱、X 射线衍射或化学分析等手段，具有可靠性高但设备依赖强、现场快速初筛成本较高的特点。机器视觉和深度学习可以从颜色、光泽、晶体形态、粒度与基质等可见外观线索中提取特征，为矿物识别与预选提供辅助信息。")
    add_body(doc, "与直接使用工业皮带图像不同，公开矿物标本照片的来源、拍摄背景、尺度和伴生关系更复杂，但也提供了较丰富的类内变化。利用此类数据建立结构清晰、来源可追溯、划分受控的识别基线，可作为后续真实矿石颗粒图像、光谱信息及工艺变量融合建模之前的研究基础。")
    add_heading(doc, "1.2 研究范围与定位", 2)
    add_note(doc, "研究定位：", "本报告研究的是公开矿物标本图像条件下的钒钛矿相关组分矿物视觉识别方法验证。模型输出反映可见外观类别，不用于直接预测钒、钛、铁品位、回收率或工业生产线分选效率。")
    add_heading(doc, "1.3 研究目标", 2)
    add_table(doc, ["序号", "研究目标", "完成方式"], [
        ("1", "建立可追溯的钒钛矿相关矿物图像数据资源", "保留来源页面、矿物标签、图片编号、筛选决策与固定清单"),
        ("2", "构建面向分选风险的四分类任务", "将非目标样本细分为含钛干扰、脉石/废石和金属光泽困难干扰"),
        ("3", "建立深度学习识别基线并评价稳定性", "采用 ResNet50、EfficientNet-B0，三随机种子训练与固定测试集评价"),
        ("4", "分析模型失效模式并提出改进方向", "混淆矩阵、错分对统计、三随机种子 Focal Loss 消融、角色感知学习和分层模型组件消融"),
    ], [1.1, 5.0, 8.0], 10.5)

    add_heading(doc, "2 总体技术路线", 1)
    add_body(doc, "本研究按照“公开图像采集—质量筛选与标注—固定划分—深度模型训练—多维评价—错误诊断”的流程开展。数据层面强调来源与筛选记录，训练层面强调固定划分和多随机种子重复，分析层面强调困难类别与具体错分方向，而非仅报告总体准确率。")
    add_figure(doc, ROUTE_IMAGE, 16.0, "图 1 本研究总体技术路线")
    add_heading(doc, "2.1 任务定义", 2)
    add_body(doc, "为避免将所有非目标矿物简单合并为同一类别，本研究设定四分类任务：目标矿物（magnetite、ilmenite、titanomagnetite 等相关组分的视觉代理）、含钛干扰矿物、脉石/废石和金属光泽困难干扰矿物。该设计使模型既需区分目标与非目标，也需面对含钛、暗色和金属光泽等更接近目标类的干扰条件。")

    add_heading(doc, "3 数据集构建与质量控制", 1)
    add_heading(doc, "3.1 数据来源与元数据", 2)
    add_body(doc, "图像主要来自 Mindat 等公开矿物信息页面。每张纳入管理的图像记录原始文件名、矿物标签、Mindat 图片编号、详情页链接、页面标题和初步筛选结论等字段。该方式保证后续可回溯图像来源与标签依据，同时避免将未经核验的网页抓取结果直接作为训练样本。")
    add_heading(doc, "3.2 类别体系与样本规模", 2)
    add_table(doc, ["任务标签", "样本数", "主要子类", "设置目的"], [
        ("目标矿物", "1,661", "magnetite 1,316；ilmenite 322；titanomagnetite 23", "模拟钒钛矿相关有用组分的视觉代理"),
        ("含钛干扰矿物", "3,268", "rutile、anatase、perovskite、titanite 等", "避免将所有含钛矿物误判为目标"),
        ("脉石/废石", "2,427", "fluorapatite、vesuvianite、quartz、feldspar、calcite、pyroxene 等", "模拟分选中的脉石与废石干扰"),
        ("金属光泽干扰", "1,173", "pyrite、hematite、goethite、chalcopyrite 等", "检验模型是否仅凭暗色或反光特征误判"),
        ("合计", "8,529", "—", "固定版本 dataset_final_v1"),
    ], [2.8, 1.5, 6.5, 3.3], 9.5)
    add_heading(doc, "3.3 图像筛选与分层复核", 2)
    add_body(doc, "预筛阶段保留矿物主体占画面比例较高、颜色/光泽/块状或晶体形态可辨、分辨率满足识别需要且页面标签清晰的图片；剔除主体不明、多矿物混杂、严重模糊、过暗/过曝、纯显微照片及标尺、硬币、人手占比过大的图片。标题出现多个矿物且主体不清晰的样本先进入 mixed_uncertain 队列，不直接作为单标签正样本。")
    add_body(doc, "数据质控采用分层抽检与复核记录机制。共建立 823 张复核队列，其中 789 张保留、32 张排除、2 张标记为 needs_expert 并不进入最终训练清单。完全重复图像、近重复图像和跨标签冲突样本均在数据划分前进行审查。复核结论已形成；为保证项目归档可追溯性，复核人、复核日期及 3 条缺失的排除原因仍需由实际复核人员补录到原始复核队列。")
    add_heading(doc, "3.4 固定划分与泄漏控制", 2)
    add_table(doc, ["集合", "图像数", "比例", "用途"], [
        ("训练集", "5,961", "69.89%", "模型参数学习与数据增强"),
        ("验证集", "1,284", "15.05%", "早停与超参数观察"),
        ("测试集", "1,284", "15.05%", "固定的最终性能评价"),
        ("合计", "8,529", "100.00%", "固定版本 dataset_final_v1"),
    ], [3.0, 2.2, 2.2, 6.7], 10.5)
    add_body(doc, "划分时以 Mindat 图片编号和近重复图像组作为基本单元，确保同一来源图片或被识别为高度相近的图片不跨训练、验证和测试集合。该策略减少了同图或近同图泄漏导致的性能虚高风险。")

    add_heading(doc, "4 模型与实验设计", 1)
    add_heading(doc, "4.1 对比模型", 2)
    add_body(doc, "选用 ResNet50 与 EfficientNet-B0 作为对比基线。二者均加载 ImageNet 预训练权重，以适应本研究公开标本图像数据量有限、成像风格多样的特点。ResNet50 用于提供经典残差网络基线，EfficientNet-B0 用于考察复合缩放结构在该任务上的表现。")
    add_heading(doc, "4.2 训练设置", 2)
    add_table(doc, ["项目", "设置"], [
        ("输入与预训练", "ImageNet 预训练权重及相应标准化"),
        ("优化器", "AdamW"),
        ("类别不平衡处理", "按类别逆频率设置损失权重"),
        ("数据增强", "随机裁剪、水平翻转、旋转和颜色扰动等训练期增强"),
        ("早停", "以验证集 Macro F1 作为主要监控指标"),
        ("重复实验", "每种模型使用 20260727、20260728、20260729 三个随机种子独立训练"),
        ("评价指标", "Accuracy、Macro Precision、Macro Recall、Macro F1 与各类别 Recall"),
    ], [4.0, 10.1], 10.5)
    add_heading(doc, "4.3 Focal Loss 三随机种子消融", 2)
    add_body(doc, "在 EfficientNet-B0 的受控条件下，将加权交叉熵替换为类别加权 Focal Loss（gamma=2.0），其余数据版本、划分、网络、增强、优化器和训练设置保持一致。两个设置均使用 20260727、20260728 和 20260729 三个随机种子独立训练。该实验用于判断统一难例强调机制能否稳定改善四类任务及目标代理风险指标。")
    add_heading(doc, "4.4 角色感知困难负样本学习", 2)
    add_body(doc, "针对基线中目标矿物与含钛干扰、金属光泽干扰之间的主要错分方向，构建角色感知困难负样本模型。模型以 EfficientNet-B0 为共享特征主干，保留四分类主任务，同时增加目标代理/非目标代理二分类辅助头；训练阶段仅对“目标矿物—含钛干扰”和“目标矿物—金属光泽干扰”施加表示分离约束。总损失由四分类加权交叉熵、二分类辅助损失和关系感知对比损失组成。该设计不依赖虚构的品位、回收率或流程成本，将改进目标限定为降低公开标本图像条件下的关键视觉混淆。")
    add_heading(doc, "4.5 理论模型与符号定义", 2)
    add_body(doc, "在现有四分类角色标签之外，数据清单还保留了 17 类细粒度 mineral_label。理论感知分层模型以 EfficientNet-B0 产生共享视觉特征 h，并从同一特征并列导出角色头、种类头、目标代理二分类头和投影头；各输出头之间不存在串联依赖。种类—角色关系仅通过固定映射矩阵 A 对种类概率进行聚合，得到角色层概率，而不是把一个输出头直接作为另一个输出头的输入。")
    add_figure(doc, THEORY_ARCHITECTURE_FIGURE, 15.4, "图 2 理论感知分层识别模型结构：共享视觉特征、四个并列输出头与种类到角色的概率聚合")
    add_table(doc, ["符号", "定义"], [
        ("x", "输入矿物图像"),
        ("S", "细粒度矿物种类集合，当前包含 17 类"),
        ("R", "选矿角色集合，当前包含目标、含钛干扰、脉石/废石和金属光泽干扰四类"),
        ("A", "固定的种类—角色映射矩阵，A 的每一列仅对应一个角色"),
        ("p_s", "种类头输出的种类概率向量"),
        ("p_r", "角色头直接输出的角色概率向量"),
        ("p_tilde_r", "由 A 对 p_s 聚合得到的角色概率向量"),
        ("q(x)", "角色预测的最大 Softmax 置信度，用于选择性识别"),
    ], [3.0, 11.1], 10.2)
    add_body(doc, "种类头产生的角色聚合概率定义为：")
    add_equation(doc, equations["aggregation"], 7.0)
    add_body(doc, "完整训练目标由角色分类、种类分类、层级一致性、目标代理二分类和困难负样本约束共同组成：")
    add_equation(doc, equations["joint_loss"], 15.0)
    add_body(doc, "其中 alpha=0.50、beta=0.10、gamma=0.25、eta=0.10；困难负样本项仅约束目标矿物—含钛干扰和目标矿物—金属光泽干扰两类预定义高风险关系。一致性损失严格采用角色头分布到种类聚合角色分布的 KL 方向：")
    add_equation(doc, equations["consistency"], 10.5)
    add_body(doc, "该形式化描述的是公开矿物图像标签层级及其视觉预测约束，不把角色输出解释为品位、回收率、流程阶段或工业处置动作。")
    add_heading(doc, "4.6 角色可识别性命题", 2)
    add_body(doc, "命题 1：对图像 x 的候选种类集合 S_i，若集合内任一种类 s 经固定映射 A 均对应同一角色 r，即对所有 s∈S_i 均有 A e_s=e_r，则角色 r 唯一确定，即使具体种类 s 仍不能唯一确定。")
    add_body(doc, "证明要点：映射 A 将候选集合中的每个种类独热向量映射为同一角色独热向量，因此种类层的不确定性在角色层被消去。唯一识别种类足以确定角色；反过来，只要全部候选种类同属一角色，确定角色并不要求先确定具体种类。故在该固定映射下，角色可识别的条件弱于种类可识别的条件。该命题是针对当前领域问题的条件化形式化，不主张一般层级学习的新理论。")
    add_heading(doc, "4.7 选择性识别规则", 2)
    add_body(doc, "对角色头输出，以最大 Softmax 概率定义置信度：")
    add_equation(doc, equations["confidence"], 8.0)
    add_body(doc, "当 q(x) 不低于阈值 tau 时保留角色预测，否则将样本标记为 defer，含义仅为建议后续检查。覆盖率与保留样本错误率分别定义为：")
    add_equation(doc, equations["coverage"], 10.5)
    add_equation(doc, equations["risk"], 13.5)

    add_heading(doc, "5 实验结果与讨论", 1)
    add_heading(doc, "5.1 总体性能比较", 2)
    add_table(doc, ["模型", "Accuracy（%）", "Macro Precision（%）", "Macro Recall（%）", "Macro F1（%）"], [
        ("ResNet50", "74.84 ± 1.56", "73.32 ± 1.87", "71.63 ± 0.78", "72.06 ± 0.92"),
        ("EfficientNet-B0", "74.69 ± 1.28", "72.71 ± 1.28", "73.17 ± 1.12", "72.87 ± 1.21"),
        ("角色感知 EfficientNet-B0", "74.79 ± 0.84", "73.28 ± 1.18", "72.87 ± 1.00", "73.02 ± 1.09"),
        ("分层一致性 EfficientNet-B0", "75.42 ± 2.59", "73.71 ± 2.85", "73.61 ± 1.78", "73.41 ± 2.40"),
    ], [2.8, 3.1, 3.6, 3.4, 2.7], 10)
    add_body(doc, "三随机种子结果显示，EfficientNet-B0 的平均 Macro F1 比 ResNet50 高 0.67 个百分点，平均 Accuracy 基本相当。角色感知模型的平均 Macro F1 较普通 EfficientNet-B0 高 0.15 个百分点；分层一致性模型提高 0.54 个百分点，但标准差为 2.40 个百分点。两种改进的均值差均小于或接近种子间波动，因此均不应表述为在总体四分类上稳定优于基线。后续类别与风险指标用于说明不同机制带来的实际取舍。")
    add_figure(doc, FIG_DIR / "fig1_model_comparison.png", 15.4, "图 3 四种模型的总体性能对比（均值±样本标准差，n=3）")
    add_heading(doc, "5.2 各类别召回率分析", 2)
    add_table(doc, ["模型", "目标矿物", "含钛干扰", "脉石/废石", "金属光泽干扰"], [
        ("ResNet50", "71.58 ± 5.18", "77.53 ± 8.32", "83.61 ± 0.82", "53.79 ± 4.26"),
        ("EfficientNet-B0", "71.98 ± 0.83", "74.34 ± 2.15", "81.97 ± 2.85", "64.39 ± 3.47"),
        ("角色感知 EfficientNet-B0", "68.39 ± 3.34", "76.71 ± 0.31", "81.24 ± 0.42", "65.15 ± 1.43"),
        ("分层一致性 EfficientNet-B0", "77.56 ± 2.19", "75.15 ± 5.54", "81.69 ± 2.84", "60.04 ± 2.15"),
    ], [2.7, 3.3, 3.3, 3.3, 3.8], 9.5)
    add_body(doc, "脉石/废石类别在各模型中均具有较高召回率，说明该类在颜色、透明度和形貌上与目标组分存在较明显差异。金属光泽干扰是更具挑战的类别：黄铁矿、赤铁矿、针铁矿和黄铜矿等可能呈现暗色、反光或块状外观，与目标矿物形成视觉混淆。角色感知模型提高了两类困难干扰的召回，但降低目标召回，体现为保守目标判定；分层一致性模型将目标类召回提升至 77.56%，但金属光泽干扰召回降至 60.04%，显示细粒度种类监督会改变类别边界而非统一改善全部类别。")
    add_figure(doc, FIG_DIR / "fig2_class_recall.png", 15.4, "图 4 四类矿物在四种模型上的测试集召回率（均值±样本标准差，n=3）")
    add_heading(doc, "5.3 目标代理风险指标", 2)
    add_table(doc, ["设置", "目标代理 F1（%）", "目标漏选率（%）", "含钛干扰误入目标（%）", "金属光泽干扰误入目标（%）"], [
        ("加权交叉熵", "70.45 ± 1.34", "28.02 ± 0.46", "9.84 ± 1.24", "11.55 ± 0.66"),
        ("加权 Focal Loss", "69.94 ± 1.20", "27.89 ± 0.80", "9.57 ± 1.41", "13.83 ± 2.62"),
        ("角色感知困难负样本学习", "70.67 ± 2.57", "31.61 ± 3.34", "7.74 ± 0.35", "9.66 ± 0.98"),
    ], [3.7, 2.7, 3.0, 3.8, 4.2], 9.2)
    add_body(doc, "上述指标将四分类预测折叠为“目标代理/非目标代理”后计算，用于刻画与预选决策相关的代理风险，不代表真实生产线的回收率、精矿品位或抛废率。角色感知模型的目标代理 F1 与基线基本相当，但目标漏选率增加；与此同时，含钛干扰和金属光泽干扰误入目标的比例分别下降 2.10 和 1.89 个百分点。这说明该模型通过更保守的目标判定降低了两类高风险视觉混淆，存在明确的精度—召回取舍。")
    add_figure(doc, FIG_DIR / "fig7_target_proxy_strategy_comparison.png", 15.4, "图 5 四种训练策略的目标代理风险指标（均值±样本标准差，n=3）")
    add_heading(doc, "5.4 混淆矩阵与主要错误方向", 2)
    add_body(doc, "以 EfficientNet-B0、随机种子 20260728 的测试结果为例，在 1,284 张固定测试图像中正确识别 968 张、错误识别 316 张。主要错误集中于含钛干扰矿物与脉石/废石、目标矿物之间的相互混淆，其中“含钛干扰→脉石/废石”为 50 张，“含钛干扰→目标矿物”为 43 张，“目标矿物→含钛干扰”为 40 张，“脉石/废石→含钛干扰”为 40 张。")
    add_figure(doc, FIG_DIR / "fig3_confusion_matrix.png", 13.5, "图 6 EfficientNet-B0（随机种子 20260728）在固定测试集上的混淆矩阵")
    add_figure(doc, FIG_DIR / "fig4_top_error_pairs.png", 15.0, "图 7 EfficientNet-B0（随机种子 20260728）的主要错分方向")
    add_heading(doc, "5.5 Focal Loss 三随机种子消融讨论", 2)
    add_table(doc, ["设置", "Accuracy（%）", "Macro F1（%）", "目标矿物 Recall（%）", "含钛干扰 Recall（%）", "金属光泽干扰 Recall（%）"], [
        ("加权交叉熵", "74.69 ± 1.28", "72.87 ± 1.21", "71.98 ± 0.83", "74.34 ± 2.15", "64.39 ± 3.47"),
        ("加权 Focal Loss（gamma=2.0）", "75.21 ± 0.74", "73.22 ± 0.83", "72.11 ± 0.80", "76.99 ± 1.78", "61.93 ± 3.41"),
    ], [3.6, 2.5, 2.5, 3.0, 3.0, 3.2], 9.2)
    add_body(doc, "三随机种子下，Focal Loss 的平均 Macro F1 比交叉熵高 0.35 个百分点，含钛干扰召回率提高 2.65 个百分点，但金属光泽干扰召回率降低 2.46 个百分点；其目标代理 F1 降低 0.51 个百分点，金属光泽干扰误入目标的比例上升。由于总体差异小于或接近种子间波动，且关键风险指标未形成一致改善，本报告不将 Focal Loss 表述为稳定优于交叉熵的结论。该结果说明统一难例权重调整不能同时解决四类之间的结构性混淆。")
    add_figure(doc, FIG_DIR / "fig5_focal_loss_ablation.png", 15.0, "图 8 加权交叉熵与加权 Focal Loss 的三随机种子消融结果")
    add_heading(doc, "5.6 角色感知困难负样本学习讨论", 2)
    add_body(doc, "角色感知模型的四分类 Macro F1 为 73.02%±1.09%，相较普通 EfficientNet-B0 的提升仅为 0.15 个百分点，不能据此主张总体识别性能显著提高。其更有意义的变化发生在风险方向：将含钛干扰和金属光泽干扰误入目标的比例分别降至 7.74%±0.35% 和 9.66%±0.98%，同时目标类召回由 71.98% 降至 68.39%。因此，该模型体现为针对高风险误入错误的保守策略原型。后续可通过调节辅助损失权重、增加真实矿石样本及结合分选专家定义的可接受漏选阈值，确定更适宜的运行点。")
    add_heading(doc, "5.7 分层一致性模型讨论", 2)
    add_table(doc, ["设置", "目标代理 F1（%）", "目标漏选率（%）", "含钛干扰误入目标（%）", "金属光泽干扰误入目标（%）"], [
        ("普通 EfficientNet-B0", "70.45 ± 1.34", "28.02 ± 0.46", "9.84 ± 1.24", "11.55 ± 0.66"),
        ("分层一致性 EfficientNet-B0", "72.04 ± 2.05", "22.44 ± 2.19", "10.93 ± 3.37", "14.39 ± 2.92"),
    ], [3.7, 2.7, 3.0, 3.8, 4.2], 9.2)
    add_body(doc, "分层模型在 17 类矿物种类头上的平均测试 Accuracy 为 56.80%±2.23%，而角色主任务的 Macro F1 为 73.41%±2.40%。这符合该数据的任务特点：网页主题矿物、伴生矿物与基质可能共存，细粒度种类标签更难由单张宏观图片稳定确定；但种类概率聚合后的角色信息仍可为粗粒度识别提供辅助。相较普通 EfficientNet-B0，分层模型把目标漏选率降低 5.58 个百分点、目标代理 F1 提高 1.59 个百分点，却使含钛干扰和金属光泽干扰误入目标的比例分别增加 1.09 和 2.84 个百分点。因此，本阶段结果只能支持“分层监督能改善目标类召回、但尚未满足困难负样本风险控制”的结论。为区分不同损失项的影响，进一步进行了层级一致性项与困难负样本约束项的三随机种子组件消融。")
    add_heading(doc, "5.8 分层模型组件消融", 2)
    component_rows = [
        (
            row["setting"],
            format_percent(row["macro_f1_mean"], row["macro_f1_sample_std"]),
            format_percent(row["target_f1_mean"], row["target_f1_sample_std"]),
            format_percent(row["target_miss_rate_mean"], row["target_miss_rate_sample_std"]),
            format_percent(row["ti_bearing_intrusion_rate_mean"], row["ti_bearing_intrusion_rate_sample_std"]),
            format_percent(row["metallic_intrusion_rate_mean"], row["metallic_intrusion_rate_sample_std"]),
        )
        for row in component_ablation
    ]
    add_table(doc, ["设置", "Macro F1（%）", "目标代理 F1（%）", "目标漏选率（%）", "含钛干扰误入目标（%）", "金属光泽干扰误入目标（%）"], component_rows, [2.8, 2.4, 2.7, 2.6, 3.3, 3.5], 8.6)
    add_body(doc, "组件消融表明，移除困难负样本约束后，Macro F1 的均值反而上升 0.21 个百分点，但目标代理 F1 下降 0.41 个百分点，含钛干扰和金属光泽干扰误入目标比例分别上升 0.54 和 0.57 个百分点；这说明该约束更可能作用于关键负样本风险，而非提升总体平均分。移除层级一致性约束后，Macro F1 基本不变，但目标代理 F1 下降 1.12 个百分点，目标漏选率、含钛干扰误入目标率和金属光泽干扰误入目标率均上升。由于每组仅含三个随机种子、均值差与样本标准差仍同量级，以上结果只能支持风险指标上的一致方向趋势，不得表述为统计显著或稳定领先。")

    add_heading(doc, "5.9 角色可识别性的受控逻辑条件验证", 2)
    candidate_summary = role_validation["summary"]
    add_table(doc, ["候选集情形", "候选集数", "种类唯一可识别率（%）", "角色唯一可识别率（%）"], [
        (
            "角色一致",
            candidate_summary["role_consistent"]["row_count"],
            f"{100 * candidate_summary['role_consistent']['species_unique_rate']:.2f}",
            f"{100 * candidate_summary['role_consistent']['role_unique_rate']:.2f}",
        ),
        (
            "角色冲突",
            candidate_summary["role_conflict"]["row_count"],
            f"{100 * candidate_summary['role_conflict']['species_unique_rate']:.2f}",
            f"{100 * candidate_summary['role_conflict']['role_unique_rate']:.2f}",
        ),
        (
            "合计",
            candidate_summary["row_count"],
            f"{100 * candidate_summary['species_unique_rate']:.2f}",
            f"{100 * candidate_summary['role_unique_rate']:.2f}",
        ),
    ], [3.5, 2.2, 4.2, 4.2], 9.8)
    add_body(doc, f"该验证使用固定随机种子 {role_validation['seed']}，按候选规模 2、3、4 构造集合，并保持原训练、验证、测试划分计数不变。角色一致候选集共 {candidate_summary['role_consistent']['row_count']} 个，角色唯一可识别率为 {100 * candidate_summary['role_consistent']['role_unique_rate']:.2f}%；角色冲突候选集共 {candidate_summary['role_conflict']['row_count']} 个，角色唯一可识别率为 {100 * candidate_summary['role_conflict']['role_unique_rate']:.2f}%。结果只验证命题 1 的逻辑条件：候选种类同属一角色时，种类歧义不妨碍角色唯一确定。它不使用模型预测或原图视觉标签推断，也不替代真实多标签标注、外部测试或工业验证。")

    add_heading(doc, "5.10 固定测试划分上的选择性识别", 2)
    selected_thresholds = {0.0, 0.5, 0.7, 0.8, 0.9, 0.95}
    threshold_rows = []
    for row in selective_validation["threshold_summary"]:
        if row["threshold"] not in selected_thresholds:
            continue
        threshold_rows.append((
            f"{row['threshold']:.2f}",
            f"{row['mean']['retained_count']:.1f} ± {row['sample_std']['retained_count']:.1f}",
            format_percent(row["mean"]["coverage"], row["sample_std"]["coverage"]),
            format_percent(row["mean"]["risk"], row["sample_std"]["risk"]),
        ))
    add_table(doc, ["置信度阈值", "平均保留样本数", "覆盖率（%）", "保留样本错误率（%）"], threshold_rows, [3.0, 3.8, 3.6, 4.2], 9.8)
    add_figure(doc, SELECTIVE_RECOGNITION_FIGURE, 15.4, "图 9 三随机种子在固定测试划分上的选择性识别覆盖率—风险关系")
    add_body(doc, f"阈值从 0.00 提高到 0.90 时，三随机种子的平均覆盖率由 100.00% 降至 {100 * threshold_090['mean']['coverage']:.2f}%，保留样本错误率由 {100 * selective_validation['threshold_summary'][0]['mean']['risk']:.2f}% 降至 {100 * threshold_090['mean']['risk']:.2f}%。这体现了固定公开图像测试划分上的覆盖率—风险取舍：覆盖率越低，延后检查的样本越多。defer 仅表示建议送检、人工复核或后续传感确认；由于当前没有真实检查成本、动作代价和流程数据，不据此声称 XRF 成本最优、工业分选有效或统计显著。")

    add_heading(doc, "6 主要创新与阶段性成果", 1)
    add_heading(doc, "6.1 主要创新", 2)
    add_body(doc, "围绕理论感知分层识别主线，本阶段新增贡献严格限定为以下三项；数据整理、基线训练和组件消融作为可复现支撑材料，不扩展为工业效果声明。")
    add_table(doc, ["序号", "创新点", "具体体现"], [
        ("1", "角色层可识别性的领域化形式化", "用固定种类—角色映射说明候选种类不唯一而角色仍可唯一确定的条件，并以受控候选集验证该逻辑条件。"),
        ("2", "种类—角色一致性与困难负样本联合学习", "在共享视觉特征上设置四个并列输出头，以 p_tilde_r=A p_s 和 KL(p_r || p_tilde_r)约束层级一致性，并保留两类预定义困难关系的表示约束。"),
        ("3", "固定划分上的选择性识别评价", "从三个完整分层模型的测试预测中复现阈值—覆盖率—保留风险关系，将低置信度样本解释为延后检查而非已验证的处置决策。"),
    ], [1.0, 5.0, 9.8], 10)
    add_heading(doc, "6.2 阶段性成果", 2)
    add_body(doc, "项目已形成数据清单、质量复核记录、训练与测试脚本、模型结果、错误分析文件和可直接用于论文/报告的图表。训练过程固定数据版本与划分，输出文件包含每次实验的指标记录与混淆矩阵。原始公开图像因授权与来源差异不在公开仓库整体分发，但可通过元数据、图片编号和下载流程进行核验与再获取。")

    add_heading(doc, "7 局限性与后续研究", 1)
    add_heading(doc, "7.1 当前局限性", 2)
    add_table(doc, ["问题", "对结果解释的影响", "处理原则"], [
        ("公开标本图像与工业现场差异", "不能将当前准确率等同于皮带分选性能或生产指标", "报告定位为公开标本图像源域验证"),
        ("titanomagnetite 样本仅 23 张", "不足以支持钛磁铁矿独立细分类性能结论", "暂作为合并目标类补充并优先扩充"),
        ("单标签与伴生矿物共存", "页面主题矿物不一定是图像唯一可见对象", "主体不清样本进入 mixed_uncertain，不直接硬标训练"),
        ("来源外泛化尚未充分验证", "摄影者、产地和背景风格仍可能引入领域偏移", "后续回填来源元数据并构建来源留出测试"),
        ("改进模型的总体收益有限", "分层模型的两项组件消融显示关键风险指标存在改善方向，但总体分数差异仍小于或接近种子间波动", "报告取舍而非宣称稳定领先，后续在目标召回约束下做风险优化和来源外验证"),
    ], [3.3, 6.0, 6.5], 9.5)
    add_heading(doc, "7.2 后续研究计划", 2)
    add_body(doc, "第一，采集由专家或化验确认的真实钒钛矿颗粒图像，建立独立外部测试集。第二，回填产地、摄影者、收藏机构或页面来源信息，构建严格的来源留出测试；上述外部验证和来源留出评价当前均未完成。第三，在已有组件消融基础上继续考察种类交叉熵和目标代理辅助头，并在专家给定的目标召回下限下研究风险运行点。第四，建立独立未知矿物图像集，比较最大 Softmax 概率、能量分数等开放集拒识方法。第五，仅在取得真实流程阶段、品位/回收率、动作代价和检查成本后，研究阶段条件化决策与真实成本矩阵；结合 XRF、磁性或化验信息的多模态确认也属于未来验证，不由当前选择性识别结果代替。")
    add_note(doc, "第一篇论文的建议主线：", "以“面向钒钛矿相关矿物识别的分层角色一致性与困难负样本学习”为主线，严谨报告其对目标召回和困难干扰风险的取舍，并以组件消融、来源外测试与开放集评价支撑结论。阶段条件化选矿决策图需要真实流程、品位、回收率和专家确认的矿物—阶段—动作映射，当前仅作为后续研究方向，不写入本阶段方法效果结论。")

    add_heading(doc, "8 结论", 1)
    conclusions = [
        "构建了面向钒钛矿相关组分矿物识别的公开标本图像四分类数据集，共 8,529 张图像，并保留来源追溯、质量筛选和固定划分信息。",
        "通过混合样本隔离、重复/近重复审查和分层复核，形成了训练、验证、测试集分别为 5,961、1,284、1,284 张的可复现实验清单。",
        "ResNet50 与 EfficientNet-B0 的三随机种子结果表明，EfficientNet-B0 的平均 Macro F1 为 72.87%±1.21%，略高于 ResNet50 的 72.06%±0.92%。",
        "脉石/废石较易识别，而含钛干扰矿物与金属光泽干扰矿物是主要难点。三随机种子 Focal Loss 消融未在总体和目标代理风险指标上形成一致改善，说明统一难例权重不足以解决类别关系混淆。",
        "角色感知困难负样本学习在总体 Macro F1 上与基线接近，但可降低含钛干扰和金属光泽干扰误入目标的比例，代价是目标类漏选增加；该方法构成后续依据实际分选风险阈值调优的可解释原型。",
        "矿物种类—选矿角色分层一致性模型将种类概率与角色概率联合约束。两项三随机种子组件消融显示：去除一致性或困难负样本约束后，完整模型在目标代理 F1 和两类关键干扰误入目标风险上均呈更优方向，但总体 Macro F1 差异仍与种子间波动相近；当前应将其视为可量化的风险改善趋势，而非已被证明的稳定性能改进。",
        f"角色可识别性命题得到受控逻辑条件验证：{role_consistent['row_count']} 个角色一致候选集的角色唯一可识别率为 {100 * role_consistent['role_unique_rate']:.2f}%，{role_conflict['row_count']} 个角色冲突候选集为 {100 * role_conflict['role_unique_rate']:.2f}%。该证据不使用视觉标签推断或模型预测。",
        f"固定测试划分上的选择性识别表明阈值提高会降低覆盖率并降低保留样本错误率；阈值 0.90 时平均覆盖率为 {100 * threshold_090['mean']['coverage']:.2f}%，保留样本错误率为 {100 * threshold_090['mean']['risk']:.2f}%。这只是固定划分证据，不是检查成本或处置策略的最优性证明。",
        "当前结果支持公开矿物标本图像上的理论形式化、受控逻辑验证与固定划分选择性识别评价，但不等同于工业现场分选、精矿品位/回收率、XRF 成本最优、外部验证或统计显著性结论。",
    ]
    for idx, text in enumerate(conclusions, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, after=5, line=1.5, first_indent=0)
        add_text(p, f"{idx}. ", 12, True, HEADING_FONT)
        add_text(p, text, 12)

    add_heading(doc, "参考文献", 1)
    references = [
        "[1] Shin Y, Shin S. Rock Classification in a Vanadiferous Titanomagnetite Deposit Based on Supervised Machine Learning[J]. Minerals, 2022, 12(4): 461.",
        "[2] Baek J, Cho S, Shin S. Vanadiferous titanomagnetite ore rock classifier using machine learning from portable X-ray fluorescence spectra[J]. Environmental Earth Sciences, 2025, 84: 368.",
        "[3] Nesteruk S, et al. MineralImage5k: A benchmark for zero-shot raw mineral visual recognition and description[J]. Computers & Geosciences, 2023, 178: 105414.",
        "[4] Wu B, et al. Mineral Identification Based on Multi-Label Image Classification[J]. Minerals, 2022, 12(11): 1338.",
        "[5] Ji X, et al. Identifying Minerals from Image Using Out-of-Distribution Artificial Intelligence-Based Model[J]. Minerals, 2024, 14(6): 627.",
        "[6] Tang H. 基于深度学习的磁铁矿图像识别及应用研究[D]. 华北理工大学, 2023.",
        "[7] 孙昊, 郑建明. 卷积神经网络在矿石识别中的应用[J]. 福建电脑, 2023.",
        "[8] Mindat.org. The mineral and locality database[EB/OL]. https://www.mindat.org/.",
        "[9] Iron Ore Image Recognition Through Multi-View Evolutionary Deep Fusion Method[J]. Future Internet, 2025, 17(12): 553. doi:10.3390/fi17120553.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        set_paragraph_format(p, after=3, line=1.3, first_indent=0)
        add_text(p, ref, 10.5)

    doc.add_page_break()
    add_heading(doc, "附录 A 数据字段与复现说明", 1)
    add_heading(doc, "A.1 关键数据字段", 2)
    add_table(doc, ["字段", "说明"], [
        ("image_id", "项目内部唯一图像编号，如 VTM-000001"),
        ("relative_path", "图像在数据版本目录中的相对路径"),
        ("mineral_label", "细粒度矿物名称标签"),
        ("four_class_label", "四分类任务标签"),
        ("mindat_photo_id", "来源图片编号，用于来源追溯与分组控制"),
        ("split_group_id", "重复或近重复控制后的分组编号"),
        ("split", "train、val 或 test 固定数据划分"),
    ], [4.2, 9.9], 10.5)
    add_heading(doc, "A.2 可复现材料", 2)
    add_body(doc, "本项目保留数据清单、采集与筛选辅助脚本、数据质量审计记录、训练脚本、实验指标、混淆矩阵、逐图预测结果和图表源数据。正式复现时应以固定清单 dataset_split_manifest_v1_0.csv 为准，并保持图片编号分组、训练随机种子和评价指标定义一致。由于公开网页图片可能具有不同作者与许可，公开共享时应优先发布元数据、脚本、图表源数据和处理规则，而不应在未核实许可的情况下整体再分发原始图片。")
    add_heading(doc, "A.3 结题前待补信息", 2)
    add_table(doc, ["项目", "当前状态", "建议补充位置"], [
        ("项目编号、承担单位、负责人", "待核验", "封面及验收报告基本信息"),
        ("任务书考核指标", "本目录未发现任务书本体", "结题验收报告“任务完成情况”逐项对照"),
        ("复核人和复核日期", "复核决策已形成，责任字段待完善", "数据质控附件与报告 3.3 节"),
        ("财务决算、成果证明与凭证", "需由项目承担单位补充", "结题验收材料附件"),
    ], [4.0, 5.5, 4.6], 10)

    doc.save(OUTPUT)
    equation_temp.cleanup()
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
