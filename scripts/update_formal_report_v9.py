from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from update_formal_report_v2 import FORMAL_REPORT, _set_run_font


FORMAL_CONFIGURATIONS = (
    "rsg_complete",
    "cgdc_complete",
    "cgdc_shared_features",
    "cgdc_unconditional",
    "cgdc_no_decomposition_loss",
)
FORMAL_SEEDS = ("20260727", "20260728", "20260729")
APPENDIX_HEADING = "附录 D CGDC-RSG-HRGV 网络理论与实验"
RPG_APPENDIX_HEADING = "附录 E RPG-HRGV 角色分区不确定性门控理论与实验"
MRPG_APPENDIX_HEADING = "附录 F M-RPG-HRGV 容量归一化单调门控理论与实验"
FIVE_SEED_APPENDIX_HEADING = "附录 G M-RPG-HRGV 五随机种子扩展验证"
RSG_THEORY_EVIDENCE_APPENDIX_HEADING = "附录 H RSG-HRGV 理论性质与证据对应图"
BACKBONE_PORTABILITY_HEADING = "H.3 主干替换不变性与 ResNet50 跨主干确认"
ROOT = Path(__file__).resolve().parents[1]
FIGURE_PATH = ROOT / "outputs" / "paper_figures" / "cgdc_rsg_hrgv_architecture.png"
RPG_FIGURE_PATH = ROOT / "outputs" / "paper_figures" / "rpg_hrgv_architecture.png"
MRPG_FIGURE_PATH = ROOT / "结题" / "图_MRPG-HRGV-Net_网络结构与理论性质.png"
RSG_THEORY_EVIDENCE_FIGURE_PATH = ROOT / "outputs" / "paper_figures_v3" / "fig_rsg_theory_evidence.png"
BACKBONE_PORTABILITY_FIGURE_PATH = (
    ROOT / "outputs" / "paper_figures_v3" / "fig_rsg_theory_evidence_portability.png"
)
DEFAULT_ANALYSIS_DIR = ROOT / "outputs" / "business_metrics" / "cgdc_rsg_hrgv" / "formal"
DEFAULT_RPG_ANALYSIS_DIR = ROOT / "outputs" / "business_metrics" / "rpg_hrgv" / "formal"
DEFAULT_FIVE_SEED_ANALYSIS_DIR = (
    ROOT / "outputs" / "business_metrics" / "target_recall_extension" / "five_seed"
)
DEFAULT_THEORY_REPLAY_ANALYSIS_DIR = (
    ROOT / "outputs" / "business_metrics" / "rsg_hrgv" / "theory_replay"
)
DEFAULT_BACKBONE_PORTABILITY_ANALYSIS_DIR = (
    ROOT / "outputs" / "business_metrics" / "rsg_hrgv" / "resnet50_portability"
)
BACKBONE_REPLAY_HEADING = "H.4 ResNet50 跨主干高精度重放"
GATE_RELIABILITY_HEADING = "H.5 门控可靠性分层诊断"
GATE_RELIABILITY_FIGURE_PATH = ROOT / "outputs" / "paper_figures_v3" / "fig_rsg_gate_reliability.png"
FORMULA_DIR = ROOT / "outputs" / "report_assets_v9"
RPG_FORMAL_CONFIGURATIONS = (
    "rsg_complete",
    "rpg_complete",
    "rpg_without_within",
    "rpg_without_between",
    "rpg_total_entropy_only",
)
MRPG_FORMAL_CONFIGURATIONS = (
    "rsg_complete",
    "rpg_complete",
    "mrpg_complete",
    "mrpg_unconstrained_between",
    "mrpg_without_between",
)
FIVE_SEED_CONFIGURATIONS = ("rsg_complete", "mrpg_complete")
FIVE_SEEDS = ("20260727", "20260728", "20260729", "20260730", "20260731")
PRIMARY_CONTRIBUTION_OLD = (
    "本次修订把理论贡献限定为可验证的三层结构：其一，固定种类—角色映射诱导逐样本风险支配和经验风险收缩；"
    "其二，KL 一致性通过 Pinsker 不等式控制双头分布差异，并由消融验证其分布效果；"
    "其三，温度校准结合多阈值修正的 Clopper–Pearson 上界，实现验证集上的选择性风险认证。"
    "来源留出和代理消融进一步限定这些结论的适用边界。"
)
PRIMARY_CONTRIBUTION_NEW = (
    "本次修订把方法贡献限定为相互衔接且可证伪的四层结构：其一，固定种类--角色映射把 17 类矿物种类后验聚合为四角色后验，"
    "并限定其只服务于公开标本图像的视觉角色识别；其二，KL 一致性通过 Pinsker 不等式控制直接角色头与种类映射头的分布差异；"
    "其三，RSG-HRGV 以直接角色专家和种类映射专家的真实类对数损失差构造软最优门控目标，并以差距权重训练门控。"
    "定理 B.1--B.3 分别给出门控误差的路由后悔上界、软目标对硬最优门控的指数逼近和后悔分支的局部梯度隔离。"
    "在两位专家概率和凸门控保持不变时，这三条性质不依赖特定视觉主干，构成主干替换不变性命题；ResNet50 三随机种子确认用于检验其经验边界。"
    "固定测试与摄影者留出确认均支持所定义平均路由后悔下降，但不构成总体分类性能优势；其四，M-RPG 的容量归一化、单调门控和凸融合性质作为理论扩展保留，"
    "五随机种子扩展未确认稳定经验增益。来源留出、代理消融及上述负结果共同限定这些结论的适用边界。"
)


def load_formal_cgdc_evidence(analysis_dir: Path) -> dict[str, object]:
    """Load a complete five-configuration, three-seed CGDC analysis summary."""
    summary_path = analysis_dir / "cgdc_three_seed_summary.json"
    if not summary_path.is_file():
        raise ValueError("Formal CGDC evidence summary is missing.")
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    if set(summary) != set(FORMAL_CONFIGURATIONS):
        raise ValueError("Formal CGDC evidence must contain all five configurations.")
    for configuration in FORMAL_CONFIGURATIONS:
        values = summary[configuration].get("macro_f1", {}).get("values", [])
        if len(values) != len(FORMAL_SEEDS):
            raise ValueError("Formal CGDC evidence must contain three registered seeds.")
    return summary


def load_formal_rpg_evidence(analysis_dir: Path) -> dict[str, object]:
    """Load the complete RPG ablation summary required for report appendix E."""
    summary_path = analysis_dir / "rpg_three_seed_summary.json"
    if not summary_path.is_file():
        raise ValueError("Formal RPG evidence summary is missing.")
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    if set(summary) != set(RPG_FORMAL_CONFIGURATIONS):
        raise ValueError("Formal RPG evidence must contain all five configurations.")
    for configuration in RPG_FORMAL_CONFIGURATIONS:
        values = summary[configuration].get("macro_f1", {}).get("values", [])
        if len(values) != len(FORMAL_SEEDS):
            raise ValueError("Formal RPG evidence must contain three registered seeds.")
    return summary


def load_formal_mrpg_evidence(analysis_dir: Path) -> dict[str, object]:
    """Load the complete M-RPG ablation summary required for report appendix F."""
    summary_path = analysis_dir / "mrpg_three_seed_summary.json"
    if not summary_path.is_file():
        raise ValueError("Formal M-RPG evidence summary is missing.")
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    if set(summary) != set(MRPG_FORMAL_CONFIGURATIONS):
        raise ValueError("Formal M-RPG evidence must contain all five configurations.")
    for configuration in MRPG_FORMAL_CONFIGURATIONS:
        values = summary[configuration].get("macro_f1", {}).get("values", [])
        if len(values) != len(FORMAL_SEEDS):
            raise ValueError("Formal M-RPG evidence must contain three registered seeds.")
    return summary


def load_five_seed_extension_evidence(analysis_dir: Path) -> tuple[dict[str, object], dict[str, object]]:
    """Load the registered five-seed extension summary and paired comparison."""
    summary_path = analysis_dir / "five_seed_summary.json"
    paired_path = analysis_dir / "paired_mrpg_complete_vs_rsg_complete.json"
    if not summary_path.is_file() or not paired_path.is_file():
        raise ValueError("Five-seed extension evidence is incomplete.")
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    if set(summary) != set(FIVE_SEED_CONFIGURATIONS):
        raise ValueError("Five-seed extension must contain RSG and M-RPG only.")
    for configuration in FIVE_SEED_CONFIGURATIONS:
        values = summary[configuration].get("macro_f1", {}).get("values", [])
        if len(values) != len(FIVE_SEEDS):
            raise ValueError("Five-seed extension evidence must contain five registered seeds.")
    paired = json.loads(paired_path.read_text(encoding="utf-8"))
    required = {
        "classification": {
            "macro_f1",
            "target_recall",
            "ti_to_target_intrusion",
            "metallic_to_target_intrusion",
        },
        "calibration": {"brier_score", "expected_calibration_error"},
    }
    for section, keys in required.items():
        if not keys <= set(paired.get(section, {})):
            raise ValueError(f"Five-seed paired evidence lacks required {section} metrics.")
    return summary, paired


def load_rsg_theory_replay_evidence(analysis_dir: Path) -> dict[str, object]:
    """Load the high-precision checkpoint replay used only for theorem consistency."""
    summary_path = analysis_dir / "theory_replay_summary.json"
    if not summary_path.is_file():
        raise ValueError("RSG theory replay evidence summary is missing.")
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    overall = summary.get("overall", {})
    required = {
        "run_count",
        "sample_count",
        "b1_max_residual",
        "b1_violation_count",
        "b2_max_residual",
        "b2_violation_count",
    }
    if not required <= set(overall) or len(summary.get("runs", [])) != 6:
        raise ValueError("RSG theory replay evidence is incomplete.")
    return summary


def load_rsg_backbone_portability_evidence(
    analysis_dir: Path,
) -> tuple[dict[str, object], dict[str, object]]:
    """Load the matched ResNet50 check for the backbone-invariance boundary."""
    summary_path = analysis_dir / "rsg_three_seed_summary.json"
    paired_path = analysis_dir / "paired_rsg_complete_vs_hrgv_reference.json"
    if not summary_path.is_file() or not paired_path.is_file():
        raise ValueError("RSG backbone portability evidence is incomplete.")
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    expected_configurations = {"hrgv_reference", "rsg_complete"}
    if set(summary) != expected_configurations:
        raise ValueError("RSG backbone portability evidence must compare HRGV and RSG only.")
    for configuration in expected_configurations:
        for metric in ("accuracy", "macro_f1", "mean_routing_regret_nll"):
            values = summary[configuration].get(metric, {}).get("values", [])
            if len(values) != len(FORMAL_SEEDS):
                raise ValueError("RSG backbone portability evidence must contain three registered seeds.")
    paired = json.loads(paired_path.read_text(encoding="utf-8"))
    if not {"accuracy", "macro_f1"} <= set(paired.get("classification", {})):
        raise ValueError("RSG backbone portability evidence lacks classification metrics.")
    if not {"difference", "ci_low", "ci_high"} <= set(paired.get("routing_regret", {})):
        raise ValueError("RSG backbone portability evidence lacks routing regret metrics.")
    return summary, paired


def load_rsg_backbone_replay_evidence(analysis_dir: Path) -> dict[str, object]:
    """Select the high-precision ResNet50 replay rows from a combined replay summary."""
    summary_path = analysis_dir / "theory_replay_summary.json"
    if not summary_path.is_file():
        raise ValueError("RSG backbone replay evidence summary is missing.")
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    numeric = summary.get("numeric_settings", {})
    if not {"float32_epsilon", "tolerance"} <= set(numeric):
        raise ValueError("RSG backbone replay numeric settings are incomplete.")
    required = {
        "protocol",
        "seed",
        "sample_count",
        "minimum_true_probability",
        "b1_max_residual",
        "b1_violation_count",
        "b2_max_residual",
        "b2_violation_count",
    }
    runs = [
        run
        for run in summary.get("runs", [])
        if run.get("protocol") == "resnet50_portability"
    ]
    if len(runs) != 3 or any(not required <= set(run) for run in runs):
        raise ValueError("RSG backbone replay evidence must contain three complete ResNet50 runs.")
    if len({str(run["seed"]) for run in runs}) != 3:
        raise ValueError("RSG backbone replay seeds must be unique.")
    return {
        "runs": sorted(runs, key=lambda run: str(run["seed"])),
        "overall": {
            "run_count": len(runs),
            "sample_count": sum(int(run["sample_count"]) for run in runs),
            "minimum_true_probability": min(
                float(run["minimum_true_probability"]) for run in runs
            ),
            "b1_max_residual": max(float(run["b1_max_residual"]) for run in runs),
            "b1_violation_count": sum(int(run["b1_violation_count"]) for run in runs),
            "b2_max_residual": max(float(run["b2_max_residual"]) for run in runs),
            "b2_violation_count": sum(int(run["b2_violation_count"]) for run in runs),
        },
        "numeric_settings": numeric,
    }


def load_rsg_gate_reliability_evidence(analysis_dir: Path) -> dict[str, object]:
    """Load the descriptive B.1/B.2 strata evidence used in appendix H.5."""
    summary_path = analysis_dir / "gate_reliability_summary.json"
    if not summary_path.is_file():
        raise ValueError("RSG gate-reliability summary is missing.")
    evidence = json.loads(summary_path.read_text(encoding="utf-8"))
    protocols = evidence.get("protocols", {})
    expected_protocols = {"fixed", "photographer_holdout", "resnet50_portability"}
    if set(protocols) != expected_protocols:
        raise ValueError("RSG gate-reliability evidence must include the three registered protocols.")
    required = {
        "sample_count",
        "mean_routing_regret_nll",
        "mean_b1_local_bound",
        "mean_soft_hard_deviation",
        "mean_b2_bound",
        "b1_local_violation_count",
        "b2_violation_count",
    }
    if any(not required <= set(protocols[name]) for name in expected_protocols):
        raise ValueError("RSG gate-reliability protocol metrics are incomplete.")
    overall = evidence.get("overall", {})
    if not {"run_count", "sample_count", "b1_local_violation_count", "b2_violation_count"} <= set(overall):
        raise ValueError("RSG gate-reliability overall metrics are incomplete.")
    if int(overall["b1_local_violation_count"]) != 0 or int(overall["b2_violation_count"]) != 0:
        raise ValueError("RSG gate-reliability evidence contains theorem violations.")
    return evidence


def _add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    _set_run_font(run)


def update_primary_contribution_statement(document: Document) -> bool:
    """Replace the earlier generic contribution summary with the evidence-aligned one."""
    for paragraph in document.paragraphs:
        if paragraph.text.strip() != PRIMARY_CONTRIBUTION_OLD:
            continue
        paragraph.clear()
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(PRIMARY_CONTRIBUTION_NEW)
        _set_run_font(run)
        return True
    return False


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _set_run_font(run)
                    run.font.size = Pt(8.5)
                    run.bold = row_index == 0
    document.add_paragraph()


def _format_percent(value: float) -> str:
    return f"{100 * value:.2f}%"


def _load_paired_comparison(analysis_dir: Path, configuration: str) -> dict[str, object]:
    path = analysis_dir / f"paired_{configuration}_vs_rsg_complete.json"
    if not path.is_file():
        raise ValueError(f"Formal paired comparison is missing: {path.name}")
    return json.loads(path.read_text(encoding="utf-8"))


def _add_architecture_figure(document: Document) -> None:
    if not FIGURE_PATH.is_file():
        raise ValueError(f"CGDC architecture figure is missing: {FIGURE_PATH}")
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(FIGURE_PATH), width=Cm(15.4))
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        caption.add_run("图 19 CGDC-RSG-HRGV-Net 的跨粒度证据分解与分歧校正结构"),
        size=9.5,
    )


def render_formulas() -> dict[str, Path]:
    """Render the report's CGDC equations as stable bitmap formula assets."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    formulas = {
        "decomposition": (
            r"$\mathbf{u}_d=\mathbf{h}+A_d(\mathbf{h}),\quad "
            r"\mathbf{u}_s=\mathbf{h}+A_s(\mathbf{h}),\quad "
            r"\mathcal{L}_{dec}=\operatorname{mean}\!\left["
            r"\cos^2(A_d(\mathbf{h}),A_s(\mathbf{h}))\right]$"
        ),
        "calibration": (
            r"$\rho=1-\exp[-D_{JS}(\mathbf{p}_d\Vert\mathbf{p}_m)],\quad "
            r"\mathbf{p}_c=\operatorname{softmax}(\log\mathbf{p}_f+\rho\tanh("
            r"\operatorname{MLP}(\mathbf{z}_c)))$"
        ),
        "bound": (
            r"$\mathbf{p}_d=\mathbf{p}_m\Rightarrow\rho=0\Rightarrow"
            r"\mathbf{p}_c=\mathbf{p}_f;\qquad "
            r"\left|\log\frac{p_{c,j}}{p_{c,k}}-\log\frac{p_{f,j}}{p_{f,k}}\right|"
            r"<2\rho\leq1$"
        ),
    }
    paths: dict[str, Path] = {}
    for name, formula in formulas.items():
        figure = plt.figure(figsize=(14, 1.05), dpi=220, facecolor="white")
        figure.text(0.5, 0.5, formula, ha="center", va="center", fontsize=16, color="#172033")
        path = FORMULA_DIR / f"cgdc_{name}.png"
        figure.savefig(path, bbox_inches="tight", pad_inches=0.12, facecolor="white")
        plt.close(figure)
        paths[name] = path
    return paths


def render_rpg_formulas() -> dict[str, Path]:
    """Render the RPG equations once as stable report assets."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    formulas = {
        "entropy": (
            r"$\mathbf{p}_m=M\mathbf{p}_s,\qquad H(S)=H(R)+H(S\mid R),\qquad "
            r"U_{between}=H(R),\quad U_{within}=H(S\mid R)$"
        ),
        "gate": (
            r"$\mathbf{z}_{RPG}=[\mathbf{h},H(\mathbf{p}_d),U_{between},U_{within},"
            r"\left|H(\mathbf{p}_d)-U_{between}\right|],\qquad g=\sigma(G(\mathbf{z}_{RPG}))$"
        ),
        "fusion": (
            r"$\mathbf{p}_f=g\mathbf{p}_d+(1-g)\mathbf{p}_m,\qquad 0\leq g\leq1$"
        ),
    }
    paths: dict[str, Path] = {}
    for name, formula in formulas.items():
        figure = plt.figure(figsize=(14, 1.05), dpi=220, facecolor="white")
        figure.text(0.5, 0.5, formula, ha="center", va="center", fontsize=16, color="#172033")
        path = FORMULA_DIR / f"rpg_{name}.png"
        figure.savefig(path, bbox_inches="tight", pad_inches=0.12, facecolor="white")
        plt.close(figure)
        paths[name] = path
    return paths


def render_mrpg_formulas() -> dict[str, Path]:
    """Render the M-RPG equations as stable report assets."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    formulas = {
        "identity": (
            r"$H(S\mid x)=H(R\mid x)+H(S\mid R,x)=U_{between}+U_{within}$"
        ),
        "normalization": (
            r"$C_{within}(p_m)=\sum_r p_m(r\mid x)\log n_r,\qquad "
            r"u_{between}=H(p_m)/\log|R|,\qquad "
            r"u_{within}=H(S\mid R,x)/C_{within}(p_m)$"
        ),
        "monotone_gate": (
            r"$g_M=\sigma\!\left(Q([\mathbf{h},H(p_d),u_{within}])"
            r"+\operatorname{softplus}(\beta)u_{between}\right),\qquad "
            r"p_f=g_Mp_d+(1-g_M)p_m$"
        ),
        "derivative": (
            r"$\frac{\partial g_M}{\partial u_{between}}="
            r"g_M(1-g_M)\operatorname{softplus}(\beta)\geq0$"
        ),
    }
    paths: dict[str, Path] = {}
    for name, formula in formulas.items():
        figure = plt.figure(figsize=(14, 1.05), dpi=220, facecolor="white")
        figure.text(0.5, 0.5, formula, ha="center", va="center", fontsize=16, color="#172033")
        path = FORMULA_DIR / f"mrpg_{name}.png"
        figure.savefig(path, bbox_inches="tight", pad_inches=0.12, facecolor="white")
        plt.close(figure)
        paths[name] = path
    return paths


def _add_formula(document: Document, path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(15.3))
    _add_body(document, caption)


def _add_theory_statement(document: Document) -> None:
    formulas = render_formulas()
    document.add_heading("D.1 网络结构与理论定位", level=2)
    _add_body(
        document,
        "CGDC-RSG-HRGV-Net 以 EfficientNet-B0 为共享视觉主干，将共享特征 h 通过直接角色残差适配器和矿物种类残差适配器分解为 u_d=h+A_d(h) 与 u_s=h+A_s(h)。直接角色专家给出 p_d，种类专家经固定种类-角色映射聚合为 p_m；既有 RSG 门控给出融合后验 p_f=g p_d+(1-g)p_m。仅在两路证据不一致时，CGDC 校正器再执行受界后验修正。",
    )
    _add_formula(document, formulas["decomposition"], "式（D-1） 跨粒度残差证据分解与适配器去相关约束")
    _add_body(
        document,
        "令 rho=1-exp[-D_JS(p_d||p_m)]，s=tanh(MLP[z_c])，其中 z_c 包含两路分解特征、逐元素乘积、绝对差、对数后验差、熵与 JS 分歧。最终校正后验为 p_c=softmax(log p_f+rho s)，总损失为 L_CGDC=L_HRGV+lambda_dec L_dec+lambda_cal L_cal，其中 L_dec=mean cos^2(A_d(h),A_s(h))，L_cal=-mean log p_c(y)。",
    )
    _add_formula(document, formulas["calibration"], "式（D-2） 分歧触发的受界后验校正")
    _add_body(
        document,
        "命题 P1（协议一致性）：当 p_d=p_m 时 rho=0，故 p_c=p_f=p_d=p_m。命题 P2（概率有效性）：p_c 为 softmax 输出，非负且四类概率和为 1。命题 P3（有界对数几率修正）：因每个校正残差满足 |s_j|<1，任意类别 j、k 有 |log[p_c(j)/p_c(k)]-log[p_f(j)/p_f(k)]|<2rho。命题 P4（全局校正预算）：对任意两路四类后验，D_JS 不超过 ln2，故 rho=1-exp[-D_JS]≤0.5，结合 P3 可得任意对数几率修正严格小于 1。P1-P4 是当前公式与计算图的确定性性质；适配器分解和校正器是否改善分类或校准，必须由以下三随机种子实验检验。",
    )
    _add_formula(document, formulas["bound"], "式（D-3） 协议一致性、分歧相关修正与全局校正预算")
    _add_architecture_figure(document)


def _add_three_seed_summary(document: Document, evidence: dict[str, object]) -> None:
    document.add_heading("D.2 三随机种子受控实验", level=2)
    labels = {
        "rsg_complete": "RSG 完整模型",
        "cgdc_complete": "CGDC 完整模型",
        "cgdc_shared_features": "共享特征消融",
        "cgdc_unconditional": "无条件校正消融",
        "cgdc_no_decomposition_loss": "取消分解损失",
    }
    rows: list[list[str]] = []
    for configuration in FORMAL_CONFIGURATIONS:
        summary = evidence[configuration]
        rows.append(
            [
                labels[configuration],
                _format_percent(summary["accuracy"]["mean"]),
                _format_percent(summary["macro_f1"]["mean"]),
                _format_percent(summary["target_recall"]["mean"]),
                _format_percent(summary["brier_score"]["mean"]),
                _format_percent(summary["expected_calibration_error"]["mean"]),
            ]
        )
    _add_table(
        document,
        ["配置", "Accuracy", "Macro F1", "目标类召回", "Brier", "ECE"],
        rows,
    )
    _add_body(
        document,
        "表中为三随机种子均值；完整配置与四个控制变量消融共享固定数据划分、图像增强、主干、训练预算和测试协议。Brier 与 ECE 均为越低越好的后验质量指标，不能替代真实选矿回收率或品位指标。",
    )


def _add_paired_evidence(document: Document, analysis_dir: Path) -> None:
    document.add_heading("D.3 成对 Bootstrap 证据与结论边界", level=2)
    rows: list[list[str]] = []
    for configuration, label in (
        ("cgdc_complete", "CGDC 完整模型"),
        ("cgdc_shared_features", "共享特征消融"),
        ("cgdc_unconditional", "无条件校正消融"),
        ("cgdc_no_decomposition_loss", "取消分解损失"),
    ):
        paired = _load_paired_comparison(analysis_dir, configuration)
        macro = paired["classification"]["macro_f1"]
        brier = paired["calibration"]["brier_score"]
        ece = paired["calibration"]["expected_calibration_error"]
        rows.append(
            [
                label,
                f"{_format_percent(macro['difference'])} [{_format_percent(macro['ci_low'])}, {_format_percent(macro['ci_high'])}]",
                f"{_format_percent(brier['difference'])} [{_format_percent(brier['ci_low'])}, {_format_percent(brier['ci_high'])}]",
                f"{_format_percent(ece['difference'])} [{_format_percent(ece['ci_low'])}, {_format_percent(ece['ci_high'])}]",
            ]
        )
    _add_table(document, ["相对 RSG", "Macro F1 差值 [95% CI]", "Brier 差值 [95% CI]", "ECE 差值 [95% CI]"], rows)
    _add_body(
        document,
        "每个区间以图片 split_group_id 聚类、随机种子与组两阶段重采样的 2,000 次 Bootstrap 得到。校正相关的经验结论仅以 Brier/ECE 差值及其区间为准；若区间跨零，报告应保留为未观察到稳定增益，而不能根据单个种子或单项指标宣称优势。",
    )
    _add_body(
        document,
        "结论边界：本附录验证的是公开矿物标本图像上的四角色闭集识别和后验校准行为，不等同于工业分选、精矿品位预测、回收率提升、钒钛元素含量检测或未知矿物拒识性能。阶段条件化选矿决策图仍作为下一篇工作与后续研究方向。",
    )


def _add_rpg_architecture_figure(document: Document) -> None:
    if not RPG_FIGURE_PATH.is_file():
        raise ValueError(f"RPG architecture figure is missing: {RPG_FIGURE_PATH}")
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(RPG_FIGURE_PATH), width=Cm(15.4))
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        caption.add_run("图 20 RPG-HRGV-Net 的角色分区不确定性门控结构"),
        size=9.5,
    )


def _add_rpg_theory_statement(document: Document) -> None:
    formulas = render_rpg_formulas()
    document.add_heading("E.1 问题设定、网络结构与可检验命题", level=2)
    _add_body(
        document,
        "RPG-HRGV-Net 保持原有 EfficientNet-B0 共享主干、直接角色专家和矿物种类专家。种类专家输出的 17 类矿物后验 p_s 通过固定的种类-角色映射矩阵 M 聚合为四角色后验 p_m=Mp_s；该映射仅编码本研究的公开图像四角色任务，不推断元素含量、品位或工业流程动作。",
    )
    _add_formula(document, formulas["entropy"], "式（E-1） 固定角色分区下的熵链式分解")
    _add_body(
        document,
        "设角色变量为 R、细粒度种类变量为 S，则 H(S)=H(R)+H(S|R)。其中 U_between=H(R) 为角色间不确定性，U_within=H(S|R) 为角色内不确定性，后者刻画同一角色内部矿物种类的歧义。RPG 门控以这两个互补量及直接专家熵作为证据，而不将单一总熵误作所有不确定性的来源。",
    )
    _add_formula(document, formulas["gate"], "式（E-2） 角色分区不确定性门控")
    _add_formula(document, formulas["fusion"], "式（E-3） 门控凸融合后验")
    _add_body(
        document,
        "命题 P-R1（精确分解）：在固定映射 M 下，式（E-1）为离散熵的链式法则，因此 U_between+U_within 与种类总熵严格一致。命题 P-R2（角色稳定性）：当种类候选仅在同一角色内变化时，角色后验 p_m 不变，U_between 不因该类内歧义而上升；因此角色决策所需的可辨识信息弱于精确矿物种类识别所需的信息。命题 P-R3（凸融合包络）：因 g 属于 [0,1]，每一角色概率 p_f(r) 位于 p_d(r) 与 p_m(r) 的闭区间内，RPG 门控不能产生超出两路专家证据范围的单角色概率。P-R1 与 P-R3 为计算图的确定性性质；P-R2 的实际收益以及分区熵是否优于消融变量，由下列三随机种子实验检验。",
    )
    _add_rpg_architecture_figure(document)


def _add_rpg_three_seed_summary(document: Document, evidence: dict[str, object]) -> None:
    document.add_heading("E.2 三随机种子分区不确定性消融", level=2)
    labels = {
        "rsg_complete": "RSG 完整模型",
        "rpg_complete": "RPG 完整模型",
        "rpg_without_within": "去除角色内不确定性",
        "rpg_without_between": "去除角色间不确定性",
        "rpg_total_entropy_only": "仅使用总熵",
    }
    rows: list[list[str]] = []
    for configuration in RPG_FORMAL_CONFIGURATIONS:
        summary = evidence[configuration]
        rows.append(
            [
                labels[configuration],
                _format_percent(summary["accuracy"]["mean"]),
                _format_percent(summary["macro_f1"]["mean"]),
                _format_percent(summary["target_recall"]["mean"]),
                _format_percent(summary["brier_score"]["mean"]),
                _format_percent(summary["expected_calibration_error"]["mean"]),
            ]
        )
    _add_table(document, ["配置", "Accuracy", "Macro F1", "目标类召回", "Brier", "ECE"], rows)
    _add_body(
        document,
        "完整 RPG、两个分区熵移除消融和仅总熵消融在固定数据划分、主干、训练预算、图像增强和随机种子下进行对比。若 RPG 完整模型相对 RSG 的成对区间未显示稳定优势，本附录仅将其报告为理论驱动的受控探索，而不作性能优越性结论。",
    )


def _add_rpg_paired_evidence(document: Document, analysis_dir: Path) -> None:
    document.add_heading("E.3 成对统计与结论边界", level=2)
    rows: list[list[str]] = []
    for configuration, label in (
        ("rpg_complete", "RPG 完整模型"),
        ("rpg_without_within", "去除角色内不确定性"),
        ("rpg_without_between", "去除角色间不确定性"),
        ("rpg_total_entropy_only", "仅使用总熵"),
    ):
        paired = _load_paired_comparison(analysis_dir, configuration)
        macro = paired["classification"]["macro_f1"]
        routing = paired["routing_regret"]
        brier = paired["calibration"]["brier_score"]
        rows.append(
            [
                label,
                f"{_format_percent(macro['difference'])} [{_format_percent(macro['ci_low'])}, {_format_percent(macro['ci_high'])}]",
                f"{_format_percent(routing['difference'])} [{_format_percent(routing['ci_low'])}, {_format_percent(routing['ci_high'])}]",
                f"{_format_percent(brier['difference'])} [{_format_percent(brier['ci_low'])}, {_format_percent(brier['ci_high'])}]",
            ]
        )
    _add_table(document, ["相对 RSG", "Macro F1 差值 [95% CI]", "门控遗憾差值 [95% CI]", "Brier 差值 [95% CI]"], rows)
    _add_body(
        document,
        "差值均为配置减去 RSG；门控遗憾以融合后验相对于两路专家中概率更优者的负对数似然损失定义，越低越好。区间通过图片 split_group_id 聚类、随机种子与组两阶段重采样的 2,000 次 Bootstrap 得到。理论命题中的信息分解恒等式不依赖实验，实验只检验所构造的门控特征是否在本数据与本协议下带来稳定的经验收益。",
    )
    _add_body(
        document,
        "结论边界：RPG-HRGV 验证的是公开矿物标本图像上的四角色闭集识别、门控行为和后验质量，不等同于工业分选、精矿品位预测、回收率提升、钒钛元素含量检测、真实工况泛化或未知矿物拒识性能。阶段条件化选矿决策图及送检风险最小化作为后续研究方向，不在本附录中声称已完成验证。",
    )


def _add_mrpg_architecture_figure(document: Document) -> None:
    if not MRPG_FIGURE_PATH.is_file():
        raise ValueError(f"M-RPG architecture figure is missing: {MRPG_FIGURE_PATH}")
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(MRPG_FIGURE_PATH), width=Cm(15.4))
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        caption.add_run("图 21 M-RPG-HRGV-Net 的容量归一化单调角色分区门控结构"),
        size=9.5,
    )


def _add_mrpg_theory_statement(document: Document) -> None:
    formulas = render_mrpg_formulas()
    document.add_heading("F.1 容量归一化单调门控", level=2)
    _add_body(
        document,
        "M-RPG-HRGV 在 RPG 的固定矿物种类-角色映射上进一步处理四个角色的种类容量不均衡问题。设 p_s 为 17 类矿物种类后验、p_m=Mp_s 为四角色后验，角色 r 所含种类数为 n_r。种类总不确定性仍满足熵链式分解，但角色内熵的最大允许值随 n_r 变化，不能直接与角色间熵比较。",
    )
    _add_formula(document, formulas["identity"], "式（F-1） 固定角色分区下保留的原始熵链式恒等式")
    _add_formula(document, formulas["normalization"], "式（F-2） 按角色种类容量归一化的两类不确定性")
    _add_body(
        document,
        "将角色间熵除以 log|R|，并将角色内条件熵除以其在当前角色后验下的容量上界 C_within(p_m)，得到均位于 [0,1] 的 u_between 与 u_within。若 C_within 为零，则定义 u_within=0。该归一化不改变式（F-1）的原始信息分解，只为门控比较提供无量纲且容量可比的输入。",
    )
    _add_formula(document, formulas["monotone_gate"], "式（F-3） 对角色间歧义单调的直接专家分配与凸融合")
    _add_formula(document, formulas["derivative"], "式（F-4） 角色间不确定性到直接专家门控的非负导数")
    _add_body(
        document,
        "命题 M-R1（精确分解）：式（F-1）为固定映射下的条件熵链式法则。命题 M-R2（归一化有界性）：u_between 与 u_within 均落在 [0,1]。命题 M-R3（单调直接专家分配）：在共享特征、直接专家熵和 u_within 固定时，式（F-4）非负，角色间歧义增加不会降低直接角色专家的门控权重。命题 M-R4（凸融合包络）：因 g_M 属于 [0,1]，最终每个角色后验仍位于两路专家同角色概率的闭区间内。M-R1 至 M-R4 是当前公式和计算图的确定性性质；模型的经验性能只由下列受控实验判断。",
    )
    _add_mrpg_architecture_figure(document)


def _add_mrpg_three_seed_summary(document: Document, evidence: dict[str, object]) -> None:
    document.add_heading("F.2 三随机种子单调性与容量归一化消融", level=2)
    labels = {
        "rsg_complete": "RSG 完整模型",
        "rpg_complete": "原始 RPG 完整模型",
        "mrpg_complete": "M-RPG 完整模型",
        "mrpg_unconstrained_between": "去除单调约束",
        "mrpg_without_between": "去除角色间不确定性",
    }
    rows: list[list[str]] = []
    for configuration in MRPG_FORMAL_CONFIGURATIONS:
        summary = evidence[configuration]
        rows.append(
            [
                labels[configuration],
                _format_percent(summary["accuracy"]["mean"]),
                _format_percent(summary["macro_f1"]["mean"]),
                _format_percent(summary["target_recall"]["mean"]),
                _format_percent(summary["brier_score"]["mean"]),
                _format_percent(summary["expected_calibration_error"]["mean"]),
            ]
        )
    _add_table(document, ["配置", "Accuracy", "Macro F1", "目标类召回", "Brier", "ECE"], rows)
    _add_body(
        document,
        "完整 M-RPG、无单调约束和无角色间不确定性三个配置使用相同数据划分、主干、训练预算、图像增强及三随机种子。原始 RPG 同时保留为容量归一化前的理论对照。",
    )


def _add_mrpg_paired_evidence(document: Document, analysis_dir: Path) -> None:
    document.add_heading("F.3 成对统计与结论边界", level=2)
    rows: list[list[str]] = []
    for configuration, label in (
        ("mrpg_complete", "M-RPG 完整模型"),
        ("mrpg_unconstrained_between", "去除单调约束"),
        ("mrpg_without_between", "去除角色间不确定性"),
    ):
        paired = _load_paired_comparison(analysis_dir, configuration)
        macro = paired["classification"]["macro_f1"]
        target = paired["classification"]["target_recall"]
        brier = paired["calibration"]["brier_score"]
        rows.append(
            [
                label,
                f"{_format_percent(macro['difference'])} [{_format_percent(macro['ci_low'])}, {_format_percent(macro['ci_high'])}]",
                f"{_format_percent(target['difference'])} [{_format_percent(target['ci_low'])}, {_format_percent(target['ci_high'])}]",
                f"{_format_percent(brier['difference'])} [{_format_percent(brier['ci_low'])}, {_format_percent(brier['ci_high'])}]",
            ]
        )
    _add_table(document, ["相对 RSG", "Macro F1 差值 [95% CI]", "目标类召回差值 [95% CI]", "Brier 差值 [95% CI]"], rows)
    direct_rows: list[list[str]] = []
    for configuration, label in (
        ("mrpg_unconstrained_between", "去除单调约束"),
        ("mrpg_without_between", "去除角色间不确定性"),
    ):
        direct_path = analysis_dir / f"paired_{configuration}_vs_mrpg_complete.json"
        if not direct_path.is_file():
            raise ValueError(f"Formal paired comparison is missing: {direct_path.name}")
        paired = json.loads(direct_path.read_text(encoding="utf-8"))
        macro = paired["classification"]["macro_f1"]
        target = paired["classification"]["target_recall"]
        brier = paired["calibration"]["brier_score"]
        direct_rows.append(
            [
                label,
                f"{_format_percent(macro['difference'])} [{_format_percent(macro['ci_low'])}, {_format_percent(macro['ci_high'])}]",
                f"{_format_percent(target['difference'])} [{_format_percent(target['ci_low'])}, {_format_percent(target['ci_high'])}]",
                f"{_format_percent(brier['difference'])} [{_format_percent(brier['ci_low'])}, {_format_percent(brier['ci_high'])}]",
            ]
        )
    _add_body(document, "相对完整 M-RPG 的直接消融：差值定义为消融配置减去完整 M-RPG。")
    _add_table(document, ["消融配置", "Macro F1 差值 [95% CI]", "目标类召回差值 [95% CI]", "Brier 差值 [95% CI]"], direct_rows)
    _add_body(
        document,
        "每个区间以图片 split_group_id 聚类、随机种子与组两阶段重采样的 2,000 次 Bootstrap 得到。理论命题 M-R1 至 M-R4 不依赖实验；经验结论只以受控三随机种子均值与成对区间为准。若完整模型与直接消融的关键区间跨零，则不能将目标类表现差异归因于单调项或角色间不确定性项。若任何关键区间跨零，则仅报告未观察到稳定经验增益，不能以单种子、单指标或理论性质替代经验性能结论。",
    )
    _add_body(
        document,
        "结论边界：M-RPG-HRGV 只验证公开矿物标本图像上的四角色闭集识别、概率融合和门控性质；不等同于工业分选、精矿品位预测、回收率提升、元素含量检测、真实工况泛化或未知矿物拒识性能。阶段条件化选矿决策图和送检代价决策仍作为后续研究方向。",
    )


def _add_five_seed_extension(document: Document, evidence: tuple[dict[str, object], dict[str, object]]) -> None:
    summary, paired = evidence
    document.add_heading("G.1 预注册的两随机种子扩展", level=2)
    _add_body(
        document,
        "为检验三随机种子观察到的目标类召回趋势是否稳健，在不改变固定数据划分、ImageNet 预训练 EfficientNet-B0 主干、RSG 残差验证器、优化器、图像增强、训练预算和 M-RPG 公式的前提下，预先登记新增随机种子 20260730 与 20260731。该附录将新增两次训练与原三次正式训练合并为五随机种子比较；它是扩展验证，不替代附录 F 的直接组件消融。",
    )
    labels = {
        "rsg_complete": "RSG 完整模型",
        "mrpg_complete": "M-RPG 完整模型",
    }
    rows: list[list[str]] = []
    for configuration in FIVE_SEED_CONFIGURATIONS:
        values = summary[configuration]
        rows.append(
            [
                labels[configuration],
                _format_percent(values["macro_f1"]["mean"]),
                _format_percent(values["target_recall"]["mean"]),
                _format_percent(values["ti_to_target_intrusion_rate"]["mean"]),
                _format_percent(values["metallic_to_target_intrusion_rate"]["mean"]),
                _format_percent(values["brier_score"]["mean"]),
                _format_percent(values["expected_calibration_error"]["mean"]),
            ]
        )
    _add_table(
        document,
        ["配置", "Macro F1", "目标类召回", "含钛类误入目标", "金属光泽类误入目标", "Brier", "ECE"],
        rows,
    )

    document.add_heading("G.2 成对区间与证据边界", level=2)
    classification = paired["classification"]
    calibration = paired["calibration"]
    rows = []
    for label, values in (
        ("Macro F1", classification["macro_f1"]),
        ("目标类召回", classification["target_recall"]),
        ("含钛类误入目标", classification["ti_to_target_intrusion"]),
        ("金属光泽类误入目标", classification["metallic_to_target_intrusion"]),
        ("Brier", calibration["brier_score"]),
        ("ECE", calibration["expected_calibration_error"]),
    ):
        rows.append(
            [
                label,
                f"{_format_percent(values['difference'])} [{_format_percent(values['ci_low'])}, {_format_percent(values['ci_high'])}]",
            ]
        )
    _add_table(document, ["M-RPG 减 RSG", "差值 [95% CI]"], rows)
    _add_body(
        document,
        "差值以 M-RPG 减 RSG 定义；目标类召回的正差值有利，误入目标比例、Brier 和 ECE 的负差值有利。所有区间均以图片 split_group_id 聚类，并在五个随机种子与图片组上进行两阶段重采样的 5,000 次 Bootstrap 得到。扩展后，Macro F1、目标类召回、两类误入目标比例及校准指标的区间均跨零。因此，本研究未观察到 M-RPG 相对 RSG 的稳定经验优势；附录 F 的 M-R1 至 M-R4 仍是公式和计算图的确定性性质，但不能据此替代经验优越性证明。",
    )
    _add_body(
        document,
        "这一扩展验证保留为可复核的反证性结果：理论结构提供了可解释的容量归一化、单调门控和凸融合约束，但在当前公开矿物标本四角色闭集协议下，尚不足以支持总体性能、目标类召回、误入控制或校准性能优于 RSG 的结论。后续需在来源留出或真实外部图像上，结合更多独立重复进一步检验其泛化价值。",
    )


def _add_rsg_theory_replay_consistency(
    document: Document, replay_evidence: dict[str, object]
) -> None:
    document.add_heading("H.2 高精度检查点重放的数值一致性验证", level=2)
    runs = replay_evidence["runs"]
    rows: list[list[str]] = []
    for run in runs:
        rows.append(
            [
                "固定测试" if run["protocol"] == "fixed" else "摄影者留出",
                str(run["seed"]).replace("seed", ""),
                str(run["sample_count"]),
                f"{float(run['minimum_true_probability']):.2e}",
                f"{float(run['b1_max_residual']):.2e}",
                str(run["b1_violation_count"]),
                f"{float(run['b2_max_residual']):.2e}",
                str(run["b2_violation_count"]),
            ]
        )
    _add_table(
        document,
        ["协议", "随机种子", "样本数", "最小真值概率", "B.1 最大残差", "B.1 违反数", "B.2 最大残差", "B.2 违反数"],
        rows,
    )
    overall = replay_evidence["overall"]
    numeric = replay_evidence["numeric_settings"]
    _add_body(
        document,
        f"为避免原始逐图 CSV 六位小数导出把极小概率写为零，使用保存的最佳检查点对固定测试三种子和摄影者留出三种子重新执行只读推理，以 15 位小数导出门控及两路真值概率。共 {overall['run_count']} 次、{overall['sample_count']} 张图像；采用 float32 概率下界 {float(numeric['float32_epsilon']):.2e}、数值容差 {float(numeric['tolerance']):.2e}。定理 B.1 最大残差为 {float(overall['b1_max_residual']):.2e}、违反数为 {overall['b1_violation_count']}；定理 B.2 最大残差为 {float(overall['b2_max_residual']):.2e}、违反数为 {overall['b2_violation_count']}。因此，高精度检查点重放支持当前 RSG-HRGV 公式在已训练模型输出上的数值一致性验证。",
    )
    _add_body(
        document,
        "该验证仅检查定理 B.1 与 B.2 所用概率、门控和软目标在实现中的数值关系；不构成新的分类性能实验，也不支持工业分选、品位、回收率、元素含量、外部泛化或未知矿物拒识结论。",
    )


def _add_rsg_theory_evidence_figure(document: Document, replay_evidence: dict[str, object]) -> None:
    if not RSG_THEORY_EVIDENCE_FIGURE_PATH.is_file():
        raise ValueError(
            "RSG theory-evidence figure is missing: "
            f"{RSG_THEORY_EVIDENCE_FIGURE_PATH}"
        )
    document.add_heading("H.1 定理、计算图与实验证据的对应关系", level=2)
    _add_body(
        document,
        "图 22 将 RSG-HRGV 的三条可证明性质与其直接对应的经验指标并列呈现。定理 B.1 约束门控偏离真实类概率更优专家时的路由后悔；定理 B.2 描述损失差驱动的软门控对硬最优选择的指数逼近；定理 B.3 限定 stop-gradient 后悔分支的局部梯度范围。它们均为当前公式与计算图的性质，而非总体分类精度的先验保证。",
    )
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(RSG_THEORY_EVIDENCE_FIGURE_PATH), width=Cm(15.8))
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        caption.add_run("图 22 RSG-HRGV 的理论性质与平均路由后悔证据对应关系"),
        size=9.5,
    )
    _add_body(
        document,
        "固定测试的 RSG-HRGV 相对 HRGV 平均路由后悔差为 -1.77 个百分点，95% Bootstrap 区间为 [-2.86, -0.69] 个百分点；摄影者留出确认集对应差值为 -3.53 个百分点，区间为 [-4.75, -2.17] 个百分点。两个预定义区间均低于零，支持在本协议下平均路由后悔下降。该结论不主张总体分类性能优越：Accuracy、Macro F1、目标类召回及两类误入目标比例的成对区间仍存在跨零情况。",
    )
    _add_rsg_theory_replay_consistency(document, replay_evidence)
    _add_body(
        document,
        "证据边界：图中的摄影者留出仍来自公开 Mindat 标本图像，不能等同于真实矿石颗粒、现场传送带、工业分选、精矿品位、回收率、元素含量检测、跨矿区泛化或未知矿物拒识。",
    )


def _add_rsg_backbone_portability_evidence(
    document: Document,
    summary: dict[str, object],
    paired: dict[str, object],
    figure_path: Path,
) -> None:
    """Append the architecture-portability check without claiming backbone superiority."""
    if not figure_path.is_file():
        raise ValueError(f"RSG backbone portability figure is missing: {figure_path}")
    document.add_heading(BACKBONE_PORTABILITY_HEADING, level=2)
    _add_body(
        document,
        "主干替换不变性命题是融合层的条件性结论：只要不同视觉主干经两位专家输出合法概率分布，并保持同一凸门控、软后悔目标和 stop-gradient 计算图，定理 B.1--B.3 的推导不依赖卷积或残差模块的具体形式。该命题不等同于任意主干具有相同或更高的分类性能，因此以 ImageNet 预训练 ResNet50 在完全匹配的数据划分、训练预算、随机种子和比较配置下进行三随机种子确认。",
    )
    rows: list[list[str]] = []
    for configuration, label in (
        ("hrgv_reference", "HRGV 参考模型"),
        ("rsg_complete", "RSG-HRGV"),
    ):
        values = summary[configuration]
        rows.append(
            [
                label,
                _format_percent(float(values["accuracy"]["mean"])),
                _format_percent(float(values["macro_f1"]["mean"])),
                _format_percent(float(values["mean_routing_regret_nll"]["mean"])),
            ]
        )
    _add_table(document, ["ResNet50 配置", "Accuracy", "Macro F1", "平均路由后悔"], rows)
    regret = paired["routing_regret"]
    macro_f1 = paired["classification"]["macro_f1"]
    accuracy = paired["classification"]["accuracy"]
    _add_body(
        document,
        "RSG-HRGV 减 HRGV 的平均路由后悔差为 "
        f"{_format_percent(float(regret['difference']))}，95% Bootstrap 区间为 "
        f"[{_format_percent(float(regret['ci_low']))}, {_format_percent(float(regret['ci_high']))}]。"
        "若该区间低于零，则支持在 ResNet50 协议下复现所定义的路由后悔下降；否则只记录为未复现的边界结果。"
    )
    classification_boundary = (
        "Macro F1 和 Accuracy 的成对区间均未显示稳定差异，故本节不主张分类性能优越。"
        if float(macro_f1["ci_low"]) <= 0 <= float(macro_f1["ci_high"])
        and float(accuracy["ci_low"]) <= 0 <= float(accuracy["ci_high"])
        else "分类指标的区间应与表中结果一并解读；本节仅用于检验融合层理论的经验边界，不用于主张主干性能优越。"
    )
    _add_body(
        document,
        "图 23 将固定 EfficientNet-B0 证据、摄影者留出证据和 ResNet50 跨主干确认并列。"
        f"{classification_boundary}"
    )
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(figure_path), width=Cm(15.8))
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        caption.add_run("图 23 RSG-HRGV 理论证据与 ResNet50 跨主干确认"),
        size=9.5,
    )
    _add_body(
        document,
        "证据边界：该跨主干检查验证的是 RSG 融合层公式在 ResNet50 输出契约下的经验一致性，"
        "不等同于主干性能优越，也不构成工业分选、品位、回收率、元素含量、真实外部泛化或未知矿物拒识结论。",
    )


def _add_rsg_backbone_replay_consistency(
    document: Document, replay_evidence: dict[str, object]
) -> None:
    """Append formula-level numerical evidence for the ResNet50 portability check."""
    document.add_heading(BACKBONE_REPLAY_HEADING, level=2)
    _add_body(
        document,
        "为排除常规逐图 CSV 概率舍入对跨主干公式检查的影响，使用 ResNet50 三个正式 RSG-HRGV 最佳检查点重新执行只读测试推理，并以 15 位小数导出门控和两路真值概率。该重放不更新任何权重，也不改变主干替换实验的分类结果。",
    )
    rows: list[list[str]] = []
    for run in replay_evidence["runs"]:
        rows.append(
            [
                str(run["seed"]).replace("seed", ""),
                str(run["sample_count"]),
                f"{float(run['minimum_true_probability']):.2e}",
                f"{float(run['b1_max_residual']):.2e}",
                str(run["b1_violation_count"]),
                f"{float(run['b2_max_residual']):.2e}",
                str(run["b2_violation_count"]),
            ]
        )
    _add_table(
        document,
        ["ResNet50 种子", "样本数", "最小真值概率", "B.1 最大残差", "B.1 违反数", "B.2 最大残差", "B.2 违反数"],
        rows,
    )
    overall = replay_evidence["overall"]
    numeric = replay_evidence["numeric_settings"]
    _add_body(
        document,
        f"共 {overall['run_count']} 次、{overall['sample_count']} 张图像；采用 float32 概率下界 {float(numeric['float32_epsilon']):.2e} 和数值容差 {float(numeric['tolerance']):.2e}。B.1 最大残差为 {float(overall['b1_max_residual']):.2e}、B.1 违反数为 {overall['b1_violation_count']}；B.2 最大残差为 {float(overall['b2_max_residual']):.2e}、B.2 违反数为 {overall['b2_violation_count']}。这说明 B.1 与 B.2 的数值关系也在 ResNet50 的高精度输出上成立。",
    )
    _add_body(
        document,
        "该结果补强的是主干替换不变性命题的公式实现证据，不构成新的分类性能实验，不推出任意视觉主干上的性能提升，也不构成工业分选、品位、回收率、元素含量、外部泛化或未知矿物拒识结论。",
    )


def _add_rsg_gate_reliability_diagnosis(
    document: Document, evidence: dict[str, object], figure_path: Path
) -> None:
    """Append a descriptive figure that connects B.1/B.2 to observed replay strata."""
    if not figure_path.is_file():
        raise ValueError(f"RSG gate-reliability figure is missing: {figure_path}")
    document.add_heading(GATE_RELIABILITY_HEADING, level=2)
    _add_body(
        document,
        "在定理 B.1 中，对每张图像取局部下界 epsilon_i=min(a_i,b_i)，即可得到不弱于统一 float32 下界的逐图推论：路由后悔不超过 |g_i-g_{o,i}| |a_i-b_i| / epsilon_i。为避免把这一推论仅作为符号公式，本节按该局部上界的排序三分位数汇总实际路由后悔；同时按专家真值概率对数差 |log a_i-log b_i| 的排序三分位数汇总 B.2 的软目标偏差及其指数界。",
    )
    labels = {
        "fixed": "固定测试",
        "photographer_holdout": "摄影者留出",
        "resnet50_portability": "ResNet50 跨主干",
    }
    rows: list[list[str]] = []
    for protocol in ("fixed", "photographer_holdout", "resnet50_portability"):
        values = evidence["protocols"][protocol]
        rows.append(
            [
                labels[protocol],
                str(values["sample_count"]),
                _format_percent(float(values["mean_routing_regret_nll"])),
                _format_percent(float(values["mean_b1_local_bound"])),
                f"{float(values['mean_soft_hard_deviation']):.3f}",
                f"{float(values['mean_b2_bound']):.3f}",
                f"{values['b1_local_violation_count']} / {values['b2_violation_count']}",
            ]
        )
    _add_table(
        document,
        ["协议", "图像数", "平均路由后悔", "平均局部 B.1 上界", "平均软硬偏差", "平均 B.2 指数界", "B.1 / B.2 违反数"],
        rows,
    )
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(figure_path), width=Cm(15.8))
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        caption.add_run("图 24 RSG-HRGV 门控可靠性分层诊断：局部 B.1 上界与 B.2 指数界"),
        size=9.5,
    )
    overall = evidence["overall"]
    _add_body(
        document,
        f"高精度重放共覆盖 {overall['run_count']} 次、{int(overall['sample_count']):,} 张图像。三种协议中，按局部 B.1 上界从低到高分层后，平均路由后悔同步升高；按专家证据差距从低到高分层后，软目标对硬最优门控的平均偏差下降，且各层 B.1 与 B.2 的违反数均为 0。这些现象与定理的条件性方向一致，但本节只是机制诊断，不替代 Accuracy、Macro F1 或工业指标的比较。",
    )
    _add_body(
        document,
        "证据边界：局部 epsilon_i 推论和分层曲线用于解释已训练模型的门控行为，不构成新增训练或新的分类性能实验；不得据此宣称工业分选、品位、回收率、元素含量、真实外部泛化或未知矿物拒识能力。",
    )


def update_report(
    input_path: Path,
    output_path: Path,
    analysis_dir: Path = DEFAULT_ANALYSIS_DIR,
    rpg_analysis_dir: Path | None = None,
    mrpg_analysis_dir: Path | None = None,
    five_seed_analysis_dir: Path | None = None,
    theory_replay_analysis_dir: Path | None = DEFAULT_THEORY_REPLAY_ANALYSIS_DIR,
    backbone_portability_analysis_dir: Path | None = None,
    backbone_portability_figure_path: Path = BACKBONE_PORTABILITY_FIGURE_PATH,
    backbone_replay_analysis_dir: Path | None = None,
    gate_reliability_analysis_dir: Path | None = None,
    gate_reliability_figure_path: Path = GATE_RELIABILITY_FIGURE_PATH,
) -> Path:
    document = Document(input_path)
    update_primary_contribution_statement(document)
    headings = {paragraph.text.strip() for paragraph in document.paragraphs}
    if APPENDIX_HEADING not in headings:
        evidence = load_formal_cgdc_evidence(analysis_dir)
        document.add_heading(APPENDIX_HEADING, level=1)
        _add_theory_statement(document)
        _add_three_seed_summary(document, evidence)
        _add_paired_evidence(document, analysis_dir)
    if rpg_analysis_dir is not None and RPG_APPENDIX_HEADING not in headings:
        rpg_evidence = load_formal_rpg_evidence(rpg_analysis_dir)
        document.add_heading(RPG_APPENDIX_HEADING, level=1)
        _add_rpg_theory_statement(document)
        _add_rpg_three_seed_summary(document, rpg_evidence)
        _add_rpg_paired_evidence(document, rpg_analysis_dir)
    if mrpg_analysis_dir is not None and MRPG_APPENDIX_HEADING not in headings:
        mrpg_evidence = load_formal_mrpg_evidence(mrpg_analysis_dir)
        document.add_heading(MRPG_APPENDIX_HEADING, level=1)
        _add_mrpg_theory_statement(document)
        _add_mrpg_three_seed_summary(document, mrpg_evidence)
        _add_mrpg_paired_evidence(document, mrpg_analysis_dir)
    if five_seed_analysis_dir is not None and FIVE_SEED_APPENDIX_HEADING not in headings:
        five_seed_evidence = load_five_seed_extension_evidence(five_seed_analysis_dir)
        document.add_heading(FIVE_SEED_APPENDIX_HEADING, level=1)
        _add_five_seed_extension(document, five_seed_evidence)
    if RSG_THEORY_EVIDENCE_APPENDIX_HEADING not in headings:
        if theory_replay_analysis_dir is None:
            raise ValueError("RSG theory replay analysis directory is required for appendix H.")
        replay_evidence = load_rsg_theory_replay_evidence(theory_replay_analysis_dir)
        document.add_heading(RSG_THEORY_EVIDENCE_APPENDIX_HEADING, level=1)
        _add_rsg_theory_evidence_figure(document, replay_evidence)
    elif "H.2 高精度检查点重放的数值一致性验证" not in headings:
        if theory_replay_analysis_dir is None:
            raise ValueError("RSG theory replay analysis directory is required for appendix H.")
        _add_rsg_theory_replay_consistency(
            document, load_rsg_theory_replay_evidence(theory_replay_analysis_dir)
        )
    if backbone_portability_analysis_dir is not None and BACKBONE_PORTABILITY_HEADING not in headings:
        portability_summary, portability_paired = load_rsg_backbone_portability_evidence(
            backbone_portability_analysis_dir
        )
        _add_rsg_backbone_portability_evidence(
            document,
            portability_summary,
            portability_paired,
            backbone_portability_figure_path,
        )
    if backbone_replay_analysis_dir is not None and BACKBONE_REPLAY_HEADING not in headings:
        _add_rsg_backbone_replay_consistency(
            document, load_rsg_backbone_replay_evidence(backbone_replay_analysis_dir)
        )
    if gate_reliability_analysis_dir is not None and GATE_RELIABILITY_HEADING not in headings:
        _add_rsg_gate_reliability_diagnosis(
            document,
            load_rsg_gate_reliability_evidence(gate_reliability_analysis_dir),
            gate_reliability_figure_path,
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Append CGDC formal evidence to the technical report.")
    parser.add_argument("--input", type=Path, default=FORMAL_REPORT)
    parser.add_argument("--output", type=Path, default=FORMAL_REPORT)
    parser.add_argument("--analysis-dir", type=Path, default=DEFAULT_ANALYSIS_DIR)
    parser.add_argument("--rpg-analysis-dir", type=Path)
    parser.add_argument("--mrpg-analysis-dir", type=Path)
    parser.add_argument("--five-seed-analysis-dir", type=Path)
    parser.add_argument("--theory-replay-analysis-dir", type=Path, default=DEFAULT_THEORY_REPLAY_ANALYSIS_DIR)
    parser.add_argument("--backbone-portability-analysis-dir", type=Path)
    parser.add_argument("--backbone-portability-figure-path", type=Path, default=BACKBONE_PORTABILITY_FIGURE_PATH)
    parser.add_argument("--backbone-replay-analysis-dir", type=Path)
    parser.add_argument("--gate-reliability-analysis-dir", type=Path)
    parser.add_argument("--gate-reliability-figure-path", type=Path, default=GATE_RELIABILITY_FIGURE_PATH)
    args = parser.parse_args()
    print(
        update_report(
            args.input.resolve(),
            args.output.resolve(),
            args.analysis_dir.resolve(),
            args.rpg_analysis_dir.resolve() if args.rpg_analysis_dir else None,
            args.mrpg_analysis_dir.resolve() if args.mrpg_analysis_dir else None,
            args.five_seed_analysis_dir.resolve() if args.five_seed_analysis_dir else None,
            args.theory_replay_analysis_dir.resolve() if args.theory_replay_analysis_dir else None,
            args.backbone_portability_analysis_dir.resolve()
            if args.backbone_portability_analysis_dir
            else None,
            args.backbone_portability_figure_path.resolve(),
            args.backbone_replay_analysis_dir.resolve()
            if args.backbone_replay_analysis_dir
            else None,
            args.gate_reliability_analysis_dir.resolve()
            if args.gate_reliability_analysis_dir
            else None,
            args.gate_reliability_figure_path.resolve(),
        )
    )


if __name__ == "__main__":
    main()
