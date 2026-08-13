from __future__ import annotations

import argparse
import json
from pathlib import Path
from statistics import mean, stdev

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FORMAL_REPORT = ROOT / "结题" / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（正式版）.docx"
ASSET_DIR = ROOT / "outputs" / "report_assets_v2"
HIERARCHY_JSON = ROOT / "outputs" / "theory_validation" / "hierarchy_consistency" / "hierarchy_consistency_summary.json"
CALIBRATION_JSON = ROOT / "outputs" / "theory_validation" / "calibrated_selective_recognition" / "calibrated_selective_recognition_summary.json"
EXPERIMENT_JSON = ROOT / "outputs" / "paper_experiments_v2" / "training_summaries" / "paper_experiment_summary.json"
SOURCE_FIGURE = ROOT / "outputs" / "paper_figures_v2" / "fig_source_holdout_comparison.png"
KL_FIGURE = ROOT / "outputs" / "paper_figures_v2" / "fig_hierarchy_kl_disagreement.png"
CALIBRATION_FIGURE = ROOT / "outputs" / "paper_figures_v2" / "fig_calibrated_reliability.png"

FONT_NAME = "宋体"
HEADING_FONT = "黑体"
FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")


def _read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def _set_run_font(run, name=FONT_NAME, size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
    tc_w.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _format_table(table, widths):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                _set_cell_shading(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(1.5)
                paragraph.paragraph_format.space_after = Pt(1.5)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    _set_run_font(run, HEADING_FONT if row_index == 0 else FONT_NAME, 9.2, row_index == 0)


def _move_before(block, anchor):
    anchor._p.addprevious(block._element)


def _add_paragraph_before(doc, anchor, text, style=None, bold_lead=None):
    paragraph = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        _set_run_font(lead, HEADING_FONT, 10.5, True)
        rest = paragraph.add_run(text[len(bold_lead):])
        _set_run_font(rest, FONT_NAME, 10.5)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run, FONT_NAME, 10.5)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(4)
    _move_before(paragraph, anchor)
    return paragraph


def _add_heading_before(doc, anchor, text, level=2):
    heading = doc.add_heading(text, level=level)
    _move_before(heading, anchor)
    return heading


def _add_caption_before(doc, anchor, text):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    _set_run_font(run, FONT_NAME, 9.5)
    _move_before(paragraph, anchor)
    return paragraph


def _add_picture_before(doc, anchor, path: Path, width_cm: float, caption: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    _move_before(paragraph, anchor)
    _add_caption_before(doc, anchor, caption)


def _add_table_before(doc, anchor, headers, rows, widths, caption):
    _add_caption_before(doc, anchor, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = str(value)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    _format_table(table, widths)
    anchor._p.addprevious(table._element)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    _move_before(spacer, anchor)
    return table


def render_equations(output_dir: Path):
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    equations = {
        "risk_contraction": (
            r"$\mathbf{1}\!\left[g(\hat{s})\ne g(s)\right]\leq "
            r"\mathbf{1}\!\left[\hat{s}\ne s\right],\qquad "
            r"\widehat{R}_{role}=\widehat{R}_{species}-\frac{1}{N}"
            r"\sum_i\mathbf{1}\!\left[\hat{s}_i\ne s_i,\ g(\hat{s}_i)=g(s_i)\right]$"
        ),
        "pinsker": (
            r"$\mathcal{L}_{cons}=D_{KL}(\widetilde{\mathbf{p}}_r\Vert\mathbf{p}_r),"
            r"\qquad \mathrm{TV}(\widetilde{\mathbf{p}}_r,\mathbf{p}_r)"
            r"\leq\sqrt{\frac{1}{2}D_{KL}(\widetilde{\mathbf{p}}_r\Vert\mathbf{p}_r)}$"
        ),
        "calibration": (
            r"$T^{*}=\arg\min_{T>0}\sum_{(x_i,y_i)\in\mathcal{V}_{fit}}"
            r"-\log\operatorname{softmax}(\mathbf{z}_i/T)_{y_i},\qquad "
            r"\mathbf{p}_{T}(x)=\operatorname{softmax}(\mathbf{z}(x)/T^{*})$"
        ),
        "certificate": (
            r"$U_{\tau}=\operatorname{Beta}^{-1}\!\left(1-\alpha/|\mathcal{T}|;"
            r"e_{\tau}+1,n_{\tau}-e_{\tau}\right),\qquad "
            r"\tau^{*}=\arg\max_{\tau\in\mathcal{T}}\mathrm{Coverage}(\tau)"
            r"\ \ \mathrm{s.t.}\ \ U_{\tau}\leq\delta$"
        ),
    }
    output_dir.mkdir(parents=True, exist_ok=True)
    paths = {}
    for name, equation in equations.items():
        figure = plt.figure(figsize=(12.8, 1.15), dpi=220, facecolor="white")
        figure.text(0.5, 0.5, equation, ha="center", va="center", fontsize=18, color="#1F2937")
        path = output_dir / f"{name}.png"
        figure.savefig(path, bbox_inches="tight", pad_inches=0.12, facecolor="white")
        plt.close(figure)
        paths[name] = path
    return paths


def render_architecture(path: Path):
    width, height = 2400, 1400
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype(str(FONT_PATH), 42)
    small = ImageFont.truetype(str(FONT_PATH), 34)
    tiny = ImageFont.truetype(str(FONT_PATH), 28)
    title = ImageFont.truetype(str(FONT_PATH), 50)

    def box(x1, y1, x2, y2, text, fill, outline="#315A7D", font_obj=small):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=fill, outline=outline, width=4)
        lines = text.split("\n")
        heights = [draw.textbbox((0, 0), line, font=font_obj)[3] for line in lines]
        y = (y1 + y2 - sum(heights) - (len(lines) - 1) * 8) / 2
        for line, line_height in zip(lines, heights):
            bbox = draw.textbbox((0, 0), line, font=font_obj)
            draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, fill="#17212B", font=font_obj)
            y += line_height + 8

    def arrow(x1, y1, x2, y2, color="#315A7D", width_px=5):
        draw.line((x1, y1, x2, y2), fill=color, width=width_px)
        import math
        angle = math.atan2(y2 - y1, x2 - x1)
        size = 18
        points = [(x2, y2),
                  (x2 - size * math.cos(angle - 0.55), y2 - size * math.sin(angle - 0.55)),
                  (x2 - size * math.cos(angle + 0.55), y2 - size * math.sin(angle + 0.55))]
        draw.polygon(points, fill=color)

    draw.text((1200, 65), "理论感知的分层矿物识别与风险认证结构", fill="#153A5B", font=title, anchor="mm")
    box(60, 500, 270, 720, "输入图像\n224×224", "#EEF5FA", font_obj=font)
    box(350, 420, 910, 800, "EfficientNet-B0 共享主干\nStem 3×3 Conv\nMBConv1/MBConv6 多尺度块\n1×1 Conv + GAP + Dropout\n1280维共享特征", "#DCEEF8", font_obj=small)
    arrow(270, 610, 350, 610)

    heads = [
        (1050, 180, 1420, 370, "角色头\n4类 Softmax", "#E8F3E8"),
        (1050, 430, 1420, 620, "种类头\n17类 Softmax", "#FFF1D6"),
        (1050, 680, 1420, 870, "目标/非目标头\n2类 Softmax", "#F6E8F3"),
        (1050, 930, 1420, 1120, "投影头\n128维归一化", "#ECE8F8"),
    ]
    for x1, y1, x2, y2, text, fill in heads:
        box(x1, y1, x2, y2, text, fill)
        arrow(910, 610, x1, (y1 + y2) // 2)

    box(1530, 400, 1900, 650, "固定映射矩阵 A\n种类概率 → 角色概率\np_map = A p_species", "#FFF8E8")
    arrow(1420, 525, 1530, 525)
    box(1980, 310, 2330, 555, "层级一致性\nDKL(p_map || p_role)\nPinsker 上界", "#FFE5E5")
    arrow(1900, 525, 1980, 455)
    arrow(1420, 275, 1980, 390)

    box(1530, 720, 1900, 930, "困难负样本关系\n目标与含钛干扰\n目标与金属光泽干扰", "#F0EAF8", font_obj=tiny)
    arrow(1420, 1025, 1530, 825)

    box(1530, 1050, 1900, 1280, "验证集温度缩放\nT* 最小化 NLL\n校准角色概率", "#E8F4EE", font_obj=tiny)
    arrow(1420, 275, 1640, 1050)
    box(1980, 1030, 2330, 1300, "独立风险认证\nCP 上界 + Bonferroni\n满足风险约束：输出\n否则：送检/拒识", "#DCEEF8", font_obj=tiny)
    arrow(1900, 1165, 1980, 1165)

    draw.text((80, 1320), "训练：四类角色 + 17类种类 + 二分类 + 困难负样本 + 层级一致性；推理：温度校准 + 预设风险约束选择性输出", fill="#334155", font=tiny)
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, quality=96)
    return path


def certificate_rows(calibration):
    rows = []
    for delta in ("0.1", "0.15", "0.2"):
        certified = []
        for seed, result in calibration["results"].items():
            certificate = result["certificates"][delta]
            if certificate["status"] == "certified":
                test = certificate["test_evaluation"]
                certified.append((seed, certificate["selected"]["threshold"], test["coverage"], test["selective_risk"]))
        if certified:
            rows.append((f"δ={float(delta):.2f}", f"{len(certified)}/3",
                         ", ".join(f"{threshold:.2f}" for _, threshold, _, _ in certified),
                         f"{100*mean(x[2] for x in certified):.2f}%",
                         f"{100*mean(x[3] for x in certified):.2f}%"))
        else:
            rows.append((f"δ={float(delta):.2f}", "0/3", "—", "—", "—"))
    return rows


def replace_paragraph_text(paragraph, text):
    paragraph.clear()
    run = paragraph.add_run(text)
    _set_run_font(run, FONT_NAME, 10.5)
    paragraph.paragraph_format.line_spacing = 1.5


def update_report(input_path: Path, output_path: Path):
    hierarchy = _read_json(HIERARCHY_JSON)
    calibration = _read_json(CALIBRATION_JSON)
    experiments = _read_json(EXPERIMENT_JSON)
    equations = render_equations(ASSET_DIR / "equations")
    architecture = render_architecture(ASSET_DIR / "fig_theory_aware_architecture_v2.png")

    document = Document(input_path)
    original_paragraphs = list(document.paragraphs)
    original_tables = list(document.tables)
    usable_width = min(
        section.page_width - section.left_margin - section.right_margin
        for section in document.sections
    )
    for shape in document.inline_shapes:
        if shape.width > usable_width:
            ratio = usable_width / shape.width
            shape.width = usable_width
            shape.height = int(shape.height * ratio)
    if any(p.text.strip() == "5.11 角色风险收缩：种类错误不必然导致角色错误" for p in document.paragraphs):
        formula_renumbering = {
            "式（16） 固定映射下的角色风险收缩关系": "式（7） 固定映射下的角色风险收缩关系",
            "式（17） 层级一致性损失与总变差上界": "式（8） 层级一致性损失与总变差上界",
            "式（18） 验证集温度缩放": "式（9） 验证集温度缩放",
            "式（19） 多阈值校正后的选择性风险证书": "式（10） 多阈值校正后的选择性风险证书",
        }
        for paragraph in document.paragraphs:
            if paragraph.text.strip() in formula_renumbering:
                replace_paragraph_text(paragraph, formula_renumbering[paragraph.text.strip()])
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path

    headings = {p.text.strip(): p for p in document.paragraphs if p.text.strip()}
    anchor = headings["5.11 实验结果的综合判断"]
    replace_paragraph_text(anchor, "5.16 实验结果的综合判断")

    full = hierarchy["aggregate"]["full"]
    no_consistency = hierarchy["aggregate"]["no_consistency"]
    source = experiments["source_holdout"]
    proxy = experiments["proxy_ablation_common_subset"]
    cal_agg = calibration["aggregate"]

    _add_heading_before(document, anchor, "5.11 角色风险收缩：种类错误不必然导致角色错误", 2)
    _add_paragraph_before(
        document, anchor,
        "设细粒度矿物种类为 s，固定的种类—选矿角色映射为 g。由于正确种类必然映射到正确角色，而错误种类若仍落在同一角色内并不会造成角色决策错误，可得到逐样本风险支配关系及其经验风险分解："
    )
    _add_picture_before(document, anchor, equations["risk_contraction"], 15.5, "式（7） 固定映射下的角色风险收缩关系")
    _add_table_before(
        document, anchor,
        ["指标", "三种子均值", "样本标准差", "解释"],
        [
            ["17类种类准确率", f"{100*full['species_accuracy']['mean']:.2f}%", f"{100*full['species_accuracy']['sample_std']:.2f}%", "细粒度识别"],
            ["直接角色头准确率", f"{100*full['direct_role_accuracy']['mean']:.2f}%", f"{100*full['direct_role_accuracy']['sample_std']:.2f}%", "四类角色决策"],
            ["种类预测映射角色准确率", f"{100*full['hard_species_role_accuracy']['mean']:.2f}%", f"{100*full['hard_species_role_accuracy']['sample_std']:.2f}%", "g(ŝ) 的硬映射"],
            ["经验角色风险收缩", f"{100*full['empirical_hard_role_risk_contraction']['mean']:.2f}个百分点", f"{100*full['empirical_hard_role_risk_contraction']['sample_std']:.2f}", "种类错但角色仍对"],
            ["种类错误中角色仍正确", f"{100*full['species_wrong_role_correct_rate_among_species_errors']['mean']:.2f}%", f"{100*full['species_wrong_role_correct_rate_among_species_errors']['sample_std']:.2f}%", "同角色混淆占比"],
        ], [4.3, 3.1, 3.0, 5.0], "表 18 固定映射下角色风险收缩的实证验证"
    )
    _add_paragraph_before(
        document, anchor,
        "结果表明，种类分类平均错误率约为43.20%，但其中43.41%的种类错误没有改变上层选矿角色，最终形成18.74±1.29个百分点的经验角色风险收缩。该结果验证的是固定映射所诱导的决策层风险收缩，不代表细粒度矿物识别问题已解决。"
    )

    _add_heading_before(document, anchor, "5.12 双头分布一致性与 Pinsker 上界", 2)
    _add_paragraph_before(
        document, anchor,
        "种类头经固定矩阵 A 聚合得到角色分布 p̃r，角色头直接输出 pr。训练中实际采用 DKL(p̃r||pr) 约束两者。Pinsker 不等式给出分布差异的可计算上界："
    )
    _add_picture_before(document, anchor, equations["pinsker"], 14.8, "式（8） 层级一致性损失与总变差上界")
    _add_table_before(
        document, anchor,
        ["设置", "平均 KL", "平均 TV", "双头预测分歧率"],
        [
            ["完整分层模型", f"{full['mean_kl_mapped_to_direct']['mean']:.4f}±{full['mean_kl_mapped_to_direct']['sample_std']:.4f}", f"{full['mean_total_variation']['mean']:.4f}±{full['mean_total_variation']['sample_std']:.4f}", f"{100*full['head_disagreement_rate']['mean']:.2f}%±{100*full['head_disagreement_rate']['sample_std']:.2f}%"],
            ["移除一致性项", f"{no_consistency['mean_kl_mapped_to_direct']['mean']:.4f}±{no_consistency['mean_kl_mapped_to_direct']['sample_std']:.4f}", f"{no_consistency['mean_total_variation']['mean']:.4f}±{no_consistency['mean_total_variation']['sample_std']:.4f}", f"{100*no_consistency['head_disagreement_rate']['mean']:.2f}%±{100*no_consistency['head_disagreement_rate']['sample_std']:.2f}%"],
        ], [4.0, 4.0, 4.0, 4.0], "表 19 层级一致性约束的分布效果"
    )
    _add_picture_before(document, anchor, KL_FIGURE, 15.2, "图 13 一致性约束对 KL 散度和双头分歧的影响")
    _add_paragraph_before(
        document, anchor,
        "一致性项使平均 KL 由0.0565降至0.0426，平均 TV 由0.0752降至0.0661，双头预测分歧率由7.42%降至6.36%。三种子上 KL 均下降，但角色准确率没有同步稳定提高，因此该约束的可辩护贡献是提高层级分布一致性，而不是保证分类精度提升。"
    )

    _add_heading_before(document, anchor, "5.13 温度校准与独立风险认证", 2)
    _add_paragraph_before(
        document, anchor,
        "为避免在测试集上反复选择阈值，将原验证集按 split_group_id 分组并按角色分层拆为温度拟合集与风险认证集，二者无分组重叠。温度参数仅在拟合集上最小化 NLL："
    )
    _add_picture_before(document, anchor, equations["calibration"], 15.5, "式（9） 验证集温度缩放")
    _add_paragraph_before(
        document, anchor,
        "对预先给定的21个阈值同时计算单侧 Clopper–Pearson 上界，并采用 Bonferroni 校正控制多阈值选择误差；随后在认证集上选择满足风险上界 δ 的最大覆盖率阈值："
    )
    _add_picture_before(document, anchor, equations["certificate"], 15.8, "式（10） 多阈值校正后的选择性风险证书")
    _add_table_before(
        document, anchor,
        ["指标", "校准前", "校准后", "变化"],
        [
            ["测试集 ECE（15箱）", f"{cal_agg['test_ece_before']['mean']:.4f}±{cal_agg['test_ece_before']['sample_std']:.4f}", f"{cal_agg['test_ece_after']['mean']:.4f}±{cal_agg['test_ece_after']['sample_std']:.4f}", "降低"],
            ["测试集 NLL", f"{cal_agg['test_nll_before']['mean']:.4f}±{cal_agg['test_nll_before']['sample_std']:.4f}", f"{cal_agg['test_nll_after']['mean']:.4f}±{cal_agg['test_nll_after']['sample_std']:.4f}", "降低"],
            ["温度 T*", "—", f"{cal_agg['temperature']['mean']:.3f}±{cal_agg['temperature']['sample_std']:.3f}", "验证集拟合"],
            ["测试集 AURC", "—", f"{cal_agg['test_aurc']['mean']:.4f}±{cal_agg['test_aurc']['sample_std']:.4f}", "越低越好"],
        ], [4.0, 4.0, 4.0, 3.8], "表 20 温度校准与选择性识别总体指标"
    )
    _add_table_before(
        document, anchor,
        ["风险约束", "取得证书种子", "认证阈值", "测试平均覆盖率", "测试平均风险"],
        certificate_rows(calibration), [3.0, 3.0, 3.8, 3.4, 3.4], "表 21 独立认证集上的风险证书及一次性测试结果"
    )
    _add_picture_before(document, anchor, CALIBRATION_FIGURE, 15.0, "图 14 温度校准前后的可靠性曲线")
    _add_paragraph_before(
        document, anchor,
        "温度缩放将测试 ECE 从0.1131降至0.0342，NLL 从0.8399降至0.6633。在δ=0.15和δ=0.20时三个种子均取得验证集风险证书；δ=0.10仅一个种子取得证书，说明严格低风险要求仍受校准样本规模和模型稳定性限制。证书只约束当前公开图像分布，不等同于工业回收率或真实送检成本最优。"
    )

    _add_heading_before(document, anchor, "5.14 摄影者留出条件下的来源外泛化", 2)
    _add_paragraph_before(
        document, anchor,
        "为检验模型是否依赖摄影背景和布光风格，从元数据中筛出5,639张具有摄影者字段的样本，以摄影者与重复图像组的并查集作为不可拆分单元，构建3,947/846/846张的训练、验证和测试划分。三个划分的摄影者及 split_group_id 均无交叉；缺失摄影者的2,890张样本不进入该严格实验。"
    )
    source_baseline = source["aggregate"]["baseline"]
    source_hier = source["aggregate"]["hierarchical"]
    _add_table_before(
        document, anchor,
        ["模型", "Accuracy", "Macro F1", "目标召回", "金属光泽干扰召回"],
        [
            ["EfficientNet-B0", f"{100*source_baseline['accuracy']['mean']:.2f}%±{100*source_baseline['accuracy']['sample_std']:.2f}%", f"{100*source_baseline['macro_f1']['mean']:.2f}%±{100*source_baseline['macro_f1']['sample_std']:.2f}%", f"{100*source_baseline['target_recall']['mean']:.2f}%±{100*source_baseline['target_recall']['sample_std']:.2f}%", f"{100*source_baseline['role_recall']['metallic_hard_negative']['mean']:.2f}%±{100*source_baseline['role_recall']['metallic_hard_negative']['sample_std']:.2f}%"],
            ["分层一致性模型", f"{100*source_hier['accuracy']['mean']:.2f}%±{100*source_hier['accuracy']['sample_std']:.2f}%", f"{100*source_hier['macro_f1']['mean']:.2f}%±{100*source_hier['macro_f1']['sample_std']:.2f}%", f"{100*source_hier['target_recall']['mean']:.2f}%±{100*source_hier['target_recall']['sample_std']:.2f}%", f"{100*source_hier['role_recall']['metallic_hard_negative']['mean']:.2f}%±{100*source_hier['role_recall']['metallic_hard_negative']['sample_std']:.2f}%"],
        ], [4.0, 3.0, 3.0, 3.0, 3.4], "表 22 摄影者留出测试的三随机种子结果"
    )
    _add_picture_before(document, anchor, SOURCE_FIGURE, 15.2, "图 15 摄影者留出测试中的基线与分层模型比较")
    paired = source["paired_hierarchical_minus_baseline"]
    _add_paragraph_before(
        document, anchor,
        f"在更严格的摄影者留出测试中，分层模型相对基线的 Macro F1 平均提高{100*paired['macro_f1']['mean']:.2f}个百分点，目标召回提高{100*paired['target_recall']['mean']:.2f}个百分点，金属光泽干扰召回提高{100*paired['role_recall']['metallic_hard_negative']['mean']:.2f}个百分点。由于只有三个随机种子且数据仍来自 Mindat，该结果支持来源外条件下的稳定改进趋势，但不构成统计显著性或跨网站泛化结论。"
    )

    _add_heading_before(document, anchor, "5.15 普通磁铁矿代理样本消融", 2)
    _add_paragraph_before(
        document, anchor,
        "35张钛磁铁矿下载图像中，12张因跨标签完全重复或近重复冲突被排除，最终数据集保留23张。为区分普通磁铁矿代理的训练价值与钛磁铁矿真实性，移除1,316张普通磁铁矿后重新训练分层模型，并在与完整模型共有的固定测试子集上比较。"
    )
    proxy_full = proxy["aggregate"]["full_common_subset"]
    proxy_none = proxy["aggregate"]["no_magnetite_proxy"]
    ilmenite = proxy["target_species"]["ilmenite"]
    titanomagnetite = proxy["target_species"]["titanomagnetite"]
    _add_table_before(
        document, anchor,
        ["设置", "共同子集 Macro F1", "目标召回", "钛铁矿召回（n=50）", "钛磁铁矿召回（n=4）"],
        [
            ["完整数据", f"{100*proxy_full['macro_f1']['mean']:.2f}%±{100*proxy_full['macro_f1']['sample_std']:.2f}%", f"{100*proxy_full['target_recall']['mean']:.2f}%±{100*proxy_full['target_recall']['sample_std']:.2f}%", f"{100*ilmenite['full_common_subset']['mean']:.2f}%±{100*ilmenite['full_common_subset']['sample_std']:.2f}%", f"{100*titanomagnetite['full_common_subset']['mean']:.2f}%±{100*titanomagnetite['full_common_subset']['sample_std']:.2f}%"],
            ["移除普通磁铁矿代理", f"{100*proxy_none['macro_f1']['mean']:.2f}%±{100*proxy_none['macro_f1']['sample_std']:.2f}%", f"{100*proxy_none['target_recall']['mean']:.2f}%±{100*proxy_none['target_recall']['sample_std']:.2f}%", f"{100*ilmenite['no_magnetite_proxy']['mean']:.2f}%±{100*ilmenite['no_magnetite_proxy']['sample_std']:.2f}%", f"{100*titanomagnetite['no_magnetite_proxy']['mean']:.2f}%±{100*titanomagnetite['no_magnetite_proxy']['sample_std']:.2f}%"],
        ], [3.5, 3.2, 2.8, 3.5, 3.5], "表 23 普通磁铁矿代理样本消融"
    )
    _add_paragraph_before(
        document, anchor,
        "移除代理样本后，共同子集总体准确率上升，但目标召回由72.22%降至45.06%，钛铁矿召回由70.67%降至44.67%。这说明类别组成变化会使总体准确率产生误导，同时普通磁铁矿可增强目标类视觉表征；然而其结果不能替代真实钛磁铁矿验证。钛磁铁矿测试仅4张，其召回率只能作描述性展示。"
    )

    # Add the complete model-and-risk pipeline after the existing baseline structure discussion.
    architecture_anchor = headings["4.6 固定映射下的角色可识别性命题"]
    _add_picture_before(document, architecture_anchor, architecture, 16.0, "图 5 理论感知的分层网络、校准与风险认证结构（修订）")
    _add_paragraph_before(
        document, architecture_anchor,
        "图5在 EfficientNet-B0 基线主干上明确给出四个任务头、固定种类—角色映射、DKL(p̃r||pr) 一致性约束以及推理期温度校准和风险认证。网络训练与选择性决策分开：训练阶段学习视觉表征和层级关系，推理阶段使用独立验证数据确定温度与风险阈值。"
    )

    # Update the abstract, stale discussion, near-term plan, and conclusions.
    replace_paragraph_text(original_paragraphs[15],
        "在固定划分上，采用 ImageNet 预训练的 ResNet50 和 EfficientNet-B0 开展三随机种子重复训练，EfficientNet-B0 的平均 Macro F1 为72.87%±1.21%。在此基础上构建共享 EfficientNet-B0 主干的分层一致性网络，同时预测4类选矿角色、17类矿物种类及目标/非目标标签，并以固定映射和困难负样本约束建立可解释的风险结构。种类预测映射到角色后产生18.74±1.29个百分点的经验风险收缩；一致性项使双头 KL 从0.0565降至0.0426。温度缩放将测试 ECE 从0.1131降至0.0342，并在独立认证子集上实现预设选择性风险控制。")
    replace_paragraph_text(original_paragraphs[16],
        "摄影者整体留出的严格实验包含5,639张具有来源字段的图像。分层模型在该测试上的 Macro F1 为69.70%±0.90%，较 EfficientNet-B0 提高2.70个百分点，并改善目标类和金属光泽干扰召回。普通磁铁矿代理消融表明，移除1,316张代理样本后目标召回由72.22%降至45.06%；钛磁铁矿最终保留23张、固定测试仅4张，因此其结果仅作描述性证据。当前工作定位为公开标本图像条件下的钒钛矿相关组分矿物视觉代理研究，不构成工业品位、回收率或现场分选验证。")
    replace_paragraph_text(original_paragraphs[17],
        "关键词：钒钛矿；矿物图像识别；分层一致性；角色风险收缩；概率校准；来源外泛化；选择性风险认证")
    replace_paragraph_text(original_paragraphs[168],
        "未校准的固定测试阈值曲线仅作为探索性结果保留。正式风险阈值已在5.13节中改为：验证集温度拟合、独立认证子集计算同时置信上界、测试集一次性评估。因而0.90不再被解释为推荐阈值，实际阈值随预设风险δ和随机种子而变化。")
    replace_paragraph_text(original_paragraphs[170],
        "当前证据链由五部分组成：多种子基线确认基本性能；固定映射证明并实证角色风险收缩；一致性消融显示双头 KL、TV 和预测分歧下降；温度校准与独立认证把描述性置信度曲线升级为预设风险约束；摄影者留出和代理消融检验来源偏差与数据构成。结果支持分层模型在严格来源留出中的改进趋势，但尚不能声称跨网站、真实矿石或工业现场有效。")
    replace_paragraph_text(original_paragraphs[174],
        "本次修订把理论贡献限定为可验证的三层结构：其一，固定种类—角色映射诱导逐样本风险支配和经验风险收缩；其二，KL 一致性通过 Pinsker 不等式控制双头分布差异，并由消融验证其分布效果；其三，温度校准结合多阈值修正的 Clopper–Pearson 上界，实现验证集上的选择性风险认证。来源留出和代理消融进一步限定这些结论的适用边界。")
    replace_paragraph_text(original_paragraphs[186],
        "第四，已利用摄影者字段建立严格来源留出测试：5,639张可用样本按摄影者与重复图像组整体划分，训练、验证和测试的摄影者及 split_group_id 均无重叠。后续仍需扩展到矿区、收藏机构和跨网站整体留出。")
    replace_paragraph_text(original_paragraphs[187],
        "第五，已完成验证集温度缩放、NLL/ECE/Brier/AURC评估和独立认证子集上的风险阈值选择。后续需要扩大校准样本规模，比较向量缩放、Dirichlet 校准和共形风险控制，并预注册更细的风险约束。")
    replace_paragraph_text(original_paragraphs[197],
        "6. 分层一致性模型不仅输出四类角色，还同时学习17类种类、目标/非目标和困难负样本关系。固定映射使种类错误中的43.41%仍保持角色正确，对应18.74±1.29个百分点的经验角色风险收缩。")
    replace_paragraph_text(original_paragraphs[198],
        "7. DKL(p̃r||pr) 一致性项将平均 KL 从0.0565降至0.0426、双头分歧率从7.42%降至6.36%；该贡献体现为层级分布一致性增强，不保证总体分类精度必然提高。")
    replace_paragraph_text(original_paragraphs[199],
        "8. 验证集温度缩放将测试 ECE 从0.1131降至0.0342、NLL 从0.8399降至0.6633；在δ=0.15和δ=0.20时三个种子均获得独立认证子集上的风险证书，实现了比测试集后验选阈更严格的选择性决策验证。")
    replace_paragraph_text(original_paragraphs[200],
        "9. 摄影者来源留出实验中，分层模型 Macro F1 为69.70%±0.90%，较基线提高2.70个百分点；代理消融表明普通磁铁矿有助于目标类表征，但最终23张钛磁铁矿及测试中的4张样本不足以支撑真实性或工业能力结论。")
    replace_paragraph_text(original_paragraphs[224],
        "项目保留数据清单、采集与筛选脚本、数据质量审计记录、训练脚本、实验指标、混淆矩阵、逐图预测和图表源数据。层级风险汇总位于 outputs/theory_validation/hierarchy_consistency/；校准风险认证汇总位于 outputs/theory_validation/calibrated_selective_recognition/；来源留出与代理消融汇总位于 outputs/paper_experiments_v2/training_summaries/；新增论文图位于 outputs/paper_figures_v2/。")

    # Renumber pre-existing captions shifted by the newly inserted architecture and experiment figures/tables.
    caption_replacements = {
        "图 5 四种模型的总体性能对比（均值±样本标准差，n=3）": "图 6 四种模型的总体性能对比（均值±样本标准差，n=3）",
        "图 6 四类矿物在四种模型上的召回率": "图 7 四类矿物在四种模型上的召回率",
        "图 7 四种训练策略的目标代理风险指标": "图 8 四种训练策略的目标代理风险指标",
        "图 8 EfficientNet-B0（随机种子 20260728）的混淆矩阵": "图 9 EfficientNet-B0（随机种子 20260728）的混淆矩阵",
        "图 9 EfficientNet-B0（随机种子 20260728）的主要错分方向": "图 10 EfficientNet-B0（随机种子 20260728）的主要错分方向",
        "图 10 加权交叉熵与加权 Focal Loss 的三随机种子消融": "图 11 加权交叉熵与加权 Focal Loss 的三随机种子消融",
        "图 11 三随机种子在固定测试划分上的选择性识别风险—覆盖率关系": "图 12 三随机种子在固定测试划分上的选择性识别风险—覆盖率关系",
        "表 18 本阶段方法贡献及其解释边界": "表 24 本阶段方法贡献及其解释边界",
        "表 19 当前局限性及处理原则": "表 25 当前局限性及处理原则",
    }
    for paragraph in original_paragraphs:
        text = paragraph.text.strip()
        if text in caption_replacements:
            replace_paragraph_text(paragraph, caption_replacements[text])

    # Update the original contribution and limitation tables using their semantic row labels.
    for table in original_tables:
        changed = False
        for row in table.rows:
            label = row.cells[0].text.strip()
            if label == "领域问题形式化":
                row.cells[2].text = "给出固定种类—角色映射下的逐样本风险支配和经验风险分解，并以三随机种子逐图预测验证18.74±1.29个百分点的角色风险收缩；结论限于当前固定映射。"
                changed = True
            elif label == "主要方法贡献":
                row.cells[2].text = "共享 EfficientNet-B0 主干联合角色、种类、目标代理和困难负样本学习，以 D_KL(Ap_s||p_r) 约束层级一致性；消融显示 KL、TV 和双头分歧下降。"
                changed = True
            elif label == "风险导向评价贡献":
                row.cells[2].text = "采用独立验证子集完成温度校准和多阈值风险认证，并以摄影者留出测试和代理样本消融检验来源偏差与数据构成。"
                changed = True
            elif label == "来源外泛化尚未验证":
                row.cells[0].text = "来源外泛化仅完成摄影者留出"
                row.cells[1].text = "尚未覆盖矿区、收藏机构或跨网站整体留出"
                row.cells[2].text = "当前报告摄影者无交叉结果；继续构建更严格外部域测试"
                changed = True
            elif label == "置信度尚未校准":
                row.cells[0].text = "校准与风险认证仍受样本规模限制"
                row.cells[1].text = "δ=0.10仅一个种子取得证书，严格风险控制不稳定"
                row.cells[2].text = "扩展认证样本并比较向量缩放、Dirichlet 校准和共形风险控制"
                changed = True
            elif label == "普通 magnetite 作为代理":
                row.cells[1].text = "代理样本提高目标召回，但不能代表真实钛磁铁矿能力"
                row.cells[2].text = "已完成使用/不使用代理消融，并单列钛磁铁矿小样本限制"
                changed = True
        if changed:
            _format_table(table, [4.0, 6.0, 6.0])

    # Insert a final boundary statement before references.
    reference_anchor = headings["参考文献"]
    _add_paragraph_before(document, reference_anchor,
        "10. 本研究的理论结果建立在固定种类—角色映射、同分布校准样本和公开矿物标本图像上；当前结论不涉及钒钛品位预测、工业回收率、真实XRF送检成本或现场生产线性能。阶段条件化选矿决策图保留为获得流程阶段与成本数据后的后续研究方向。")

    # Add calibration reference and ask Word/LibreOffice to refresh fields on open.
    appendix_anchor = next(p for p in original_paragraphs if p.text.strip().startswith("附录 A"))
    _add_paragraph_before(document, appendix_anchor,
        "[19] Guo C, Pleiss G, Sun Y, Weinberger K Q. On Calibration of Modern Neural Networks[C]//Proceedings of ICML. 2017: 1321-1330.")
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def parse_args():
    parser = argparse.ArgumentParser(description="Update the formal technical report with v2 theory experiments.")
    parser.add_argument("--input", type=Path, default=FORMAL_REPORT)
    parser.add_argument("--output", type=Path, default=FORMAL_REPORT)
    return parser.parse_args()


def main():
    args = parse_args()
    print(update_report(args.input.resolve(), args.output.resolve()))


if __name__ == "__main__":
    main()
