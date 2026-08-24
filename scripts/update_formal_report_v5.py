from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from update_formal_report_v2 import (
    FORMAL_REPORT,
    _add_heading_before,
    _add_paragraph_before,
    _add_picture_before,
    _add_table_before,
    _set_run_font,
    replace_paragraph_text,
)


ROOT = Path(__file__).resolve().parents[1]
RSG_SUMMARY_JSON = ROOT / "outputs" / "business_metrics" / "rsg_hrgv" / "formal" / "rsg_three_seed_summary.json"
RSG_PAIRED_JSON = (
    ROOT / "outputs" / "business_metrics" / "rsg_hrgv" / "formal" / "paired_rsg_complete_vs_hrgv_reference.json"
)
ARCHITECTURE = ROOT / "outputs" / "paper_figures_v2" / "fig_hrgv_architecture.png"
ASSET_DIR = ROOT / "outputs" / "report_assets_v5"


def _pct(mean: float, std: float | None = None) -> str:
    if std is None:
        return f"{100.0 * mean:.2f}%"
    return f"{100.0 * mean:.2f}% ± {100.0 * std:.2f}%"


def _metric(summary: dict, setting: str, key: str) -> str:
    value = summary[setting][key]
    return _pct(float(value["mean"]), float(value["sample_std"]))


def _remove_figure_and_caption(caption) -> None:
    caption_element = caption._p
    previous = caption_element.getprevious()
    parent = caption_element.getparent()
    parent.remove(caption_element)
    if previous is not None and previous.xpath(".//w:drawing"):
        previous.getparent().remove(previous)


def _insert_numbered_conclusion(document, anchor, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(run)
    paragraph.paragraph_format.line_spacing = 1.5
    anchor._p.addprevious(paragraph._element)


def render_rsg_equations(output_dir: Path) -> dict[str, Path]:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    equations = {
        "regret": (
            r"$\ell_d=-\log p_d(y\mid x),\quad \ell_m=-\log p_m(y\mid x),\quad "
            r"\Delta=\ell_m-\ell_d,\quad g^*=\sigma(\Delta/T_r),\quad "
            r"w=\tanh(|\Delta|/T_w),\quad \mathcal{L}_{reg}=w\,\mathrm{BCE}(g,g^*)$"
        ),
        "bounds": (
            r"$0\leq-\log p_g+\log\max(a,b)\leq "
            r"\frac{|g-g_o|\,|a-b|}{\varepsilon},\qquad "
            r"|g^*-g_o|\leq\exp(-|\Delta|/T_r),\qquad "
            r"\nabla_{\theta_{expert}}\mathcal{L}_{reg}=0$"
        ),
    }
    output_dir.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    for name, equation in equations.items():
        figure = plt.figure(figsize=(15.4, 1.15), dpi=220, facecolor="white")
        figure.text(0.5, 0.5, equation, ha="center", va="center", fontsize=16.0, color="#1F2937")
        path = output_dir / f"rsg_{name}.png"
        figure.savefig(path, bbox_inches="tight", pad_inches=0.12, facecolor="white")
        plt.close(figure)
        paths[name] = path
    return paths


def _rsg_rows(summary: dict) -> list[list[str]]:
    rows = []
    for key, label in (
        ("hrgv_reference", "耦合残差 HRGV（参考）"),
        ("rsg_complete", "RSG-HRGV（完整）"),
    ):
        rows.append(
            [
                label,
                _metric(summary, key, "accuracy"),
                _metric(summary, key, "macro_f1"),
                _metric(summary, key, "target_recall"),
                _metric(summary, key, "ti_to_target_intrusion_rate"),
                _metric(summary, key, "metallic_to_target_intrusion_rate"),
                _metric(summary, key, "one_right_gate_selection_accuracy"),
                _metric(summary, key, "mean_routing_regret_nll"),
            ]
        )
    return rows


def update_report(input_path: Path, output_path: Path) -> Path:
    summary = json.loads(RSG_SUMMARY_JSON.read_text(encoding="utf-8"))
    paired = json.loads(RSG_PAIRED_JSON.read_text(encoding="utf-8"))
    equations = render_rsg_equations(ASSET_DIR)
    document = Document(input_path)

    if any(p.text.strip() == "4.6.1 后悔监督门控（RSG）" for p in document.paragraphs):
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path

    headings = {p.text.strip(): p for p in document.paragraphs if p.text.strip()}
    method_anchor = headings["4.7 固定映射下的角色可识别性命题"]

    old_caption = next(
        p for p in document.paragraphs if p.text.strip() == "图 5 HRGV-Net 的层级证据融合与困难负样本风险验证结构"
    )
    _remove_figure_and_caption(old_caption)
    _add_picture_before(
        document,
        method_anchor,
        ARCHITECTURE,
        16.1,
        "图 5 RSG-HRGV-Net 的分层证据、后悔监督门控与困难负样本风险验证结构",
    )

    model_intro = next(p for p in document.paragraphs if p.text.startswith("既有分层模型能够利用17类矿物种类监督"))
    replace_paragraph_text(
        model_intro,
        "既有分层模型能够利用17类矿物种类监督提升目标召回，但固定测试集的成对分析表明，两类困难负样本误入目标的比例也随之上升。为直接处理这一失效模式，本研究在 HRGV-Net 的层级双专家、双验证器与中性区残差校正基础上，引入后悔监督门控，形成 RSG-HRGV-Net。网络以 EfficientNet-B0 为共享主干，设置直接角色专家和矿物种类专家；后者经固定映射矩阵 A 聚合为角色后验。可靠性门控根据共享特征、两个专家的熵及 Jensen-Shannon 分歧进行逐图融合，训练期再使用真实类别下的专家相对损失给予门控方向性监督。随后由“目标—含钛干扰”和“目标—金属光泽干扰”两个验证器对目标后验进行中性区残差校正。本方法仍是公开矿物标本图像上的视觉代理任务，不等同于工业分选回收率验证。",
    )

    _add_heading_before(document, method_anchor, "4.6.1 后悔监督门控（RSG）", 3)
    _add_paragraph_before(
        document,
        method_anchor,
        "常规融合损失只要求融合后验正确，却不会直接告诉门控“本图像更应信任哪一位专家”。RSG 以真实角色 y 下的直接角色专家与种类映射专家的负对数损失差为可观测的训练期证据，构造连续门控目标。令 g 表示直接角色专家的融合权重，则：",
    )
    _add_picture_before(document, method_anchor, equations["regret"], 16.0, "式（17） RSG 的样本级后悔差、软门控目标与加权监督")
    _add_paragraph_before(
        document,
        method_anchor,
        "当直接角色专家更可靠时，Δ 为正且 g* 趋近 1；当种类映射专家更可靠时，g* 趋近 0。差距权重 w 使两位专家几乎等价的样本不会主导门控学习。正式配置取 λg=0.1、Tr=0.2、Tw=0.5；为防止专家通过人为改变后验迎合门控标签，g* 与 w 由停止梯度的专家后验构造。该隔离只作用于后悔监督分支，双验证器仍通过原有耦合路径参与共享表征训练。",
    )
    _add_paragraph_before(
        document,
        method_anchor,
        "命题6（门控误差的路由后悔上界）",
        bold_lead="命题6（门控误差的路由后悔上界）",
    )
    _add_paragraph_before(
        document,
        method_anchor,
        "令 a=pd(y|x)、b=pm(y|x)，且 a,b≥ε>0；令 go=I(a≥b) 为硬最优门控，pg=ga+(1−g)b。则融合相对选择真实类别概率更大专家的对数后悔由式（18）上界。证明使用 −log 在 [ε,1] 上的 Lipschitz 性质。它说明，门控误差仅在专家对真实类别的证据差距较大时造成较高后悔，因此差距加权具有明确的优化含义。",
    )
    _add_paragraph_before(
        document,
        method_anchor,
        "命题7（软目标对硬门控的指数逼近）",
        bold_lead="命题7（软目标对硬门控的指数逼近）",
    )
    _add_paragraph_before(
        document,
        method_anchor,
        "对 Δ≠0，软目标 g* 与硬最优门控 go 的距离不超过 exp(−|Δ|/Tr)。因此，当两专家明显分出优劣时，连续目标逼近硬选择；当它们接近时，目标保持不确定性而非制造过强监督。该命题对应软目标、硬目标和无差距加权的受控预实验设计。",
    )
    _add_paragraph_before(
        document,
        method_anchor,
        "命题8（后悔监督的局部梯度隔离）",
        bold_lead="命题8（后悔监督的局部梯度隔离）",
    )
    _add_paragraph_before(
        document,
        method_anchor,
        "若构造 g* 与 w 的专家后验使用停止梯度，则后悔监督关于专家参数的梯度严格为0，而关于门控参数的梯度可非零。该结论仅限定 Lreg 分支，不切断角色、种类、验证器和对比约束对共享主干的训练。它解释了为何“局部梯度隔离”不同于将整个验证器与主干隔离。",
    )
    _add_picture_before(document, method_anchor, equations["bounds"], 16.0, "式（18） RSG 的路由后悔上界、软目标逼近与局部梯度隔离")
    _add_paragraph_before(
        document,
        method_anchor,
        "命题6至8分别给出门控应优化何种风险、连续目标为何合理以及梯度如何受控；它们不等价于“任何门控网络均会提高总体准确率”。因此，后续实验同时报告常规四分类指标和以真实类别专家后验定义的路由指标，并用成对簇 Bootstrap 区间区分机制证据与总体任务分数。",
    )

    headings = {p.text.strip(): p for p in document.paragraphs if p.text.strip()}
    result_anchor = headings["5.18 实验结果的综合判断"]
    replace_paragraph_text(result_anchor, "5.19 实验结果的综合判断")
    _add_heading_before(document, result_anchor, "5.18 RSG-HRGV-Net 后悔监督门控正式验证", 2)
    _add_paragraph_before(
        document,
        result_anchor,
        "RSG-HRGV 与参考 HRGV 使用完全相同的数据清单、固定划分、耦合残差验证器推理图和三个随机种子，仅在门控训练中加入冻结的软后悔目标与差距加权。软目标、硬目标、无差距加权和允许后悔监督反传至专家的单种子受控预实验用于冻结配置；下表和统计推断只使用参考 HRGV 与完整 RSG-HRGV 的三随机种子正式对照。",
    )
    _add_table_before(
        document,
        result_anchor,
        ["设置", "Accuracy", "Macro F1", "目标召回", "含钛误入", "金属误入", "一对一错路由", "平均路由后悔"],
        _rsg_rows(summary),
        [2.45, 1.50, 1.50, 1.55, 1.55, 1.55, 2.05, 1.65],
        "表 26 RSG-HRGV-Net 与 HRGV 的三随机种子正式比较",
    )

    classification = paired["classification"]
    routing = paired["routing_regret"]
    rsg = summary["rsg_complete"]
    ref = summary["hrgv_reference"]
    gate_delta = 100.0 * (
        float(rsg["one_right_gate_selection_accuracy"]["mean"])
        - float(ref["one_right_gate_selection_accuracy"]["mean"])
    )
    regret_delta = 100.0 * float(routing["difference"])
    regret_ci_low = 100.0 * float(routing["ci_low"])
    regret_ci_high = 100.0 * float(routing["ci_high"])
    _add_paragraph_before(
        document,
        result_anchor,
        f"完整 RSG-HRGV 的平均 Accuracy 为{_metric(summary, 'rsg_complete', 'accuracy')}，与参考 HRGV 的{_metric(summary, 'hrgv_reference', 'accuracy')}基本相当；Macro F1 由{_metric(summary, 'hrgv_reference', 'macro_f1')}变为{_metric(summary, 'rsg_complete', 'macro_f1')}。目标召回和两类误入目标比例在均值上均呈有利变化，但 Accuracy、Macro F1、目标召回以及两类误入指标的 2,000 次成对簇 Bootstrap 区间均跨0。因此，当前证据不能写成 RSG 对总体四分类性能具有显著优势。",
    )
    _add_paragraph_before(
        document,
        result_anchor,
        f"RSG 的机制结果更明确：在“恰有一位专家正确”的可路由样本上，门控选择正确专家的比例从{_pct(float(ref['one_right_gate_selection_accuracy']['mean']), float(ref['one_right_gate_selection_accuracy']['sample_std']))}升至{_pct(float(rsg['one_right_gate_selection_accuracy']['mean']), float(rsg['one_right_gate_selection_accuracy']['sample_std']))}，提高{gate_delta:.2f}个百分点；平均路由后悔从{_pct(float(ref['mean_routing_regret_nll']['mean']), float(ref['mean_routing_regret_nll']['sample_std']))}降至{_pct(float(rsg['mean_routing_regret_nll']['mean']), float(rsg['mean_routing_regret_nll']['sample_std']))}。后者的成对差异（RSG−HRGV）为{regret_delta:.2f}个百分点，95%簇 Bootstrap 区间为[{regret_ci_low:.2f}, {regret_ci_high:.2f}]个百分点，未跨0，方向性重采样概率为{float(routing['probability_favorable']):.4f}。这支持 RSG 在本固定测试协议下确实降低了式（18）所定义的路由后悔。",
    )
    _add_paragraph_before(
        document,
        result_anchor,
        "该结果形成了可审计的理论—实验闭环：命题6给出门控误差与路由后悔的上界，命题7给出软目标在专家差异明确时逼近硬选择的理由，命题8限定后悔监督不改写专家证据；正式实验则表明，所设计的局部监督稳定改善了该路由目标。由于分类终端指标的区间仍跨0，论文应将创新表述为“针对跨粒度专家证据选择的后悔监督机制及其可检验路由风险改善”，而不是宣称普适的分类精度突破。",
    )

    replacements = {
        "表 26 本阶段方法贡献及其解释边界": "表 27 本阶段方法贡献及其解释边界",
        "表 27 当前局限性及处理原则": "表 28 当前局限性及处理原则",
    }
    for paragraph in document.paragraphs:
        if paragraph.text.strip() in replacements:
            replace_paragraph_text(paragraph, replacements[paragraph.text.strip()])

    keyword_anchor = next(p for p in document.paragraphs if p.text.strip().startswith("关键词："))
    _add_paragraph_before(
        document,
        keyword_anchor,
        "进一步的 RSG-HRGV-Net 将“应融合哪些证据”从间接终端监督转为基于真实类别专家损失差的后悔监督。三随机种子正式比较中，RSG 的总体 Accuracy 为76.45%±0.98%、Macro F1 为74.76%±0.90%，与参考 HRGV 的差异区间均跨0；但平均路由后悔下降1.77个百分点，95%簇 Bootstrap 区间为[−2.86, −0.69]个百分点。该网络创新的实证支撑来自路由机制，而非对总体分类精度的过度宣传。",
    )

    old_last_conclusion = next(
        p for p in document.paragraphs if p.text.strip().startswith("12. 本研究的理论结果建立在固定种类—角色映射")
    )
    replace_paragraph_text(old_last_conclusion, old_last_conclusion.text.replace("12. 本研究", "13. 本研究", 1))
    _insert_numbered_conclusion(
        document,
        old_last_conclusion,
        "12. RSG-HRGV-Net 在 HRGV 的双专家、验证器与残差校正基础上，引入由样本级专家相对损失构造的软后悔门控目标，并以差距权重和局部停止梯度控制学习范围。三随机种子正式对照表明，该设计将平均路由后悔降低1.77个百分点，95%簇 Bootstrap 区间为[−2.86, −0.69]个百分点；总体 Accuracy 和 Macro F1 的差异区间均跨0。因此，本研究将其定位为可证明、可测量的证据路由创新，不将其解释为分类精度的普适显著提升。",
    )

    for paragraph in document.paragraphs:
        if "HRGV 三随机种子汇总、消融差值、验证器 AUC" in paragraph.text:
            replace_paragraph_text(
                paragraph,
                paragraph.text
                + " RSG-HRGV 的冻结配置、受控预实验、三随机种子逐图预测与成对统计位于 outputs/training/rsg_controlled/、outputs/business_metrics/rsg_hrgv/formal/ 和 docs/experiment_records/2026-08-22_rsg_hrgv.md。",
            )
            break

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Integrate RSG-HRGV theory and formal evidence into the report.")
    parser.add_argument("--input", type=Path, default=FORMAL_REPORT)
    parser.add_argument("--output", type=Path, default=FORMAL_REPORT)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    print(update_report(args.input.resolve(), args.output.resolve()))


if __name__ == "__main__":
    main()
