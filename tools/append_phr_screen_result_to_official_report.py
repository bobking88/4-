"""Append the audited PHR validation-screen outcome to the official report."""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "结题" / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（正式版）.docx"
DECISION = ROOT / "outputs" / "training" / "phr_hrgv_screen_v2" / "screen_decision.json"
TITLE = "I.6 PHR 验证集筛选结果与边界结论"


def _set_font(run: Any, name: str, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for field in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{field}"), name)


def _add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.4
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    _set_font(paragraph.add_run(text), "宋体", 11)


def _metric(decision: dict[str, Any], configuration: str) -> dict[str, float]:
    key = f"{configuration}/seed{decision['screen_seed']}"
    raw = decision["evidence"]["runs"][key]["metrics"]
    return {name: float(value) for name, value in raw.items()}


def _add_result_table(document: Document, reference: dict[str, float], candidate: dict[str, float]) -> None:
    headers = ("配置", "Macro F1", "目标类召回", "Ti 误入目标", "金属误入目标", "Ti/金属边遗憾")
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            _set_font(run, "宋体", 8.5, True)
    rows = (("RSG 参考", reference), ("完整 PHR", candidate))
    for label, metrics in rows:
        values = (
            label,
            f"{metrics['macro_f1']:.4f}",
            f"{metrics['target_recall']:.4f}",
            f"{metrics['ti_to_target_intrusion_rate']:.4f}",
            f"{metrics['metallic_to_target_intrusion_rate']:.4f}",
            f"{metrics['phr_ti_mean_margin_regret']:.4f}/{metrics['phr_metallic_mean_margin_regret']:.4f}",
        )
        for cell, text in zip(table.add_row().cells, values):
            cell.text = text
            for run in cell.paragraphs[0].runs:
                _set_font(run, "宋体", 8.5)


def append_screen_result(report: Path, decision: dict[str, Any]) -> bool:
    document = Document(report)
    if any(paragraph.text.strip() == TITLE for paragraph in document.paragraphs):
        return False
    if decision.get("promote_to_formal") is not False or decision.get("criterion_ids") != []:
        raise ValueError("This appendix writer is restricted to the audited PHR non-promotion decision.")

    reference = _metric(decision, "rsg_reference")
    candidate = _metric(decision, "phr_complete")
    delta_f1 = candidate["macro_f1"] - reference["macro_f1"]
    delta_recall = candidate["target_recall"] - reference["target_recall"]
    delta_accuracy = candidate["accuracy"] - reference["accuracy"]

    heading = document.add_paragraph(style="Heading 2")
    heading.paragraph_format.keep_with_next = True
    _set_font(heading.add_run(TITLE), "黑体", 12, True)
    _add_body(document, "PHR 的八项预注册验证集筛选已完成。筛选只使用冻结清单中的验证集，自动决定文件对运行状态、环境、指标、逐图预测、登记配置和源文件哈希进行了绑定。其 criterion_ids 为空且 promote_to_formal=false，因此完整 PHR 未晋级到锁定测试集上的三随机种子正式比较。")
    _add_result_table(document, reference, candidate)
    _add_body(document, f"完整 PHR 的 Macro F1 相对 RSG 参考变化为 {delta_f1:+.4f}，Accuracy 变化为 {delta_accuracy:+.4f}，目标类召回变化为 {delta_recall:+.4f}。虽然两条成对边遗憾从 {reference['phr_ti_mean_margin_regret']:.4f}/{reference['phr_metallic_mean_margin_regret']:.4f} 降至 {candidate['phr_ti_mean_margin_regret']:.4f}/{candidate['phr_metallic_mean_margin_regret']:.4f}，但目标类召回下降，且 Ti 与金属光泽干扰误入目标的比例均上升。因此，它未满足既定的总体性能、局部风险或风险取舍三条任一晋级规则。")
    _add_body(document, "该结果支持本附录已明确的理论边界：两条成对边界上的局部遗憾降低，并不推出 Softmax 耦合后的全局四分类风险、Macro F1 或目标类代理指标改善。PHR 的解析遗憾恒等式、软目标损失界和最小范数约束修正仍作为可复核的候选机制保留；但在当前数据与协议下，PHR 只能作为机制性负结果和后续研究线索，不能写作本报告的性能创新成果。")

    temporary = report.with_suffix(".phr_screen_result.tmp.docx")
    document.save(temporary)
    temporary.replace(report)
    return True


def main(argv: Sequence[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--report", type=Path, default=REPORT)
    parser.add_argument("--decision", type=Path, default=DECISION)
    args = parser.parse_args(argv)
    decision = json.loads(args.decision.read_text(encoding="utf-8"))
    changed = append_screen_result(args.report, decision)
    print(f"updated={changed} report={args.report}")


if __name__ == "__main__":
    main()
