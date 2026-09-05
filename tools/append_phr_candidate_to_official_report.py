"""Append the pre-registered PHR candidate-network appendix to the official report.

The appendix is deliberately framed as a candidate extension.  It carries the
formal definitions and implementation checks needed for future paper work but
does not claim a classification improvement before the registered experiment
has completed.
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "结题" / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（正式版）.docx"
BACKUP = ROOT / "结题" / "历史版本" / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（正式版_20260902_追加PHR前）.docx"
FIGURE = ROOT / "outputs" / "paper_figures_v3" / "fig_phr_routing_architecture.png"


def set_run_font(run, name: str, size: float = 11.0, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for field in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{field}"), name)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.4
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    set_run_font(paragraph.add_run(text), "宋体", 11)


def add_equation(document: Document, text: str, number: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    equation = paragraph.add_run(text)
    set_run_font(equation, "Cambria Math", 11)
    paragraph.add_run("    ")
    label = paragraph.add_run(f"({number})")
    set_run_font(label, "Times New Roman", 10.5)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(text), "黑体", 14 if level == 1 else 12, True)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(5)
    set_run_font(paragraph.add_run(text), "宋体", 10.5)


def build() -> None:
    if not REPORT.exists():
        raise FileNotFoundError(REPORT)
    if not FIGURE.exists():
        raise FileNotFoundError(FIGURE)
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    document = Document(REPORT)
    document.add_page_break()
    add_heading(document, "附录 I 候选网络扩展：成对困难负样本遗憾路由（PHR）", 1)
    add_body(document, "本附录记录后续论文补实验所采用的候选网络扩展。PHR 的目标不是以未经证实的方式替代当前 RSG-HRGV 主模型，而是针对已观察到的两条高风险混淆边界，即目标矿物与含钛干扰矿物、目标矿物与金属光泽干扰矿物，分别学习路由门控并施加最小范数的成对边界修正。以下公式、结构图与数值性质核验为预注册实验提供可复现定义；在验证集筛选和正式三随机种子测试完成前，本附录不作任何性能提升结论。")

    add_heading(document, "I.1 成对边界的概率、对数几率与融合门控", 2)
    add_body(document, "设 T、I、G、M 分别表示目标矿物、含钛干扰、脉石和金属光泽干扰；p_d 与 p_m 分别为直接角色头和种类概率映射后的角色后验。对 q∈{I,M}，PHR 不直接重写四分类标签，而是在两条预定义边界上计算目标相对负类的对数几率。")
    add_equation(document, "m_(d,q)=log(p_(d,T)/p_(d,q)),   m_(m,q)=log(p_(m,T)/p_(m,q))", 1)
    add_equation(document, "m_(f,q)=g_q m_(d,q)+(1-g_q)m_(m,q),   g_q=sigma(phi_q([f,m_(d,q),m_(m,q)]))", 2)
    add_body(document, "其中 f 为共享骨干的视觉特征，g_q∈[0,1] 是第 q 条边界的路由门控，sigma 为 Sigmoid 函数。两条门控互不强制相等，因此可分别处理‘目标—含钛干扰’与‘目标—金属光泽干扰’的视觉证据差异。")

    add_heading(document, "I.2 软专家与遗憾目标", 2)
    add_body(document, "对训练样本，只在真实角色属于当前成对边界的情形启用该边界。设 Delta_q 为直接角色证据与种类映射证据在第 q 条边界上的有向差异，使用温度 tau 构造软专家门控目标；其对应的精确边际遗憾定义如下。")
    add_equation(document, "g_(o,q)=sigma(Delta_q/tau),   r_q=|g_q-g_(o,q)| |Delta_q|", 3)
    add_equation(document, "L_PHR=(sum_(q in {I,M}) w_q r_q)/(sum_(q in {I,M}) w_q 1[eligible_q])", 4)
    add_body(document, "在二元对数损失下，门控偏离软专家所造成的边际损失超额可由 r_q 上界；这一结论只针对预定义的成对对数几率边界，不推出四分类 Macro F1、外部泛化或工业分选指标必然提升。数值实现以 10,000 组 float64 合成样本核验遗憾恒等式、损失上界和符号保持条件，结果仅用于公式实现自检。")

    add_heading(document, "I.3 最小范数耦合修正与性质边界", 2)
    add_body(document, "PHR 将两条成对边界的修正量同时投影到四类 logit，以最小二范数改变满足约束。令 delta_I=m_(f,I)-m_(0,I)、delta_M=m_(f,M)-m_(0,M)，其中 m_0 为 RSG-HRGV 融合后的基线边际，则一组闭式修正为：")
    add_equation(document, "a_T=(delta_I+delta_M)/3,   a_I=(delta_M-2delta_I)/3,   a_M=(delta_I-2delta_M)/3,   a_G=0", 5)
    add_equation(document, "ell_PHR=ell_0+(a_T,a_I,a_G,a_M)", 6)
    add_body(document, "该投影满足目标—含钛干扰和目标—金属光泽干扰的指定边际变化，并在所有满足约束的 logit 修正中具有最小二范数。需特别说明：a_G=0 只表示脉石 logit 不变；Softmax 归一化后，脉石概率仍可能变化。因此，局部成对边际改善不等价于全局四分类风险、校准误差或 Macro F1 的必然改善。")

    add_heading(document, "I.4 网络结构与可审计验证方案", 2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(FIGURE), width=Cm(15.2))
    add_caption(document, "图 27 PHR 候选网络扩展：两条成对边界的门控、遗憾监督与最小范数修正。虚线部分仅在训练期使用。")
    add_body(document, "筛选阶段固定数据清单，仅使用验证集，并预注册八组设置：RSG 参考、完整 PHR、固定 0.5 门控、硬门控目标、非加权损失、耦合特征、仅含钛边界和仅金属光泽边界。完整 PHR 只有满足预先定义的 Macro F1、成对遗憾或风险取舍准则之一时，才会晋级至三随机种子的锁定测试集比较。正式分析同时核验清单哈希、脚本哈希、运行状态、验证集晋级决定和逐图预测，避免把旧试跑、测试集调参或不完整运行写入结论。")

    add_heading(document, "I.5 与本报告主线的关系", 2)
    add_body(document, "当前技术报告的已完成主线仍为角色感知分层一致性与困难负样本学习。PHR 作为可检验的网络创新候选，补足了‘两类困难干扰具有不同边界、应分别路由’的形式化定义与可复现实验协议。若后续正式实验未通过验证集晋级或未在锁定测试上形成可信证据，应将其保留为负结果或后续工作，而不能写入本报告的已取得性能成果。")

    temporary = REPORT.with_suffix(".phr_appendix.tmp.docx")
    document.save(temporary)
    temporary.replace(REPORT)
    print(REPORT)


if __name__ == "__main__":
    build()
