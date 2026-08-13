from __future__ import annotations

import hashlib
import sys
import unittest
import zipfile
from pathlib import Path
from unittest import mock


PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "scripts"))


class TechnicalReportIntegrationTests(unittest.TestCase):
    def test_theory_inputs_are_loaded_from_reproducible_outputs(self) -> None:
        from build_technical_report import load_theory_report_inputs

        inputs = load_theory_report_inputs()

        role_summary = inputs["role_identifiability"]["summary"]
        self.assertEqual(role_summary["role_consistent"]["row_count"], 48)
        self.assertEqual(role_summary["role_consistent"]["role_unique_rate"], 1.0)
        self.assertEqual(role_summary["role_conflict"]["row_count"], 51)
        self.assertEqual(role_summary["role_conflict"]["role_unique_rate"], 0.0)

        thresholds = inputs["selective_recognition"]["threshold_summary"]
        threshold_090 = next(row for row in thresholds if row["threshold"] == 0.9)
        self.assertAlmostEqual(threshold_090["mean"]["coverage"], 0.6087746625)
        self.assertAlmostEqual(threshold_090["mean"]["risk"], 0.1171331941)

        self.assertEqual(
            [row["setting"] for row in inputs["component_ablation"]],
            ["完整分层模型", "去除困难负样本约束", "去除层级一致性约束"],
        )

    def test_equation_specs_use_required_species_to_role_kl_direction(self) -> None:
        from build_technical_report import theory_equation_specs

        equations = theory_equation_specs()

        self.assertIn(r"\widetilde{\mathbf{p}}_r = \mathbf{A}\mathbf{p}_s", equations["aggregation"])
        self.assertIn(
            r"D_{\mathrm{KL}}\!\left(\widetilde{\mathbf{p}}_r\,\Vert\,\mathbf{p}_r\right)",
            equations["consistency"],
        )

    def test_mathtext_candidate_collection_skips_inaccessible_windows_aliases(self) -> None:
        from build_technical_report import _mathtext_python_candidates

        with mock.patch("build_technical_report.Path.exists", side_effect=OSError("inaccessible")):
            self.assertEqual(_mathtext_python_candidates(), [])

    def test_generated_report_contains_required_sections_and_figures(self) -> None:
        from docx import Document

        report_path = PROJECT_ROOT / "结题" / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（正式版）.docx"
        document = Document(report_path)
        report_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertGreater(len(document.paragraphs), 130)
        self.assertGreaterEqual(len(document.inline_shapes), 10)
        self.assertGreaterEqual(len(document.tables), 15)
        for required in (
            "矿物种类—角色分层一致性模型",
            "角色可识别性命题",
            "固定测试划分上的选择性识别",
            "角色风险收缩：种类错误不必然导致角色错误",
            "温度校准与独立风险认证",
            "摄影者留出条件下的来源外泛化",
            "outputs/theory_validation/hierarchy_consistency/",
            "outputs/theory_validation/calibrated_selective_recognition/",
            "outputs/paper_experiments_v2/training_summaries/",
        ):
            self.assertIn(required, report_text)
        for prohibited in (
            "实现工业分选",
            "提高精矿品位",
            "提高回收率",
            "XRF 成本最优策略得到验证",
            "外部验证证明",
            "具有统计显著性",
        ):
            self.assertNotIn(prohibited, report_text)

        required_figures = (
            PROJECT_ROOT / "outputs" / "paper_figures_v2" / "fig_hierarchy_kl_disagreement.png",
            PROJECT_ROOT / "outputs" / "paper_figures_v2" / "fig_calibrated_reliability.png",
            PROJECT_ROOT / "outputs" / "paper_figures_v2" / "fig_source_holdout_comparison.png",
        )
        expected_hashes = {hashlib.sha256(path.read_bytes()).hexdigest() for path in required_figures}
        with zipfile.ZipFile(report_path) as archive:
            embedded_hashes = {
                hashlib.sha256(archive.read(name)).hexdigest()
                for name in archive.namelist()
                if name.startswith("word/media/")
            }
        self.assertLessEqual(expected_hashes, embedded_hashes)


if __name__ == "__main__":
    unittest.main()
