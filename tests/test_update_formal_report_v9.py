from __future__ import annotations

import json
import hashlib
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]


class FormalReportV9EvidenceTests(unittest.TestCase):
    @staticmethod
    def _write_complete_evidence(analysis_dir: Path) -> None:
        configurations = (
            "rsg_complete",
            "cgdc_complete",
            "cgdc_shared_features",
            "cgdc_unconditional",
            "cgdc_no_decomposition_loss",
        )
        metrics = (
            "accuracy",
            "macro_f1",
            "target_recall",
            "ti_to_target_intrusion_rate",
            "metallic_to_target_intrusion_rate",
            "brier_score",
            "expected_calibration_error",
        )
        summary = {
            configuration: {
                metric: {"mean": 0.70, "sample_std": 0.01, "values": [0.69, 0.70, 0.71]}
                for metric in metrics
            }
            for configuration in configurations
        }
        (analysis_dir / "cgdc_three_seed_summary.json").write_text(
            json.dumps(summary), encoding="utf-8"
        )
        for configuration in configurations[1:]:
            (analysis_dir / f"paired_{configuration}_vs_rsg_complete.json").write_text(
                json.dumps(
                    {
                        "classification": {"macro_f1": {"difference": 0.01, "ci_low": -0.01, "ci_high": 0.02}},
                        "routing_regret": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                        "calibration": {
                            "brier_score": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                            "expected_calibration_error": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                        },
                    }
                ),
                encoding="utf-8",
            )

    @staticmethod
    def _write_complete_rpg_evidence(analysis_dir: Path) -> None:
        configurations = (
            "rsg_complete",
            "rpg_complete",
            "rpg_without_within",
            "rpg_without_between",
            "rpg_total_entropy_only",
        )
        metrics = (
            "accuracy",
            "macro_f1",
            "target_recall",
            "brier_score",
            "expected_calibration_error",
        )
        summary = {
            configuration: {
                metric: {"mean": 0.70, "sample_std": 0.01, "values": [0.69, 0.70, 0.71]}
                for metric in metrics
            }
            for configuration in configurations
        }
        (analysis_dir / "rpg_three_seed_summary.json").write_text(
            json.dumps(summary), encoding="utf-8"
        )
        for configuration in configurations[1:]:
            (analysis_dir / f"paired_{configuration}_vs_rsg_complete.json").write_text(
                json.dumps(
                    {
                        "classification": {"macro_f1": {"difference": 0.01, "ci_low": -0.01, "ci_high": 0.02}},
                        "routing_regret": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                        "calibration": {
                            "brier_score": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                            "expected_calibration_error": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                        },
                    }
                ),
                encoding="utf-8",
            )

    @staticmethod
    def _write_complete_mrpg_evidence(analysis_dir: Path) -> None:
        configurations = (
            "rsg_complete",
            "rpg_complete",
            "mrpg_complete",
            "mrpg_unconstrained_between",
            "mrpg_without_between",
        )
        metrics = (
            "accuracy",
            "macro_f1",
            "target_recall",
            "brier_score",
            "expected_calibration_error",
        )
        summary = {
            configuration: {
                metric: {"mean": 0.70, "sample_std": 0.01, "values": [0.69, 0.70, 0.71]}
                for metric in metrics
            }
            for configuration in configurations
        }
        (analysis_dir / "mrpg_three_seed_summary.json").write_text(
            json.dumps(summary), encoding="utf-8"
        )
        for configuration in configurations[2:]:
            (analysis_dir / f"paired_{configuration}_vs_rsg_complete.json").write_text(
                json.dumps(
                    {
                        "classification": {
                            "macro_f1": {"difference": 0.01, "ci_low": -0.01, "ci_high": 0.02},
                            "target_recall": {"difference": 0.02, "ci_low": -0.01, "ci_high": 0.04},
                        },
                        "routing_regret": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                        "calibration": {
                            "brier_score": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                            "expected_calibration_error": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                        },
                    }
                ),
                encoding="utf-8",
            )
        for configuration in configurations[3:]:
            (analysis_dir / f"paired_{configuration}_vs_mrpg_complete.json").write_text(
                json.dumps(
                    {
                        "classification": {
                            "macro_f1": {"difference": 0.01, "ci_low": -0.01, "ci_high": 0.02},
                            "target_recall": {"difference": 0.02, "ci_low": -0.01, "ci_high": 0.04},
                        },
                        "routing_regret": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                        "calibration": {
                            "brier_score": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                            "expected_calibration_error": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                        },
                    }
                ),
                encoding="utf-8",
            )

    @staticmethod
    def _write_complete_backbone_portability_evidence(analysis_dir: Path) -> None:
        metrics = (
            "accuracy",
            "macro_f1",
            "mean_routing_regret_nll",
        )
        summary = {
            configuration: {
                metric: {"mean": 0.70 if metric != "mean_routing_regret_nll" else 0.08,
                         "sample_std": 0.01,
                         "values": [0.69, 0.70, 0.71]}
                for metric in metrics
            }
            for configuration in ("hrgv_reference", "rsg_complete")
        }
        (analysis_dir / "rsg_three_seed_summary.json").write_text(
            json.dumps(summary), encoding="utf-8"
        )
        paired = {
            "classification": {
                "accuracy": {"difference": 0.00, "ci_low": -0.01, "ci_high": 0.01},
                "macro_f1": {"difference": 0.01, "ci_low": -0.01, "ci_high": 0.02},
            },
            "routing_regret": {"difference": -0.02, "ci_low": -0.03, "ci_high": -0.01},
        }
        (analysis_dir / "paired_rsg_complete_vs_hrgv_reference.json").write_text(
            json.dumps(paired), encoding="utf-8"
        )

    @staticmethod
    def _write_complete_backbone_replay_evidence(analysis_dir: Path) -> None:
        runs = []
        for seed in ("20260727", "20260728", "20260729"):
            runs.append(
                {
                    "protocol": "resnet50_portability",
                    "seed": f"seed{seed}",
                    "sample_count": 1284,
                    "minimum_true_probability": 1.2e-7,
                    "mean_routing_regret_nll": 0.08,
                    "mean_b1_upper_bound": 170000.0,
                    "b1_max_residual": 0.0,
                    "b1_violation_count": 0,
                    "b2_max_residual": 5.0e-8,
                    "b2_violation_count": 0,
                }
            )
        (analysis_dir / "theory_replay_summary.json").write_text(
            json.dumps(
                {
                    "runs": runs,
                    "overall": {
                        "run_count": 3,
                        "sample_count": 3852,
                        "minimum_true_probability": 1.2e-7,
                        "b1_max_residual": 0.0,
                        "b1_violation_count": 0,
                        "b2_max_residual": 5.0e-8,
                        "b2_violation_count": 0,
                    },
                    "numeric_settings": {
                        "float32_epsilon": 1.1920928955078125e-7,
                        "tolerance": 2e-6,
                    },
                }
            ),
            encoding="utf-8",
        )

    @staticmethod
    def _write_complete_gate_reliability_evidence(analysis_dir: Path) -> None:
        protocols = {}
        for name, count in (("fixed", 3852), ("photographer_holdout", 2538), ("resnet50_portability", 3852)):
            protocols[name] = {
                "run_count": 3,
                "sample_count": count,
                "mean_routing_regret_nll": 0.08,
                "mean_b1_local_bound": 0.25,
                "mean_oracle_margin": 0.22,
                "mean_soft_hard_deviation": 0.36,
                "mean_b2_bound": 0.68,
                "mean_exact_decomposition_abs_residual": 4e-8,
                "exact_decomposition_max_abs_residual": 8e-7,
                "exact_decomposition_violation_count": 0,
                "b1_local_max_residual": 0.0,
                "b1_local_violation_count": 0,
                "b2_max_residual": 5e-8,
                "b2_violation_count": 0,
            }
        (analysis_dir / "gate_reliability_summary.json").write_text(
            json.dumps(
                {
                    "protocols": protocols,
                    "overall": {
                        "run_count": 9,
                        "sample_count": 10242,
                        "mean_exact_decomposition_abs_residual": 4e-8,
                        "exact_decomposition_max_abs_residual": 8e-7,
                        "exact_decomposition_violation_count": 0,
                        "b1_local_max_residual": 0.0,
                        "b1_local_violation_count": 0,
                        "b2_max_residual": 5e-8,
                        "b2_violation_count": 0,
                    },
                    "numeric_settings": {
                        "float32_epsilon": 1.1920928955078125e-7,
                        "tolerance": 2e-6,
                        "strata_count": 3,
                    },
                    "claim_boundary": "mechanism diagnosis only",
                }
            ),
            encoding="utf-8",
        )

    @staticmethod
    def _write_complete_five_seed_evidence(analysis_dir: Path) -> None:
        metrics = (
            "accuracy",
            "macro_f1",
            "target_recall",
            "ti_to_target_intrusion_rate",
            "metallic_to_target_intrusion_rate",
            "brier_score",
            "expected_calibration_error",
        )
        summary = {
            configuration: {
                metric: {
                    "mean": 0.70,
                    "sample_std": 0.01,
                    "values": [0.68, 0.69, 0.70, 0.71, 0.72],
                }
                for metric in metrics
            }
            for configuration in ("rsg_complete", "mrpg_complete")
        }
        (analysis_dir / "five_seed_summary.json").write_text(
            json.dumps(summary), encoding="utf-8"
        )
        (analysis_dir / "paired_mrpg_complete_vs_rsg_complete.json").write_text(
            json.dumps(
                {
                    "classification": {
                        "macro_f1": {"difference": 0.01, "ci_low": -0.01, "ci_high": 0.02},
                        "target_recall": {"difference": 0.02, "ci_low": -0.01, "ci_high": 0.04},
                        "ti_to_target_intrusion": {"difference": 0.01, "ci_low": -0.02, "ci_high": 0.03},
                        "metallic_to_target_intrusion": {"difference": -0.01, "ci_low": -0.04, "ci_high": 0.02},
                    },
                    "calibration": {
                        "brier_score": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                        "expected_calibration_error": {"difference": -0.01, "ci_low": -0.02, "ci_high": 0.00},
                    },
                }
            ),
            encoding="utf-8",
        )
    def test_requires_all_formal_configurations_and_three_seeds(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import load_formal_cgdc_evidence

        with tempfile.TemporaryDirectory() as temp_dir:
            analysis_dir = Path(temp_dir)
            configurations = (
                "rsg_complete",
                "cgdc_complete",
                "cgdc_shared_features",
                "cgdc_unconditional",
                "cgdc_no_decomposition_loss",
            )
            summary = {
                configuration: {"macro_f1": {"values": [0.70, 0.71]}}
                for configuration in configurations
            }
            (analysis_dir / "cgdc_three_seed_summary.json").write_text(
                json.dumps(summary), encoding="utf-8"
            )
            with self.assertRaisesRegex(ValueError, "three registered seeds"):
                load_formal_cgdc_evidence(analysis_dir)

    def test_appends_cgdc_appendix_once_with_architecture_and_boundaries(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            analysis_dir = temp_root / "analysis"
            analysis_dir.mkdir()
            self._write_complete_evidence(analysis_dir)
            document = Document()
            document.add_paragraph("existing report")
            document.save(source)

            update_report(source, output, analysis_dir)
            update_report(output, output, analysis_dir)

            rendered = Document(output)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertEqual(text.count("附录 D CGDC-RSG-HRGV 网络理论与实验"), 1)
            self.assertIn("命题 P1", text)
            self.assertIn("命题 P4", text)
            self.assertIn("≤0.5", text)
            self.assertIn("Brier", text)
            self.assertIn("不等同于工业分选", text)

    def test_embeds_rendered_cgdc_formula_bundle(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import FORMULA_DIR, update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            analysis_dir = temp_root / "analysis"
            analysis_dir.mkdir()
            self._write_complete_evidence(analysis_dir)
            document = Document()
            document.save(source)

            update_report(source, output, analysis_dir)

            expected = {
                hashlib.sha256(path.read_bytes()).hexdigest()
                for path in FORMULA_DIR.glob("cgdc_*.png")
            }
            self.assertGreaterEqual(len(expected), 3)
            with zipfile.ZipFile(output) as archive:
                embedded = {
                    hashlib.sha256(archive.read(name)).hexdigest()
                    for name in archive.namelist()
                    if name.startswith("word/media/")
                }
            self.assertLessEqual(expected, embedded)

    def test_appends_rpg_appendix_after_cgdc_without_duplication(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            cgdc_dir = temp_root / "cgdc"
            rpg_dir = temp_root / "rpg"
            cgdc_dir.mkdir()
            rpg_dir.mkdir()
            self._write_complete_evidence(cgdc_dir)
            self._write_complete_rpg_evidence(rpg_dir)
            Document().save(source)

            update_report(source, output, cgdc_dir, rpg_dir)
            update_report(output, output, cgdc_dir, rpg_dir)

            rendered = Document(output)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertEqual(text.count("附录 D CGDC-RSG-HRGV 网络理论与实验"), 1)
            self.assertEqual(text.count("附录 E RPG-HRGV 角色分区不确定性门控理论与实验"), 1)
            self.assertIn("H(S)=H(R)+H(S|R)", text)
            self.assertIn("角色间不确定性", text)
            self.assertIn("角色内不确定性", text)
            self.assertIn("不等同于工业分选", text)

    def test_appends_mrpg_appendix_with_normalization_and_monotonicity(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            cgdc_dir = temp_root / "cgdc"
            mrpg_dir = temp_root / "mrpg"
            cgdc_dir.mkdir()
            mrpg_dir.mkdir()
            self._write_complete_evidence(cgdc_dir)
            self._write_complete_mrpg_evidence(mrpg_dir)
            Document().save(source)

            update_report(source, output, cgdc_dir, mrpg_analysis_dir=mrpg_dir)
            update_report(output, output, cgdc_dir, mrpg_analysis_dir=mrpg_dir)

            rendered = Document(output)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertEqual(text.count("附录 F M-RPG-HRGV 容量归一化单调门控理论与实验"), 1)
            self.assertIn("容量归一化", text)
            self.assertIn("单调直接专家分配", text)
            self.assertIn("命题 M-R3", text)
            self.assertIn("相对完整 M-RPG", text)
            self.assertIn("不等同于工业分选", text)

    def test_appends_five_seed_extension_appendix_once(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            cgdc_dir = temp_root / "cgdc"
            extension_dir = temp_root / "extension"
            cgdc_dir.mkdir()
            extension_dir.mkdir()
            self._write_complete_evidence(cgdc_dir)
            self._write_complete_five_seed_evidence(extension_dir)
            Document().save(source)

            update_report(source, output, cgdc_dir, five_seed_analysis_dir=extension_dir)
            update_report(output, output, cgdc_dir, five_seed_analysis_dir=extension_dir)

            rendered = Document(output)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertEqual(text.count("附录 G M-RPG-HRGV 五随机种子扩展验证"), 1)
            self.assertIn("预注册的两随机种子扩展", text)
            self.assertIn("未观察到 M-RPG 相对 RSG 的稳定经验优势", text)
            self.assertIn("目标类召回", text)

    def test_appends_rsg_theory_evidence_figure_once_with_claim_boundary(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            cgdc_dir = temp_root / "cgdc"
            cgdc_dir.mkdir()
            self._write_complete_evidence(cgdc_dir)
            Document().save(source)

            update_report(source, output, cgdc_dir)
            update_report(output, output, cgdc_dir)

            rendered = Document(output)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertEqual(text.count("附录 H RSG-HRGV 理论性质与证据对应图"), 1)
            self.assertIn("平均路由后悔", text)
            self.assertIn("不主张总体分类性能优越", text)
            self.assertIn("高精度检查点重放", text)
            self.assertIn("数值一致性验证", text)
            self.assertGreaterEqual(len(rendered.inline_shapes), 1)

    def test_updates_main_contribution_with_rsg_theory_and_mrpg_boundary(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import PRIMARY_CONTRIBUTION_OLD, update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            cgdc_dir = temp_root / "cgdc"
            cgdc_dir.mkdir()
            self._write_complete_evidence(cgdc_dir)
            document = Document()
            document.add_paragraph(PRIMARY_CONTRIBUTION_OLD)
            document.save(source)

            update_report(source, output, cgdc_dir)
            update_report(output, output, cgdc_dir)

            text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
            self.assertNotIn(PRIMARY_CONTRIBUTION_OLD, text)
            self.assertIn("RSG-HRGV", text)
            self.assertIn("定理 B.1--B.3", text)
            self.assertIn("主干替换不变性", text)
            self.assertIn("不构成总体分类性能优势", text)

    def test_appends_backbone_portability_evidence_with_claim_boundary(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import RSG_THEORY_EVIDENCE_FIGURE_PATH, update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            cgdc_dir = temp_root / "cgdc"
            portability_dir = temp_root / "portability"
            cgdc_dir.mkdir()
            portability_dir.mkdir()
            self._write_complete_evidence(cgdc_dir)
            self._write_complete_backbone_portability_evidence(portability_dir)
            Document().save(source)

            update_report(
                source,
                output,
                cgdc_dir,
                backbone_portability_analysis_dir=portability_dir,
                backbone_portability_figure_path=RSG_THEORY_EVIDENCE_FIGURE_PATH,
            )
            update_report(
                output,
                output,
                cgdc_dir,
                backbone_portability_analysis_dir=portability_dir,
                backbone_portability_figure_path=RSG_THEORY_EVIDENCE_FIGURE_PATH,
            )

            rendered = Document(output)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertEqual(text.count("H.3 主干替换不变性与 ResNet50 跨主干确认"), 1)
            self.assertIn("\u878d\u5408\u5c42", text)
            self.assertIn("不等同于主干性能优越", text)
            self.assertIn("平均路由后悔", text)

    def test_appends_backbone_replay_consistency_once(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            cgdc_dir = temp_root / "cgdc"
            replay_dir = temp_root / "replay"
            cgdc_dir.mkdir()
            replay_dir.mkdir()
            self._write_complete_evidence(cgdc_dir)
            self._write_complete_backbone_replay_evidence(replay_dir)
            Document().save(source)

            update_report(source, output, cgdc_dir, backbone_replay_analysis_dir=replay_dir)
            update_report(output, output, cgdc_dir, backbone_replay_analysis_dir=replay_dir)

            text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
            self.assertEqual(text.count("H.4 ResNet50 跨主干高精度重放"), 1)
            self.assertIn("3 次、3852 张图像", text)
            self.assertIn("B.1 违反数为 0", text)
            self.assertIn("不构成新的分类性能实验", text)

    def test_appends_gate_reliability_diagnosis_once(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import RSG_THEORY_EVIDENCE_FIGURE_PATH, update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            cgdc_dir = temp_root / "cgdc"
            reliability_dir = temp_root / "reliability"
            cgdc_dir.mkdir()
            reliability_dir.mkdir()
            self._write_complete_evidence(cgdc_dir)
            self._write_complete_gate_reliability_evidence(reliability_dir)
            Document().save(source)

            update_report(
                source,
                output,
                cgdc_dir,
                gate_reliability_analysis_dir=reliability_dir,
                gate_reliability_figure_path=RSG_THEORY_EVIDENCE_FIGURE_PATH,
            )
            update_report(
                output,
                output,
                cgdc_dir,
                gate_reliability_analysis_dir=reliability_dir,
                gate_reliability_figure_path=RSG_THEORY_EVIDENCE_FIGURE_PATH,
            )

            text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
            self.assertEqual(text.count("H.5 门控可靠性分层诊断"), 1)
            self.assertEqual(text.count("H.6 凸融合路由后悔精确分解的数值验证"), 1)
            self.assertIn("epsilon_i=min", text)
            self.assertIn("10,242", text)
            self.assertIn("机制诊断", text)
            self.assertIn("8.00e-07", text)

    def test_appends_exact_regret_curve_once_with_formula_boundary(self) -> None:
        import sys

        sys.path.insert(0, str(PROJECT_ROOT / "scripts"))
        from update_formal_report_v9 import RSG_THEORY_EVIDENCE_FIGURE_PATH, update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source = temp_root / "source.docx"
            output = temp_root / "output.docx"
            cgdc_dir = temp_root / "cgdc"
            reliability_dir = temp_root / "reliability"
            cgdc_dir.mkdir()
            reliability_dir.mkdir()
            self._write_complete_evidence(cgdc_dir)
            self._write_complete_gate_reliability_evidence(reliability_dir)
            Document().save(source)

            update_report(
                source,
                output,
                cgdc_dir,
                gate_reliability_analysis_dir=reliability_dir,
                exact_decomposition_figure_path=RSG_THEORY_EVIDENCE_FIGURE_PATH,
            )
            update_report(
                output,
                output,
                cgdc_dir,
                gate_reliability_analysis_dir=reliability_dir,
                exact_decomposition_figure_path=RSG_THEORY_EVIDENCE_FIGURE_PATH,
            )

            rendered = Document(output)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertEqual(text.count("H.7 凸融合路由后悔的理论曲线"), 1)
            self.assertIn("r=-log(1-delta d/M)", text)
            self.assertIn("解析理论曲线", text)
            self.assertIn("不构成新的分类性能实验", text)
            self.assertGreaterEqual(len(rendered.inline_shapes), 2)


if __name__ == "__main__":
    unittest.main()
