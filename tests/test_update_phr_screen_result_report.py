from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "tools"))


class PHRScreenResultReportTests(unittest.TestCase):
    def test_appends_nonpromotion_result_once(self) -> None:
        from append_phr_screen_result_to_official_report import append_screen_result

        decision = {
            "screen_seed": 20260728,
            "promote_to_formal": False,
            "criterion_ids": [],
            "evidence": {
                "runs": {
                    "rsg_reference/seed20260728": {
                        "metrics": {
                            "macro_f1": 0.7485,
                            "accuracy": 0.7664,
                            "target_recall": 0.7360,
                            "ti_to_target_intrusion_rate": 0.0918,
                            "metallic_to_target_intrusion_rate": 0.1130,
                            "phr_ti_mean_margin_regret": 0.4187,
                            "phr_metallic_mean_margin_regret": 0.6374,
                        }
                    },
                    "phr_complete/seed20260728": {
                        "metrics": {
                            "macro_f1": 0.7414,
                            "accuracy": 0.7570,
                            "target_recall": 0.6880,
                            "ti_to_target_intrusion_rate": 0.0959,
                            "metallic_to_target_intrusion_rate": 0.1243,
                            "phr_ti_mean_margin_regret": 0.1813,
                            "phr_metallic_mean_margin_regret": 0.2350,
                        }
                    },
                }
            },
        }
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "report.docx"
            document = Document()
            document.add_paragraph("existing report")
            document.save(path)

            self.assertTrue(append_screen_result(path, decision))
            self.assertFalse(append_screen_result(path, decision))

            report_text = "\n".join(paragraph.text for paragraph in Document(path).paragraphs)
            table_text = "\n".join(
                cell.text for table in Document(path).tables for row in table.rows for cell in row.cells
            )
            self.assertEqual(report_text.count("I.6 PHR 验证集筛选结果与边界结论"), 1)
            self.assertIn("未晋级", report_text)
            self.assertIn("0.7485", table_text)
            self.assertIn("0.7414", table_text)


if __name__ == "__main__":
    unittest.main()
