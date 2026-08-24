from __future__ import annotations

import hashlib
import re
import unittest
import zipfile
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT = PROJECT_ROOT / "结题" / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（正式版）.docx"
PAIRED_FIGURE = PROJECT_ROOT / "outputs" / "paper_figures_v3" / "fig_paired_cluster_effects.png"


class FormalReportV3Tests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.document = Document(REPORT)
        cls.text = "\n".join(paragraph.text for paragraph in cls.document.paragraphs)

    def test_paired_inference_section_and_evidence_are_present(self) -> None:
        required = (
            "5.16 成对簇 Bootstrap 与配对统计推断",
            "5.19 实验结果的综合判断",
            "目标召回提高9.16个百分点",
            "95%簇 Bootstrap 区间为[5.18, 13.28]个百分点",
            "三组经 Holm 校正的 p 值均大于0.05",
            "式（11） 两阶段成对簇 Bootstrap 的模型效应估计",
            "表 24 固定测试集上的成对统计推断",
            "图 16 分层模型相对基线的成对效应及95%簇 Bootstrap 区间",
        )
        for item in required:
            self.assertIn(item, self.text)

    def test_new_figure_is_embedded_and_captions_remain_unique(self) -> None:
        expected_hash = hashlib.sha256(PAIRED_FIGURE.read_bytes()).hexdigest()
        with zipfile.ZipFile(REPORT) as archive:
            embedded_hashes = {
                hashlib.sha256(archive.read(name)).hexdigest()
                for name in archive.namelist()
                if name.startswith("word/media/")
            }
        self.assertIn(expected_hash, embedded_hashes)

        identifiers = []
        for paragraph in self.document.paragraphs:
            if paragraph.style.name != "Caption":
                continue
            match = re.match(r"^(图|表)\s+([0-9]+)\b", paragraph.text.strip())
            if match:
                identifiers.append(match.group(1) + match.group(2))
        self.assertEqual(len(identifiers), len(set(identifiers)))

    def test_claims_remain_within_the_evidence_boundary(self) -> None:
        prohibited = (
            "分层模型总体显著优于基线",
            "已证明工业分选有效",
            "实现精矿品位预测",
        )
        for item in prohibited:
            self.assertNotIn(item, self.text)


if __name__ == "__main__":
    unittest.main()
