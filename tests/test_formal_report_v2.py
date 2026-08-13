from __future__ import annotations

import re
import unittest
import zipfile
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT = PROJECT_ROOT / "结题" / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（正式版）.docx"


class FormalReportV2Tests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.document = Document(REPORT)

    def test_v2_sections_and_evidence_are_present(self) -> None:
        text = "\n".join(paragraph.text for paragraph in self.document.paragraphs)
        required = (
            "5.11 角色风险收缩：种类错误不必然导致角色错误",
            "5.12 双头分布一致性与 Pinsker 上界",
            "5.13 温度校准与独立风险认证",
            "5.14 摄影者留出条件下的来源外泛化",
            "5.15 普通磁铁矿代理样本消融",
            "18.74±1.29个百分点",
            "最终数据集保留23张",
        )
        for item in required:
            self.assertIn(item, text)

    def test_figure_and_table_numbers_are_unique(self) -> None:
        identifiers = []
        for paragraph in self.document.paragraphs:
            if paragraph.style.name != "Caption":
                continue
            match = re.match(r"^(图|表)\s+([0-9]+)\b", paragraph.text.strip())
            if match:
                identifiers.append(match.group(1) + match.group(2))
        self.assertEqual(len(identifiers), len(set(identifiers)))

    def test_images_fit_inside_page_width(self) -> None:
        section = self.document.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        self.assertGreaterEqual(len(self.document.inline_shapes), 25)
        for shape in self.document.inline_shapes:
            self.assertLessEqual(shape.width, usable_width)

    def test_docx_package_relationship_targets_exist(self) -> None:
        with zipfile.ZipFile(REPORT) as archive:
            names = set(archive.namelist())
            self.assertIn("word/document.xml", names)
            self.assertIn("word/_rels/document.xml.rels", names)
            self.assertGreaterEqual(sum(name.startswith("word/media/") for name in names), 25)


if __name__ == "__main__":
    unittest.main()
