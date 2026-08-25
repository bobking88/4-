from __future__ import annotations

import unittest
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT = PROJECT_ROOT / "结题" / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（正式版）.docx"


class FormalReportV6Tests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        document = Document(REPORT)
        cls.text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        )

    def test_source_holdout_rsg_evidence_and_boundaries_are_present(self) -> None:
        for expected in (
            "5.19 RSG-HRGV-Net 的摄影者留出泛化验证",
            "表 27 摄影者留出下 RSG-HRGV-Net 与 HRGV 的独立三随机种子比较",
            "图 18 摄影者留出条件下 RSG-HRGV 的角色指标与路由机制比较",
            "95%成对簇 Bootstrap 区间为[-4.75, -2.17]个百分点",
            "不能声称 RSG 在来源变化下稳定提高总体角色分类",
            "不将其替代真实颗粒、矿石品位或工业回收率验证",
        ):
            self.assertIn(expected, self.text)


if __name__ == "__main__":
    unittest.main()
