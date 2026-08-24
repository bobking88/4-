from __future__ import annotations

import hashlib
import re
import unittest
import zipfile
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT = PROJECT_ROOT / "结题" / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（正式版）.docx"
ARCHITECTURE = PROJECT_ROOT / "outputs" / "paper_figures_v2" / "fig_hrgv_architecture.png"


class FormalReportV5Tests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.document = Document(REPORT)
        paragraphs = [paragraph.text for paragraph in cls.document.paragraphs]
        cells = [cell.text for table in cls.document.tables for row in table.rows for cell in row.cells]
        cls.text = "\n".join(paragraphs + cells)

    def test_rsg_theory_and_formal_evidence_are_present(self) -> None:
        required = (
            "4.6.1 后悔监督门控（RSG）",
            "命题6（门控误差的路由后悔上界）",
            "命题7（软目标对硬门控的指数逼近）",
            "命题8（后悔监督的局部梯度隔离）",
            "5.18 RSG-HRGV-Net 后悔监督门控正式验证",
            "表 26 RSG-HRGV-Net 与 HRGV 的三随机种子正式比较",
            "图 5 RSG-HRGV-Net 的分层证据、后悔监督门控与困难负样本风险验证结构",
            "45.99% ± 4.82%",
            "58.01% ± 1.76%",
            "95%簇 Bootstrap 区间为[-2.86, -0.69]个百分点",
        )
        for item in required:
            self.assertIn(item, self.text)

    def test_rsg_boundary_and_caption_ids_are_valid(self) -> None:
        self.assertIn("总体 Accuracy 和 Macro F1 的差异区间均跨0", self.text)
        self.assertIn("不将其解释为分类精度的普适显著提升", self.text)
        identifiers = []
        for paragraph in self.document.paragraphs:
            if paragraph.style.name != "Caption":
                continue
            match = re.match(r"^(图|表)\s+([0-9]+)\b", paragraph.text.strip())
            if match:
                identifiers.append(match.group(1) + match.group(2))
        self.assertEqual(len(identifiers), len(set(identifiers)))

    def test_updated_architecture_is_embedded(self) -> None:
        expected_hash = hashlib.sha256(ARCHITECTURE.read_bytes()).hexdigest()
        with zipfile.ZipFile(REPORT) as archive:
            embedded_hashes = {
                hashlib.sha256(archive.read(name)).hexdigest()
                for name in archive.namelist()
                if name.startswith("word/media/")
            }
        self.assertIn(expected_hash, embedded_hashes)


if __name__ == "__main__":
    unittest.main()
