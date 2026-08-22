from __future__ import annotations

import hashlib
import re
import unittest
import zipfile
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT = PROJECT_ROOT / "结题" / "基于深度学习的钒钛矿相关矿物图像识别方法研究_技术报告（正式版）.docx"
ARCHITECTURE = PROJECT_ROOT / "outputs" / "paper_figures_v2" / "fig_hrgv_architecture.png"
EFFECTS = PROJECT_ROOT / "outputs" / "paper_figures_v2" / "fig_hrgv_vs_hierarchical_effects.png"


class FormalReportV4Tests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.document = Document(REPORT)
        paragraph_text = [paragraph.text for paragraph in cls.document.paragraphs]
        table_text = [cell.text for table in cls.document.tables for row in table.rows for cell in row.cells]
        cls.text = "\n".join(paragraph_text + table_text)

    def test_hrgv_method_theory_and_results_are_present(self) -> None:
        required = (
            "4.6 层级风险门控验证网络（HRGV-Net）",
            "5.17 HRGV-Net 三随机种子实验与组件验证",
            "5.18 实验结果的综合判断",
            "命题1（凸融合对数损失上界）",
            "命题2（中性区残差校正的恒等性与单调性）",
            "命题3（验证式目标接纳的误收上界）",
            "命题4（带拒识的贝叶斯最优选择性决策）",
            "命题5（验证器梯度隔离）",
            "76.48% ± 0.23%",
            "74.59% ± 0.34%",
            "Accuracy 提高1.09个百分点",
            "95%区间为[0.08, 2.15]个百分点",
            "不能解释为总体性能全面显著提高",
            "图 5 HRGV-Net 的层级证据融合与困难负样本风险验证结构",
            "图 17 HRGV-Net 相对既有分层模型的成对效应及95%簇 Bootstrap 区间",
            "表 25 HRGV-Net 三随机种子组件消融",
        )
        for item in required:
            self.assertIn(item, self.text)

    def test_hrgv_figures_are_embedded_and_caption_ids_are_unique(self) -> None:
        expected_hashes = {
            hashlib.sha256(ARCHITECTURE.read_bytes()).hexdigest(),
            hashlib.sha256(EFFECTS.read_bytes()).hexdigest(),
        }
        with zipfile.ZipFile(REPORT) as archive:
            embedded_hashes = {
                hashlib.sha256(archive.read(name)).hexdigest()
                for name in archive.namelist()
                if name.startswith("word/media/")
            }
        self.assertLessEqual(expected_hashes, embedded_hashes)

        identifiers = []
        for paragraph in self.document.paragraphs:
            if paragraph.style.name != "Caption":
                continue
            match = re.match(r"^(图|表)\s+([0-9]+)\b", paragraph.text.strip())
            if match:
                identifiers.append(match.group(1) + match.group(2))
        self.assertEqual(len(identifiers), len(set(identifiers)))

    def test_theory_claims_are_not_overstated(self) -> None:
        required_boundaries = (
            "命题本身是计算图上的确定性性质",
            "阶段条件化选矿决策图仍保留为后续研究方向",
            "不等同于工业分选回收率验证",
        )
        for item in required_boundaries:
            self.assertIn(item, self.text)

        prohibited = (
            "HRGV-Net全面显著优于所有基线",
            "首次提出一般层级分类理论",
            "已验证工业回收率提升",
            "已实现精矿品位预测",
        )
        for item in prohibited:
            self.assertNotIn(item, self.text)


if __name__ == "__main__":
    unittest.main()
