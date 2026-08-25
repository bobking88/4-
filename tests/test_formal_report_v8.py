from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "scripts"))


class FormalReportV8Tests(unittest.TestCase):
    def test_appends_theory_ablation_appendix_once(self) -> None:
        from update_formal_report_v8 import update_report

        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "source.docx"
            output = Path(temp_dir) / "output.docx"
            document = Document()
            document.add_paragraph("existing report")
            document.save(source)

            update_report(source, output)
            update_report(output, output)

            rendered = Document(output)
            headings = [paragraph.text for paragraph in rendered.paragraphs]
            self.assertEqual(headings.count("附录 C RSG-HRGV 三随机种子理论消融"), 1)
            self.assertTrue(any("硬目标" in text for text in headings))


if __name__ == "__main__":
    unittest.main()
